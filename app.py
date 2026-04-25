import streamlit as st
import trafilatura
from youtube_transcript_api import YouTubeTranscriptApi
from urllib.parse import urlparse, parse_qs
import io
from pathlib import Path
import os
import html as html_lib
from docx import Document
from openai import OpenAI
import httpx
import requests
import json
from bs4 import BeautifulSoup
import re
import streamlit.components.v1 as components
import base64
import hashlib
import textwrap
from datetime import datetime
import time
import ctypes
try:
    import pandas as pd
except Exception:
    pd = None

from podcast_audio import (
    DEFAULT_TTS_MODEL,
    DEFAULT_TTS_VOICE,
    VOICE_LABELS as PODCAST_VOICE_LABELS,
    PodcastAudioError,
    synthesize_podcast,
)
from podcast_script import get_podcast_sys_prompt, normalize_podcast_segments, parse_podcast_script_segments


# ==========================================
# 0. 提示词与草稿持久化管理 (JSON 存储)
# ==========================================
APP_DIR = os.path.dirname(os.path.abspath(__file__))
RUNTIME_AUDIO_DIR = os.path.join(APP_DIR, "runtime_audio")
PODCAST_OUTPUT_DIR = os.path.join(RUNTIME_AUDIO_DIR, "podcasts")
PODCAST_CACHE_DIR = os.path.join(RUNTIME_AUDIO_DIR, "tts_cache")

PROMPTS_PATH_CANDIDATES = [
    os.path.join(APP_DIR, "prompts.json"),
    os.path.join(APP_DIR, "prompt.json"),
    os.path.join(APP_DIR, "Prompt.json"),
    "prompts.json",
    "prompt.json",
    "Prompt.json",
]

PROMPTS_FILE = PROMPTS_PATH_CANDIDATES[0]
DRAFT_FILE = os.path.join(APP_DIR, "draft_state.json")
TASK_QUEUE_FILE = os.path.join(APP_DIR, "task_queue_state.json")
AI_DIAGNOSTIC_LOG = os.path.join(APP_DIR, "ai_diagnostics.log")
PROMPTS_LOAD_REPORT = ""
PENDING_DRAFT_RESTORE_KEY = "_pending_draft_restore"
DRAFT_RESTORE_NOTICE_KEY = "_draft_restore_notice"
TASK_QUEUE_NOTICE_KEY = "_task_queue_notice"

DRAFT_STATE_KEYS = [
    'current_step', 'article_url', 'video_url', 'source_content',
    'source_images', 'extraction_success', 'draft_article',
    'review_feedback', 'review_actions', 'accepted_review_items', 'modified_article', 'final_article', 'title_candidates', 'highlighted_article', 'spoken_script',
    'podcast_enabled', 'podcast_duration', 'podcast_script_raw', 'podcast_script_segments',
    'podcast_audio_path', 'podcast_audio_manifest', 'podcast_last_error', 'podcast_voice',
    'podcast_tts_provider', 'podcast_tts_api_key_present',
    'chat_history', 'image_keywords', 'selected_role', 'target_article_words',
    'source_images_all', 'selected_source_image_ids', 'article_versions',
    'active_article_version_id', 'de_ai_model', 'de_ai_variant', 'de_ai_temperature',
    'de_ai_prompt_template', 'term_rules_enabled', 'term_rules_scope',
    'banned_terms_text', 'replacement_terms_text', 'default_replacement_terms_text',
    'suggested_replacement_terms_text', 'article_banned_terms_text', 'article_replacement_terms_text',
    'article_default_replacement_terms_text', 'article_suggested_replacement_terms_text',
    'term_scan_result', 'term_scan_summary',
    'pending_ai_stage', 'last_completed_ai_stage',
    'last_completed_ai_target_step', 'last_ai_error', 'recovered_ai_notice',
    'obsidian_enabled', 'obsidian_vault_path', 'obsidian_max_hits',
    'obsidian_show_hits', 'obsidian_hits', 'obsidian_research_brief',
    'obsidian_retrieval_error', 'obsidian_query_terms', 'obsidian_wiki_root',
    'obsidian_last_indexed_at', 'obsidian_retrieval_signature',
    'obsidian_influence_map', 'obsidian_influence_summary', 'obsidian_influence_signature',
    'evidence_map', 'evidence_summary', 'evidence_signature'
]

TASK_TEMPLATE_CONFIG_KEYS = [
    'selected_role', 'target_article_words',
    'de_ai_model', 'de_ai_variant', 'de_ai_temperature',
    'term_rules_enabled', 'term_rules_scope',
    'banned_terms_text', 'default_replacement_terms_text', 'suggested_replacement_terms_text',
    'article_banned_terms_text', 'article_default_replacement_terms_text', 'article_suggested_replacement_terms_text',
    'podcast_enabled', 'podcast_duration', 'podcast_voice',
    'obsidian_enabled', 'obsidian_vault_path', 'obsidian_max_hits', 'obsidian_show_hits',
]

TASK_STATUS_LABELS = {
    "pending": "待处理",
    "in_progress": "处理中",
    "needs_review": "待复核",
    "completed": "已完成",
    "failed": "失败",
}



def get_prompt_file_candidates():
    candidates = []
    env_prompt_path = os.environ.get("PROMPTS_FILE", "").strip()
    if env_prompt_path:
        candidates.append(env_prompt_path)

    candidates.extend(PROMPTS_PATH_CANDIDATES)

    roots = [
        APP_DIR,
        os.getcwd(),
        os.path.dirname(APP_DIR),
        os.path.dirname(os.path.dirname(APP_DIR)),
    ]
    for root in roots:
        if not root:
            continue
        for name in ("prompts.json", "prompt.json", "Prompt.json"):
            candidates.append(os.path.join(root, name))

    try:
        app_depth = APP_DIR.rstrip(os.sep).count(os.sep)
        for walk_root, _, files in os.walk(APP_DIR):
            current_depth = walk_root.rstrip(os.sep).count(os.sep) - app_depth
            if current_depth > 2:
                continue
            for file_name in files:
                if file_name.lower() in {"prompts.json", "prompt.json"}:
                    candidates.append(os.path.join(walk_root, file_name))
    except Exception:
        pass

    unique_candidates = []
    seen = set()
    for path in candidates:
        normalized = os.path.abspath(path)
        if normalized not in seen:
            seen.add(normalized)
            unique_candidates.append(normalized)
    return unique_candidates


def build_exception_chain(exc):
    chain = []
    current = exc
    seen = set()
    while current and id(current) not in seen:
        seen.add(id(current))
        label = type(current).__name__
        message = str(current).strip()
        chain.append(f"{label}: {message}" if message else label)
        current = getattr(current, "__cause__", None) or getattr(current, "__context__", None)
    return chain


def format_llm_exception(exc):
    chain = build_exception_chain(exc)
    if not chain:
        return "Unknown error."
    if len(chain) == 1:
        return chain[0]
    return " -> ".join(chain)


def build_proxy_diagnostic_snapshot():
    env_proxy_map = {}
    for env_name in ("HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY", "NO_PROXY"):
        env_value = os.environ.get(env_name, "").strip()
        if env_value:
            env_proxy_map[env_name] = env_value

    return {
        "http_proxy": env_proxy_map.get("HTTP_PROXY", ""),
        "https_proxy": env_proxy_map.get("HTTPS_PROXY", ""),
        "all_proxy": env_proxy_map.get("ALL_PROXY", ""),
        "no_proxy": env_proxy_map.get("NO_PROXY", ""),
        "env_proxy_keys": sorted(env_proxy_map.keys()),
        "trust_env_disabled": True,
    }


def write_ai_diagnostic_log(stage_name, base_url, model_name, error_message, user_content, image_count=0, history_count=0, extra_details=None):
    payload = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "stage": stage_name or "",
        "base_url": base_url or "",
        "model": model_name or "",
        "error": error_message,
        "image_count": image_count,
        "history_count": history_count,
        "user_content_chars": len(user_content or ""),
    }
    if extra_details:
        payload.update(extra_details)
    try:
        with open(AI_DIAGNOSTIC_LOG, "a", encoding="utf-8") as f:
            f.write(json.dumps(payload, ensure_ascii=False) + "\n")
    except Exception:
        pass

DEFAULT_GLOBAL_PROMPT = """【全局强制写作规范（最高优先级）】
1. 切断 AI 八股句式：坚决禁用“不是……而是”、“不仅……甚至”、“总而言之”、“在这个瞬息万变的时代”、“正如前文所述”等强烈的机械感过渡句和排比句。
2. 禁用伪高级“黑话”：严禁滥用带双引号的互联网/营销词汇（如“赋能”、“底层逻辑”、“打法”、“组合拳”、“降维打击”）。遇到专业概念，请用人话直白解释，不要故作高深。
3. 拒绝强行升华：文章结尾禁止进行“爹味说教”或喊口号式的价值升华，客观给出冷酷的结论或留白即可。
4. 打破匀速节奏：多用短句！避免一口气读不完的复杂长句。允许出现少量口语化的标点停顿，模仿人类写稿时真实的“呼吸感”和偶尔的“毒舌”感。"""

def load_prompts():
    default_data = {
        "editors": {
            "发行主编": "你是一位资深的海外发行主编。请深度分析素材，重点关注买量、ROI与发行策略...",
            "研发主编": "你是一位硬核游戏制作人。请深度拆解素材，重点关注核心循环、系统设计与工业化管线...",
            "游戏快讯编辑": "你是一位敏锐的游戏媒体编辑。请将素材提炼为通俗易懂、具有爆点的新闻快讯...",
            "客观转录编辑": "你是一位专业速记员。请剥离所有主观情绪，将素材客观、结构化地转录并总结..."
        },
        "reviewer": "你是一个极其严苛的资深游戏媒体主编兼风控专家。请严格核查初稿中的事实错误、逻辑漏洞及AI幻觉...",
        "global_instruction": DEFAULT_GLOBAL_PROMPT
    }

    global PROMPTS_FILE
    global PROMPTS_LOAD_REPORT

    parse_errors = []
    for candidate in get_prompt_file_candidates():
        if not os.path.exists(candidate):
            continue
        try:
            with open(candidate, "r", encoding="utf-8-sig") as f:
                data = json.load(f)

            editors = data.get("editors", {}) if isinstance(data, dict) else {}
            if not isinstance(editors, dict) or len(editors) == 0:
                raise ValueError("缺少 editors 配置或 editors 为空")

            if "global_instruction" not in data:
                data["global_instruction"] = DEFAULT_GLOBAL_PROMPT

            PROMPTS_FILE = candidate
            PROMPTS_LOAD_REPORT = f"loaded:{PROMPTS_FILE}:editors={len(editors)}"
            return data
        except Exception as e:
            parse_errors.append(f"{candidate} -> {str(e)}")

    PROMPTS_FILE = PROMPTS_PATH_CANDIDATES[0]
    PROMPTS_LOAD_REPORT = "fallback_default"
    if parse_errors:
        PROMPTS_LOAD_REPORT += " | " + " || ".join(parse_errors[:3])

    save_prompts(default_data)
    return default_data

def save_prompts(data):
    with open(PROMPTS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
def clone_json_data(data):
    try:
        return json.loads(json.dumps(data, ensure_ascii=False))
    except Exception:
        return data


def build_draft_data():
    return {key: clone_json_data(st.session_state[key]) for key in DRAFT_STATE_KEYS if key in st.session_state}


def queue_draft_restore(draft_data, notice=""):
    snapshot = clone_json_data(draft_data or {})
    st.session_state[PENDING_DRAFT_RESTORE_KEY] = snapshot
    if notice:
        st.session_state[DRAFT_RESTORE_NOTICE_KEY] = notice
    try:
        with open(DRAFT_FILE, "w", encoding="utf-8") as file_obj:
            json.dump(snapshot, file_obj, ensure_ascii=False, indent=2)
    except Exception as exc:
        print(f"读取草稿失败: {exc}")


def apply_draft_data(draft_data):
    for key, value in (draft_data or {}).items():
        st.session_state[key] = clone_json_data(value)

    current_role = st.session_state.get("selected_role", "")
    if isinstance(current_role, str):
        st.session_state.selected_role_widget = current_role
    st.session_state.target_article_words_slider = st.session_state.get("target_article_words", 1500)
    for stale_key in [key for key in list(st.session_state.keys()) if key.startswith("source_image_pick_")]:
        del st.session_state[stale_key]


def build_task_fallback_name(entity_id, default_suffix="001"):
    suffix = str(entity_id or "").replace("T", "") or default_suffix
    return f"{chr(0x4EFB)}{chr(0x52A1)} {suffix}"


def build_task_title(task_snapshot, fallback_name="\u4efb\u52a1"):
    snapshot = task_snapshot or {}
    titles = normalize_title_candidates(snapshot.get("title_candidates", []))
    if titles:
        return titles[0][:60]

    for field_name in ("article_url", "video_url"):
        raw_value = str(snapshot.get(field_name, "") or "")
        items = [line.strip() for line in raw_value.splitlines() if line.strip()]
        if items:
            first_url = items[0]
            host = urlparse(first_url).netloc.replace("www.", "") or first_url
            if len(items) > 1:
                return f"{host} \u7b49{len(items)}\u6761"[:60]
            return host[:60]

    article_candidates = [
        get_article_body_text((snapshot.get("final_article") or "").strip()),
        get_article_body_text((snapshot.get("modified_article") or "").strip()),
        get_article_body_text((snapshot.get("draft_article") or "").strip()),
        (snapshot.get("source_content") or "").strip(),
    ]
    for candidate in article_candidates:
        lines = [line.strip() for line in str(candidate).splitlines() if line.strip()]
        if lines:
            clean_line = re.sub(r"^[#>*\-\d\.)\s]+", "", lines[0]).strip()
            if clean_line:
                return clean_line[:60]

    return fallback_name


def derive_task_status(task_snapshot, explicit_error=""):
    snapshot = task_snapshot or {}
    last_error = (explicit_error or snapshot.get("last_ai_error", "") or "").strip()
    if last_error:
        return "failed"
    if (snapshot.get("final_article") or "").strip():
        return "completed"
    if snapshot.get("review_feedback") or snapshot.get("review_actions"):
        return "needs_review"
    if (
        snapshot.get("current_step", 1) > 1
        or (snapshot.get("source_content") or "").strip()
        or (snapshot.get("draft_article") or "").strip()
        or (snapshot.get("modified_article") or "").strip()
    ):
        return "in_progress"
    return "pending"


def build_task_metrics(task_snapshot):
    snapshot = task_snapshot or {}
    article_text = get_article_body_text(
        (snapshot.get("final_article") or snapshot.get("modified_article") or snapshot.get("draft_article") or "").strip()
    )
    return {
        "word_count": len((article_text or "").strip()),
        "version_count": len(snapshot.get("article_versions", []) or []),
        "has_highlight": bool((snapshot.get("highlighted_article") or "").strip()),
        "has_podcast": bool((snapshot.get("podcast_script_raw") or "").strip() or (snapshot.get("podcast_audio_path") or "").strip()),
    }


def build_task_queue_metrics(tasks):
    summary = {
        "total": 0,
        "pending": 0,
        "in_progress": 0,
        "needs_review": 0,
        "completed": 0,
        "failed": 0,
    }
    for task in tasks or []:
        summary["total"] += 1
        status = task.get("status", "pending")
        if status not in summary:
            continue
        summary[status] += 1
    return summary


def build_task_template_config(task_snapshot):
    snapshot = task_snapshot or {}
    return {key: clone_json_data(snapshot.get(key)) for key in TASK_TEMPLATE_CONFIG_KEYS if key in snapshot}


def apply_task_template_config(task_snapshot, template_config):
    snapshot = clone_json_data(task_snapshot or {})
    for key, value in (template_config or {}).items():
        snapshot[key] = clone_json_data(value)
    return snapshot


def get_next_entity_id(prefix, items):
    highest = 0
    for item in items or []:
        match = re.search(r"(\d+)$", str(item.get("id", "")))
        if match:
            highest = max(highest, int(match.group(1)))
    return f"{prefix}{highest + 1:03d}"


def get_task_by_id(task_id):
    for task in st.session_state.get("task_queue", []) or []:
        if task.get("id") == task_id:
            return task
    return None


def get_template_by_id(template_id):
    for template in st.session_state.get("task_templates", []) or []:
        if template.get("id") == template_id:
            return template
    return None


def is_placeholder_task_name(task_name):
    clean_name = str(task_name or "").strip()
    if not clean_name:
        return True
    placeholder_prefixes = (
        "任务 ",
        "?? ",
        "浠诬姟 ",
        "ÈÎÎñ ",
        "??? ",
    )
    return clean_name.startswith(placeholder_prefixes)


def refresh_task_record(task_record, task_snapshot=None):
    snapshot = clone_json_data(task_snapshot if task_snapshot is not None else task_record.get("snapshot", {}))
    task_record["snapshot"] = snapshot
    task_record["status"] = derive_task_status(snapshot, task_record.get("last_error", ""))
    task_record["current_step"] = int(snapshot.get("current_step", 1) or 1)
    task_record["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    task_record["resume_stage"] = snapshot.get("pending_ai_stage") or snapshot.get("last_completed_ai_stage") or ""
    task_record["last_error"] = (snapshot.get("last_ai_error", "") or "").strip()
    task_record["metrics"] = build_task_metrics(snapshot)
    if is_placeholder_task_name(task_record.get("name", "")):
        fallback_name = build_task_fallback_name(task_record.get('id', ''))
        task_record["name"] = build_task_title(snapshot, fallback_name=fallback_name)
    return task_record


def read_task_queue_data():
    if not os.path.exists(TASK_QUEUE_FILE):
        return None
    with open(TASK_QUEUE_FILE, "r", encoding="utf-8") as file_obj:
        queue_data = json.load(file_obj)
    if not isinstance(queue_data, dict):
        raise ValueError("Task queue payload must be a JSON object.")
    return queue_data


def save_task_queue_state():
    queue_payload = {
        "active_task_id": st.session_state.get("active_task_id", ""),
        "tasks": clone_json_data(st.session_state.get("task_queue", [])),
        "templates": clone_json_data(st.session_state.get("task_templates", [])),
    }
    with open(TASK_QUEUE_FILE, "w", encoding="utf-8") as file_obj:
        json.dump(queue_payload, file_obj, ensure_ascii=False, indent=2)


def init_task_queue_state():
    if st.session_state.get("_task_queue_loaded"):
        return

    tasks = []
    templates = []
    active_task_id = ""
    try:
        queue_data = read_task_queue_data() or {}
        tasks = [item for item in queue_data.get("tasks", []) if isinstance(item, dict)]
        templates = [item for item in queue_data.get("templates", []) if isinstance(item, dict)]
        active_task_id = queue_data.get("active_task_id", "") or ""
    except Exception:
        tasks = []
        templates = []
        active_task_id = ""

    normalized_tasks = []
    queue_changed = False
    for task in tasks:
        original_name = task.get("name", "")
        refresh_task_record(task)
        normalized_tasks.append(task)
        if task.get("name", "") != original_name:
            queue_changed = True

    st.session_state.task_queue = normalized_tasks
    st.session_state.task_templates = templates
    st.session_state.active_task_id = active_task_id
    st.session_state._task_queue_loaded = True
    if queue_changed:
        save_task_queue_state()


def persist_active_task_snapshot(draft_data=None):
    init_task_queue_state()
    active_task_id = st.session_state.get("active_task_id", "")
    if not active_task_id:
        return

    tasks = st.session_state.get("task_queue", []) or []
    snapshot = clone_json_data(draft_data if draft_data is not None else build_draft_data())
    task_record = next((item for item in tasks if item.get("id") == active_task_id), None)
    if task_record is None:
        task_record = {
            "id": active_task_id,
            "name": build_task_title(snapshot, fallback_name=build_task_fallback_name(active_task_id)),
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        tasks.append(task_record)
    refresh_task_record(task_record, snapshot)
    st.session_state.task_queue = tasks
    save_task_queue_state()


def build_blank_task_snapshot(base_snapshot=None):
    snapshot = clone_json_data(base_snapshot or {})
    reset_defaults = {
        "current_step": 1,
        "article_url": "",
        "video_url": "",
        "source_content": "",
        "source_images": [],
        "source_images_all": [],
        "selected_source_image_ids": [],
        "extraction_success": False,
        "draft_article": "",
        "review_feedback": "",
        "review_actions": [],
        "accepted_review_items": [],
        "modified_article": "",
        "final_article": "",
        "title_candidates": [],
        "highlighted_article": "",
        "spoken_script": "",
        "podcast_script_raw": "",
        "podcast_script_segments": [],
        "podcast_audio_path": "",
        "podcast_audio_manifest": {},
        "podcast_last_error": "",
        "chat_history": [],
        "image_keywords": "",
        "article_versions": [],
        "active_article_version_id": None,
        "term_scan_result": [],
        "term_scan_summary": {},
        "pending_ai_stage": "",
        "last_completed_ai_stage": "",
        "last_completed_ai_target_step": 0,
        "last_ai_error": "",
        "recovered_ai_notice": "",
        "obsidian_hits": [],
        "obsidian_research_brief": "",
        "obsidian_retrieval_error": "",
        "obsidian_query_terms": [],
        "obsidian_wiki_root": "",
        "obsidian_last_indexed_at": "",
        "obsidian_retrieval_signature": "",
        "obsidian_influence_map": [],
        "obsidian_influence_summary": "",
        "obsidian_influence_signature": "",
        "evidence_map": [],
        "evidence_summary": "",
        "evidence_signature": "",
    }
    for key, value in reset_defaults.items():
        snapshot[key] = clone_json_data(value)
    return snapshot


def draft_has_meaningful_content(draft_data):
    if not isinstance(draft_data, dict):
        return False
    if int(draft_data.get("current_step", 1) or 1) > 1:
        return True

    text_keys = (
        "article_url",
        "video_url",
        "source_content",
        "draft_article",
        "review_feedback",
        "modified_article",
        "final_article",
        "highlighted_article",
        "spoken_script",
        "podcast_script_raw",
        "image_keywords",
        "obsidian_research_brief",
    )
    for key in text_keys:
        if str(draft_data.get(key, "") or "").strip():
            return True

    list_keys = (
        "source_images",
        "source_images_all",
        "selected_source_image_ids",
        "title_candidates",
        "review_actions",
        "accepted_review_items",
        "article_versions",
        "podcast_script_segments",
        "chat_history",
        "obsidian_hits",
        "evidence_map",
    )
    for key in list_keys:
        if draft_data.get(key):
            return True
    return False


def snapshots_equivalent(left_snapshot, right_snapshot):
    try:
        return json.dumps(left_snapshot or {}, ensure_ascii=False, sort_keys=True) == json.dumps(right_snapshot or {}, ensure_ascii=False, sort_keys=True)
    except Exception:
        return clone_json_data(left_snapshot or {}) == clone_json_data(right_snapshot or {})


def should_offer_draft_restore():
    try:
        draft_data = read_draft_data()
    except Exception:
        return False
    if not draft_has_meaningful_content(draft_data):
        return False
    return not snapshots_equivalent(draft_data, build_draft_data())


def reset_active_task_to_blank():
    init_task_queue_state()
    active_task_id = st.session_state.get("active_task_id", "")
    if not active_task_id:
        return False

    task_record = get_task_by_id(active_task_id)
    if not task_record:
        return False

    blank_snapshot = build_blank_task_snapshot(build_draft_data())
    task_record["name"] = build_task_fallback_name(active_task_id)
    refresh_task_record(task_record, blank_snapshot)
    st.session_state.task_queue = st.session_state.get("task_queue", []) or []
    st.session_state[TASK_QUEUE_NOTICE_KEY] = f"已重置当前任务：{task_record.get('name', active_task_id)}"
    save_task_queue_state()
    queue_draft_restore(blank_snapshot)
    return True


def ensure_task_queue_bootstrap():
    init_task_queue_state()
    tasks = st.session_state.get("task_queue", []) or []
    active_task_id = st.session_state.get("active_task_id", "")

    if tasks:
        if not active_task_id or not any(task.get("id") == active_task_id for task in tasks):
            st.session_state.active_task_id = tasks[0].get("id", "")
        if not os.path.exists(DRAFT_FILE):
            active_task = get_task_by_id(st.session_state.get("active_task_id", ""))
            if active_task and isinstance(active_task.get("snapshot"), dict):
                apply_draft_data(active_task.get("snapshot", {}))
                save_draft()
        return

    initial_snapshot = build_draft_data()
    task_id = get_next_entity_id("T", tasks)
    task_record = {
        "id": task_id,
        "name": build_task_title(initial_snapshot, fallback_name=build_task_fallback_name(task_id)),
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    refresh_task_record(task_record, initial_snapshot)
    st.session_state.task_queue = [task_record]
    st.session_state.active_task_id = task_id
    save_task_queue_state()


def create_task_from_current_state(*, clone_current=False, template_id="", task_name=""):
    init_task_queue_state()
    current_snapshot = build_draft_data()
    task_snapshot = clone_json_data(current_snapshot) if clone_current else build_blank_task_snapshot(current_snapshot)

    template = get_template_by_id(template_id) if template_id else None
    if template:
        task_snapshot = apply_task_template_config(task_snapshot, template.get("config_snapshot", {}))

    tasks = st.session_state.get("task_queue", []) or []
    task_id = get_next_entity_id("T", tasks)
    task_record = {
        "id": task_id,
        "name": (task_name or "").strip() or build_task_title(task_snapshot, fallback_name=build_task_fallback_name(task_id)),
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    refresh_task_record(task_record, task_snapshot)
    tasks.append(task_record)
    st.session_state.task_queue = tasks
    st.session_state.active_task_id = task_id
    st.session_state[TASK_QUEUE_NOTICE_KEY] = f"已创建新任务：{task_record['name']}"
    save_task_queue_state()
    queue_draft_restore(task_snapshot)
    return task_id


def switch_to_task(task_id):
    init_task_queue_state()
    task_record = get_task_by_id(task_id)
    if not task_record:
        return False

    current_active_id = st.session_state.get("active_task_id", "")
    if current_active_id and current_active_id != task_id:
        persist_active_task_snapshot()

    st.session_state.active_task_id = task_id
    st.session_state[TASK_QUEUE_NOTICE_KEY] = f"???????{task_record.get('name', task_id)}"
    save_task_queue_state()
    queue_draft_restore(task_record.get("snapshot", {}))
    return True


def resume_task(task_id=""):
    init_task_queue_state()
    target_task_id = task_id or st.session_state.get("active_task_id", "")
    task_record = get_task_by_id(target_task_id)
    if not task_record:
        return False

    snapshot = clone_json_data(task_record.get("snapshot", {}))
    target_step = int(task_record.get("current_step", snapshot.get("current_step", 1) or 1))
    snapshot["current_step"] = target_step
    snapshot["last_ai_error"] = ""
    st.session_state.active_task_id = target_task_id
    st.session_state.last_ai_error = ""
    queue_draft_restore(snapshot)
    return True


def delete_task(task_id):
    init_task_queue_state()
    tasks = st.session_state.get("task_queue", []) or []
    if len(tasks) <= 1:
        return False

    task_record = get_task_by_id(task_id)
    if not task_record:
        return False

    remaining_tasks = [task for task in tasks if task.get("id") != task_id]
    if len(remaining_tasks) == len(tasks):
        return False

    st.session_state.task_queue = remaining_tasks
    active_task_id = st.session_state.get("active_task_id", "")
    if active_task_id == task_id:
        next_task = remaining_tasks[0] if remaining_tasks else None
        st.session_state.active_task_id = next_task.get("id", "") if next_task else ""
        if next_task:
            queue_draft_restore(next_task.get("snapshot", {}))

    st.session_state[TASK_QUEUE_NOTICE_KEY] = f"已删除任务：{task_record.get('name', task_id)}"
    save_task_queue_state()
    return True


def save_current_config_as_template(template_name):
    clean_name = (template_name or "").strip()
    if not clean_name:
        return None

    init_task_queue_state()
    templates = st.session_state.get("task_templates", []) or []
    config_snapshot = build_task_template_config(build_draft_data())
    existing_template = next((item for item in templates if (item.get("name") or "").strip() == clean_name), None)
    now_text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if existing_template:
        existing_template["config_snapshot"] = config_snapshot
        existing_template["updated_at"] = now_text
        template_id = existing_template.get("id")
    else:
        template_id = get_next_entity_id("TP", templates)
        templates.append({
            "id": template_id,
            "name": clean_name,
            "created_at": now_text,
            "updated_at": now_text,
            "config_snapshot": config_snapshot,
        })

    st.session_state.task_templates = templates
    save_task_queue_state()
    return template_id


def apply_template_to_tasks(template_id, task_ids=None):
    template = get_template_by_id(template_id)
    if not template:
        return 0

    target_ids = list(task_ids or [])
    if not target_ids and st.session_state.get("active_task_id"):
        target_ids = [st.session_state.get("active_task_id")]
    if not target_ids:
        return 0

    templates_config = template.get("config_snapshot", {})
    updated_count = 0
    tasks = st.session_state.get("task_queue", []) or []
    active_task_id = st.session_state.get("active_task_id", "")
    active_snapshot = None

    for task_record in tasks:
        if task_record.get("id") not in target_ids:
            continue
        snapshot = apply_task_template_config(task_record.get("snapshot", {}), templates_config)
        refresh_task_record(task_record, snapshot)
        updated_count += 1
        if task_record.get("id") == active_task_id:
            active_snapshot = snapshot

    st.session_state.task_queue = tasks
    save_task_queue_state()
    if active_snapshot is not None:
        queue_draft_restore(active_snapshot)
    return updated_count


def build_batch_export_markdown(tasks):
    export_parts = []
    for task_record in tasks or []:
        snapshot = task_record.get("snapshot", {}) or {}
        article_text = (snapshot.get("final_article") or snapshot.get("modified_article") or snapshot.get("draft_article") or "").strip()
        highlighted_article = (snapshot.get("highlighted_article") or "").strip()
        part_lines = [
            f"# {task_record.get('name', task_record.get('id', '任务'))}",
            "",
            f"- 任务 ID：{task_record.get('id', '')}",
            f"- 状态：{TASK_STATUS_LABELS.get(task_record.get('status', 'pending'), task_record.get('status', 'pending'))}",
            f"- 最后更新时间：{task_record.get('updated_at', '')}",
            "",
            "## 正文",
            "",
            article_text or "暂无正文内容。",
        ]
        if highlighted_article:
            part_lines.extend(["", "## 高亮阅读版", "", highlighted_article])
        export_parts.append("\n".join(part_lines).strip())
    return "\n\n---\n\n".join(part for part in export_parts if part).strip()


def save_draft():
    draft_data = build_draft_data()
    try:
        with open(DRAFT_FILE, "w", encoding="utf-8") as file_obj:
            json.dump(draft_data, file_obj, ensure_ascii=False, indent=2)
        persist_active_task_snapshot(draft_data)
    except Exception as exc:
        print(f"\u8bfb\u53d6\u8349\u7a3f\u5931\u8d25: {exc}")


def read_draft_data():
    if not os.path.exists(DRAFT_FILE):
        return None

    with open(DRAFT_FILE, "r", encoding="utf-8") as file_obj:
        draft_data = json.load(file_obj)

    if not isinstance(draft_data, dict):
        raise ValueError("Draft payload must be a JSON object.")

    return draft_data


def load_draft():
    try:
        draft_data = read_draft_data()
        if draft_data is None:
            return False
        st.session_state[PENDING_DRAFT_RESTORE_KEY] = draft_data
        st.session_state[DRAFT_RESTORE_NOTICE_KEY] = "已恢复上次草稿内容。"
        return True
    except Exception as exc:
        print(f"读取草稿失败: {exc}")
        return False


def apply_pending_draft_restore():
    pending_draft = st.session_state.pop(PENDING_DRAFT_RESTORE_KEY, None)
    if not isinstance(pending_draft, dict):
        return

    apply_draft_data(pending_draft)


def clear_draft():
    if os.path.exists(DRAFT_FILE):
        try:
            os.remove(DRAFT_FILE)
        except Exception:
            pass


def ensure_podcast_runtime_dirs():
    os.makedirs(PODCAST_OUTPUT_DIR, exist_ok=True)
    os.makedirs(PODCAST_CACHE_DIR, exist_ok=True)


def reset_podcast_outputs(delete_audio=False):
    audio_path = st.session_state.get("podcast_audio_path", "")
    if delete_audio and audio_path and os.path.exists(audio_path):
        try:
            os.remove(audio_path)
        except OSError:
            pass
    st.session_state.podcast_script_raw = ""
    st.session_state.podcast_script_segments = []
    st.session_state.podcast_audio_path = ""
    st.session_state.podcast_audio_manifest = {}
    st.session_state.podcast_last_error = ""


def generate_podcast_script_for_current_article(api_key, base_url, model_name, podcast_duration):
    final_article = get_article_body_text((st.session_state.get("final_article") or "").strip())
    if not final_article:
        reset_podcast_outputs(delete_audio=True)
        return False

    raw_script = call_llm(
        api_key=api_key,
        base_url=base_url,
        model_name=model_name,
        system_prompt=get_podcast_sys_prompt(podcast_duration),
        user_content=f"请将以下文章改写为适合 {podcast_duration} 的单人中文播客解说 JSON：\n\n{final_article}",
    )

    try:
        segments = parse_podcast_script_segments(raw_script)
    except ValueError as exc:
        reset_podcast_outputs(delete_audio=True)
        st.session_state.podcast_script_raw = raw_script
        st.session_state.podcast_last_error = str(exc)
        save_draft()
        return False

    st.session_state.podcast_script_raw = raw_script
    st.session_state.podcast_script_segments = segments
    st.session_state.podcast_audio_path = ""
    st.session_state.podcast_audio_manifest = {}
    st.session_state.podcast_last_error = ""
    save_draft()
    return True

def sanitize_editor_prompt(prompt_text):
    if not isinstance(prompt_text, str):
        return ""
    cleaned_lines = []
    for line in prompt_text.splitlines():
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append(line)
            continue

        has_digit = any(ch.isdigit() for ch in stripped)
        should_remove = (
            "\u76ee\u6807\u5b57\u6570" in stripped
            or "\u5b57\u6570\u8981\u6c42" in stripped
            or "\u603b\u5b57\u6570" in stripped
            or ("\u63a7\u5236\u5728" in stripped and "\u5b57" in stripped and has_digit)
            or ("\u4e0d\u8d85" in stripped and "\u5b57" in stripped and has_digit)
            or ("\u4e0d\u8d85\u8fc7" in stripped and "\u5b57" in stripped and has_digit)
            or ("\u4e25\u7981\u8d85\u8fc7" in stripped and "\u5b57" in stripped and has_digit)
            or ("\u6574\u4f53\u4e0d\u8d85" in stripped and "\u5b57" in stripped and has_digit)
        )
        if should_remove:
            continue
        cleaned_lines.append(line)
    cleaned_prompt = "\n".join(cleaned_lines)
    while "\n\n\n" in cleaned_prompt:
        cleaned_prompt = cleaned_prompt.replace("\n\n\n", "\n\n")
    return cleaned_prompt.strip()

def get_target_article_words():
    raw_value = st.session_state.get("target_article_words", 1500)
    try:
        value = int(raw_value)
    except Exception:
        value = 1500
    return max(200, min(5000, value))

def sync_target_article_words():
    raw_value = st.session_state.get("target_article_words_slider", get_target_article_words())
    try:
        value = int(raw_value)
    except Exception:
        value = 1500
    st.session_state.target_article_words = max(200, min(5000, value))
    save_draft()

def sync_selected_role():
    selected_role = st.session_state.get("selected_role_widget", "")
    if isinstance(selected_role, str) and selected_role.strip():
        st.session_state.selected_role = selected_role
        save_draft()



def build_target_length_instruction():
    target_words = get_target_article_words()
    return (
        f"[Global Length Requirement | Highest Priority]\n"
        f"Target length for this run: about {target_words} words (allow +/-10%).\n"
        "If any role prompt contains a fixed word-count target, ignore it and follow this run-level target instead."
    )


ARTICLE_TITLE_MARKER = "\u3010\u5907\u9009\u6807\u9898\u3011"
ARTICLE_BODY_MARKER = "\u3010\u6b63\u6587\u3011"
PURE_TITLE_MARKER = "\u3010\u7eaf\u51c0\u6807\u9898\u7ec4\u3011"
PURE_BODY_MARKER = "\u3010\u7eaf\u51c0\u5b9a\u7a3f\u3011"
HIGHLIGHT_MARKER = "\u3010\u9ad8\u4eae\u9605\u8bfb\u7248\u3011"


def build_article_structure_instruction(target_words=None):
    if target_words is None:
        target_words = get_target_article_words()
    try:
        target_words = int(target_words)
    except Exception:
        target_words = get_target_article_words()

    if target_words < 1200:
        heading_instruction = "Use exactly 3 `##` sections for a short article."
    elif target_words <= 2500:
        heading_instruction = "Use 3-5 `##` sections so the middle of the article is clearly structured."
    else:
        heading_instruction = "Use 4-5 `##` sections by default; add `###` only when one section is genuinely complex."

    return "\n".join([
        "[Article Structure Protocol | Highest Priority]",
        "The article body must be written in Markdown.",
        "Use a short opening lede of 1-2 paragraphs, then move into the main analysis.",
        heading_instruction,
        "Each `##` section should usually contain 2-4 natural paragraphs.",
        "Section titles must be informative and specific; do not use empty headings like `Background`, `Summary`, or `Some Thoughts`.",
        "Keep the structure clear without making every section mechanically equal.",
    ])


def build_reviewer_structure_instruction():
    structure_check_marker = "\u3010\u7ed3\u6784\u6027\u68c0\u67e5\u3011"
    background_check_marker = "\u3010\u80cc\u666f\u5b8c\u6574\u6027\u68c0\u67e5\u3011"
    title_check_marker = "\u3010\u6807\u9898\u7ec4\u68c0\u67e5\u3011"
    return "\n".join([
        "[Structural Review Requirements]",
        f"You must add a dedicated `{structure_check_marker}` section to the review report.",
        f"Inside that section, you must explicitly cover `{background_check_marker}` and `{title_check_marker}`.",
        "Check whether the opening 1-2 paragraphs clearly tell the reader which game, event, company, or controversy the article is about.",
        "Check whether the opening also explains why this case matters now before the analysis begins.",
        "Check whether 3-5 candidate titles are still present, aligned with the body, and not drifting in framing.",
        "Check whether any `##` section is too long, whether headings are too vague, and whether the article still reads like one uninterrupted wall of text.",
        "If the structure is good, say so explicitly. If it is not, give concrete repair instructions in the revision feedback.",
        "Inside `\u3010\u4fee\u6539\u4efb\u52a1\u6e05\u5355\u3011`, every task must be a top-level Markdown bullet item.",
        "Use this stable shape for each task block: `- \u4efb\u52a1 1\uff5c\u4e00\u53e5\u8bdd\u6458\u8981`, then `\u95ee\u9898\u7c7b\u578b\uff1a...`, `\u5bf9\u5e94\u4f4d\u7f6e\u6216\u539f\u53e5\uff1a...`, `\u4e3a\u4ec0\u4e48\u8981\u6539\uff1a...`, `\u5e94\u8be5\u600e\u4e48\u6539\uff1a...`.",
    ])


def build_reviewer_system_prompt(reviewer_prompt, anti_hallucination_instruction=""):
    humanizer_review_instruction = "\n".join([
        "[Humanizer Review Requirements]",
        "Inside your expression review, explicitly check whether the draft still shows obvious AI writing patterns or connector-word stacking.",
        "Flag formulaic sentence shapes such as `不仅仅是`, `不是……而是……`, and `值得一提的是` when they make the article feel templated.",
        "Flag empty uplift, promotional tone, fake-depth phrasing, vague authority claims, and generic positive endings that do not add real information.",
        "Flag overly neat three-part parallel structures or quote-like lines that feel manufactured instead of earned.",
        "When you identify these issues, explain why they hurt credibility and give concrete rewrite instructions instead of abstract style criticism.",
    ])
    prompt_parts = [
        reviewer_prompt.strip() if isinstance(reviewer_prompt, str) else "",
        build_reviewer_structure_instruction(),
        humanizer_review_instruction,
        anti_hallucination_instruction.strip() if isinstance(anti_hallucination_instruction, str) else "",
    ]
    return "\n\n".join([part for part in prompt_parts if part])

def build_article_output_instruction():
    return "\n".join([
        "[Opening Background Protocol | Highest Priority]",
        "Paragraph 1 must tell the reader exactly which game, event, company, or controversy this article is about.",
        "Paragraph 2 should explain the immediate background, the current debate, or why this case is worth analyzing now.",
        "Do not start with conclusions, monetization analysis, or system-level breakdown before the object/event is introduced.",
        "[Structured Output Protocol | Strict]",
        "Keep the title group alive at every stage. Do not swallow it into the body.",
        "All candidate titles and the body must be written in Simplified Chinese.",
        f"Output exactly two blocks in this order:\n{ARTICLE_TITLE_MARKER}\n1. ...\n2. ...\n3. ...\nProvide 3-5 candidate titles.\n\n{ARTICLE_BODY_MARKER}\nWrite the full article body here.",
    ])


def build_editor_system_prompt(editor_prompt, global_instruction):
    prompt_parts = [
        sanitize_editor_prompt(editor_prompt),
        build_target_length_instruction(),
        build_article_structure_instruction(),
        build_article_output_instruction(),
        global_instruction.strip() if isinstance(global_instruction, str) else "",
    ]
    return "\n\n".join([part for part in prompt_parts if part])


def build_modification_system_prompt(global_instruction, term_rules_instruction=""):
    base_prompt = (
        "你是一名资深中文游戏内容编辑，负责根据审稿意见对文章做定向修订。"
        "请在保留核心事实、分析骨架、标题组和文章结构的前提下完成修改，不要重写成另一篇完全不同的稿子。"
    )
    prompt_parts = [
        base_prompt,
        build_target_length_instruction(),
        build_article_structure_instruction(),
        build_article_output_instruction(),
        term_rules_instruction.strip() if isinstance(term_rules_instruction, str) else "",
        global_instruction.strip() if isinstance(global_instruction, str) else "",
    ]
    return "\n\n".join([part for part in prompt_parts if part])


def parse_banned_terms_text(text):
    if not isinstance(text, str):
        return []
    terms = []
    seen = set()
    for raw_line in text.splitlines():
        term = raw_line.strip()
        if not term or term in seen:
            continue
        seen.add(term)
        terms.append(term)
    return terms


def parse_replacement_terms_text(text):
    if not isinstance(text, str):
        return {}, []
    replacements = {}
    invalid_lines = []
    separators = ("=>", "->", "→")
    for raw_line in text.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            continue
        separator = next((sep for sep in separators if sep in stripped), None)
        if not separator:
            invalid_lines.append(stripped)
            continue
        source, target = stripped.split(separator, 1)
        source = source.strip()
        target = target.strip()
        if not source or not target or source == target:
            invalid_lines.append(stripped)
            continue
        replacements[source] = target
    return replacements, invalid_lines


def merge_term_rules(global_banned_terms, global_replacements, article_banned_terms=None, article_replacements=None):
    merged_banned_terms = []
    seen = set()
    for term in list(global_banned_terms or []) + list(article_banned_terms or []):
        clean_term = (term or "").strip()
        if not clean_term or clean_term in seen:
            continue
        seen.add(clean_term)
        merged_banned_terms.append(clean_term)

    merged_replacements = {}
    for source, target in (global_replacements or {}).items():
        clean_source = (source or "").strip()
        clean_target = (target or "").strip()
        if clean_source and clean_target and clean_source != clean_target:
            merged_replacements[clean_source] = clean_target
    for source, target in (article_replacements or {}).items():
        clean_source = (source or "").strip()
        clean_target = (target or "").strip()
        if clean_source and clean_target and clean_source != clean_target:
            merged_replacements[clean_source] = clean_target
    return merged_banned_terms, merged_replacements


def merge_replacement_terms(global_replacements=None, article_replacements=None):
    merged_replacements = {}
    for replacements in (global_replacements or {}, article_replacements or {}):
        for source, target in replacements.items():
            clean_source = (source or "").strip()
            clean_target = (target or "").strip()
            if clean_source and clean_target and clean_source != clean_target:
                merged_replacements[clean_source] = clean_target
    return merged_replacements


def build_term_rules_instruction(banned_terms=None, replacement_terms=None):
    resolved_banned_terms = [term for term in (banned_terms or []) if isinstance(term, str) and term.strip()]
    resolved_replacement_terms = {
        source: target
        for source, target in (replacement_terms or {}).items()
        if isinstance(source, str) and source.strip() and isinstance(target, str) and target.strip()
    }
    if not resolved_banned_terms and not resolved_replacement_terms:
        return ""

    lines = ["【个人词表约束】"]
    if resolved_banned_terms:
        lines.append("以下表达请尽量不要出现：")
        lines.extend([f"- {term}" for term in resolved_banned_terms])
    if resolved_replacement_terms:
        lines.append("如果需要表达相近意思，优先替换为：")
        lines.extend([f"- {source} -> {target}" for source, target in resolved_replacement_terms.items()])
    lines.append("这些要求只作用于表达层，不要因此改动事实、逻辑、结构和标题组。")
    return "\n".join(lines)


def build_term_rules_preview_text(banned_terms=None, default_replacement_terms=None, suggested_replacement_terms=None):
    preview_blocks = []
    primary_block = build_term_rules_instruction(banned_terms, default_replacement_terms)
    if primary_block:
        preview_blocks.append(primary_block)

    resolved_suggested_replacement_terms = {
        source: target
        for source, target in (suggested_replacement_terms or {}).items()
        if isinstance(source, str) and source.strip() and isinstance(target, str) and target.strip()
    }
    if resolved_suggested_replacement_terms:
        lines = [
            "【建议替换词表】",
            "以下表达主要用于最终稿阶段的人工优化提示，不强制注入当前 prompt：",
        ]
        lines.extend([f"- {source} -> {target}" for source, target in resolved_suggested_replacement_terms.items()])
        preview_blocks.append("\n".join(lines))

    return "\n\n".join([block for block in preview_blocks if block])


def scan_article_terms(article_text, banned_terms=None, replacement_terms=None, suggested_replacement_terms=None):
    banned_set = {term for term in (banned_terms or []) if isinstance(term, str) and term.strip()}
    replacement_terms = replacement_terms or {}
    suggested_replacement_terms = suggested_replacement_terms or {}
    ordered_terms = []
    seen = set()
    for term in list(banned_terms or []) + list(replacement_terms.keys()) + list(suggested_replacement_terms.keys()):
        clean_term = (term or "").strip()
        if not clean_term or clean_term in seen:
            continue
        seen.add(clean_term)
        ordered_terms.append(clean_term)
    if not ordered_terms:
        return []

    article_body = get_article_body_text(article_text)
    paragraphs = [paragraph.strip() for paragraph in re.split(r"\n\s*\n", article_body) if paragraph.strip()]
    if not paragraphs and article_body.strip():
        paragraphs = [article_body.strip()]

    scan_result = []
    for term in ordered_terms:
        count = 0
        paragraph_indexes = []
        for index, paragraph in enumerate(paragraphs, start=1):
            term_hits = paragraph.count(term)
            if term_hits:
                count += term_hits
                paragraph_indexes.append(index)
        if count:
            level = "default" if term in replacement_terms else "suggested" if term in suggested_replacement_terms else ""
            replacement = replacement_terms.get(term) or suggested_replacement_terms.get(term, "")
            scan_result.append({
                "type": "banned" if term in banned_set else "replacement",
                "level": level,
                "term": term,
                "replacement": replacement,
                "count": count,
                "paragraph_indexes": paragraph_indexes,
            })
    return scan_result


def summarize_term_scan(scan_result):
    result = scan_result if isinstance(scan_result, list) else []
    return {
        "matched_terms": len(result),
        "total_hits": sum(int(item.get("count", 0) or 0) for item in result),
        "banned_terms": sum(1 for item in result if item.get("type") == "banned"),
        "default_replacement_terms": sum(
            1 for item in result if item.get("type") == "replacement" and item.get("level") == "default"
        ),
        "suggested_replacement_terms": sum(
            1 for item in result if item.get("type") == "replacement" and item.get("level") == "suggested"
        ),
    }



def get_expected_h2_range(target_words=None):
    resolved_words = target_words
    if resolved_words in (None, "", 0):
        getter = globals().get("get_target_article_words")
        if callable(getter):
            resolved_words = getter()
    try:
        word_count = int(resolved_words or 0)
    except (TypeError, ValueError):
        word_count = 0
    if word_count <= 1200:
        return (3, 3)
    if word_count >= 2500:
        return (4, 5)
    return (3, 5)


def split_body_paragraphs(article_text):
    article_body = get_article_body_text(article_text)
    paragraphs = [paragraph.strip() for paragraph in re.split(r"\n\s*\n", article_body) if paragraph.strip()]
    if not paragraphs and article_body.strip():
        paragraphs = [article_body.strip()]
    return paragraphs


def extract_markdown_h2_sections(article_text):
    article_body = get_article_body_text(article_text)
    if not article_body.strip():
        return []

    sections = []
    current_section = None
    paragraph_buffer = []

    def flush_paragraph_buffer():
        nonlocal paragraph_buffer, current_section
        if current_section is None:
            paragraph_buffer = []
            return
        paragraph_text = "\n".join(paragraph_buffer).strip()
        if paragraph_text:
            current_section["paragraphs"].append(paragraph_text)
        paragraph_buffer = []

    for raw_line in article_body.splitlines():
        heading_match = re.match(r"^\s*##\s+(.+?)\s*$", raw_line)
        if heading_match:
            flush_paragraph_buffer()
            current_section = {
                "heading": heading_match.group(1).strip(),
                "paragraphs": [],
            }
            sections.append(current_section)
            continue

        if not raw_line.strip():
            flush_paragraph_buffer()
            continue

        if current_section is not None:
            paragraph_buffer.append(raw_line.strip())

    flush_paragraph_buffer()

    for section in sections:
        section["paragraph_count"] = len(section.get("paragraphs", []))
    return sections


def build_publish_quality_gate_report(
    article_text,
    title_candidates=None,
    highlighted_article="",
    term_scan_summary=None,
    target_words=None,
):
    explicit_titles = normalize_title_candidates(title_candidates)
    parsed_titles, _ = split_structured_article_sections(article_text)
    resolved_titles = explicit_titles or parsed_titles
    title_count = len(resolved_titles)

    article_body = get_article_body_text(article_text)
    expected_h2_range = get_expected_h2_range(target_words)
    min_h2, max_h2 = expected_h2_range

    intro_source = re.split(r"^\s*##\s+.+$", article_body, maxsplit=1, flags=re.MULTILINE)[0]
    intro_paragraphs = [paragraph.strip() for paragraph in re.split(r"\n\s*\n", intro_source) if paragraph.strip()]

    h2_sections = extract_markdown_h2_sections(article_text)
    h2_count = len(h2_sections)
    term_summary = term_scan_summary if isinstance(term_scan_summary, dict) else {}
    humanizer_fail_patterns = [
        ("\u6a21\u677f\u8f6c\u6298\u53e5", r"\u4e0d(?:\u662f|\u4ec5|\u4ec5\u4ec5\u662f)[^\u3002\uff01\uff1f\\n]{0,24}(?:\u800c\u662f|\u66f4\u662f)"),
        ("\u7a7a\u6cdb\u5347\u534e\u53e5", r"(?:\u8fd9|\u5b83)(?:\u4e0d|\u4e0d\u4ec5|\u4e0d\u4ec5\u4ec5)\u662f[^\u3002\uff01\uff1f\\n]{0,24}(?:\u800c\u662f|\u66f4\u662f)"),
    ]
    humanizer_warn_patterns = [
        ("AI\u8fde\u63a5\u8bcd", r"(?:\u6b64\u5916|\u4e0e\u6b64\u540c\u65f6|\u503c\u5f97\u4e00\u63d0\u7684\u662f|\u4ece\u67d0\u79cd\u610f\u4e49\u4e0a\u8bf4|\u4e0d\u53ef\u5426\u8ba4|\u663e\u800c\u6613\u89c1|\u4f17\u6240\u5468\u77e5)"),
        ("\u5ba3\u4f20\u8154\u8bcd\u6c47", r"(?:\u81f3\u5173\u91cd\u8981|\u5173\u952e\u6027\u7684|\u6301\u4e45\u7684|\u5145\u6ee1\u6d3b\u529b\u7684|\u5b8c\u7f8e\u878d\u5408|\u4ee4\u4eba\u60ca\u53f9|\u53f2\u65e0\u524d\u4f8b|\u524d\u6240\u672a\u6709)"),
        ("\u6a21\u7cca\u5f52\u56e0", r"(?:\u884c\u4e1a\u62a5\u544a\u663e\u793a|\u4e13\u5bb6\u8ba4\u4e3a|\u6709\u5206\u6790\u8ba4\u4e3a|\u4e00\u4e9b\u6279\u8bc4\u8005\u8ba4\u4e3a|\u89c2\u5bdf\u4eba\u58eb\u6307\u51fa)"),
        ("\u7a7a\u6cdb\u7ed3\u5c3e", r"(?:\u8fd9\u4e5f\u63d0\u9192\u6211\u4eec|\u672a\u6765\u53ef\u671f|\u503c\u5f97\u6df1\u601d|\u4e0d\u5bb9\u5ffd\u89c6)"),
    ]
    humanizer_fail_hits = [label for label, pattern in humanizer_fail_patterns if re.search(pattern, article_body)]
    humanizer_warn_hits = [label for label, pattern in humanizer_warn_patterns if re.search(pattern, article_body)]

    items = []

    if 3 <= title_count <= 5:
        items.append({"key": "titles", "label": "标题组", "status": "pass", "detail": f"已保留 {title_count} 个备选标题。"})
    else:
        items.append({"key": "titles", "label": "标题组", "status": "fail", "detail": f"当前标题组只有 {title_count} 个，发布前需要补到 3-5 个。"})

    if not intro_paragraphs:
        items.append({"key": "intro", "label": "开头导语", "status": "fail", "detail": "开头缺少 1-2 段导语，读者还没搞清对象和背景。"})
    elif len(intro_paragraphs) <= 2:
        items.append({"key": "intro", "label": "开头导语", "status": "pass", "detail": f"导语段落数为 {len(intro_paragraphs)}，符合当前结构要求。"})
    else:
        items.append({"key": "intro", "label": "开头导语", "status": "warn", "detail": f"导语段落数为 {len(intro_paragraphs)}，建议压缩到 1-2 段。"})

    if min_h2 <= h2_count <= max_h2:
        items.append({"key": "h2_count", "label": "二级标题", "status": "pass", "detail": f"当前共有 {h2_count} 个 `##`，落在建议范围 {min_h2}-{max_h2} 内。"})
    else:
        items.append({"key": "h2_count", "label": "二级标题", "status": "fail", "detail": f"当前共有 {h2_count} 个 `##`，建议范围应为 {min_h2}-{max_h2} 个。"})

    if h2_count < 2:
        items.append({"key": "section_balance", "label": "小节均衡度", "status": "warn", "detail": "当前还没有足够的 `##` 小节，暂时无法判断小节均衡度。"})
    else:
        imbalanced_sections = [
            section.get("heading", "未命名小节")
            for section in h2_sections
            if section.get("paragraph_count", 0) < 2 or section.get("paragraph_count", 0) > 4
        ]
        if imbalanced_sections:
            items.append({"key": "section_balance", "label": "小节均衡度", "status": "warn", "detail": "以下小节段落数失衡，建议再修一下：" + "、".join(imbalanced_sections) + "。"})
        else:
            items.append({"key": "section_balance", "label": "小节均衡度", "status": "pass", "detail": "各个 `##` 小节段落数基本均衡。"})

    if (highlighted_article or "").strip():
        items.append({"key": "highlight", "label": "高亮阅读版", "status": "pass", "detail": "高亮阅读版已生成。"})
    else:
        items.append({"key": "highlight", "label": "高亮阅读版", "status": "warn", "detail": "当前版本还没有高亮阅读版视图。"})

    banned_hits = int(term_summary.get("banned_terms", 0) or 0)
    default_hits = int(term_summary.get("default_replacement_terms", 0) or 0)
    suggested_hits = int(term_summary.get("suggested_replacement_terms", 0) or 0)
    if banned_hits > 0:
        items.append({"key": "term_rules", "label": "\u8bcd\u8868\u98ce\u9669", "status": "fail", "detail": f"\u4ecd\u547d\u4e2d {banned_hits} \u4e2a\u7981\u7528\u8bcd\uff0c\u53d1\u5e03\u524d\u5efa\u8bae\u5148\u5904\u7406\u3002"})
    elif default_hits > 0 or suggested_hits > 0:
        items.append({"key": "term_rules", "label": "\u8bcd\u8868\u98ce\u9669", "status": "warn", "detail": f"\u4ecd\u547d\u4e2d\u9ed8\u8ba4\u66ff\u6362 {default_hits} \u4e2a\u3001\u5efa\u8bae\u66ff\u6362 {suggested_hits} \u4e2a\u3002"})
    else:
        items.append({"key": "term_rules", "label": "\u8bcd\u8868\u98ce\u9669", "status": "pass", "detail": "\u5f53\u524d\u6ca1\u6709\u547d\u4e2d\u7981\u7528\u8bcd\u6216\u66ff\u6362\u8bcd\u98ce\u9669\u3002"})

    if humanizer_fail_hits or len(humanizer_warn_hits) >= 3:
        humanizer_detail = "\u5b58\u5728\u660e\u663e\u7684 AI \u5199\u4f5c\u75d5\u8ff9\uff0c\u5efa\u8bae\u518d\u505a\u4e00\u8f6e\u53bb\u6a21\u677f\u5316\u91cd\u5199\u3002"
        if humanizer_fail_hits:
            humanizer_detail += " \u91cd\u70b9\u95ee\u9898\uff1a" + "\u3001".join(humanizer_fail_hits) + "\u3002"
        elif humanizer_warn_hits:
            humanizer_detail += " \u89e6\u53d1\u4fe1\u53f7\uff1a" + "\u3001".join(humanizer_warn_hits[:4]) + "\u3002"
        items.append({"key": "humanizer_risk", "label": "AI\u75d5\u8ff9\u98ce\u9669", "status": "fail", "detail": humanizer_detail})
    elif humanizer_warn_hits:
        items.append({"key": "humanizer_risk", "label": "AI\u75d5\u8ff9\u98ce\u9669", "status": "warn", "detail": "\u4ecd\u6709\u4e00\u4e9b\u516c\u5f0f\u5316\u8fde\u63a5\u8bcd\u3001\u5ba3\u4f20\u8154\u6216\u7a7a\u6cdb\u7ed3\u5c3e\u75d5\u8ff9\uff0c\u5efa\u8bae\u518d\u987a\u4e00\u8f6e\u8868\u8fbe\u3002"})
    else:
        items.append({"key": "humanizer_risk", "label": "AI\u75d5\u8ff9\u98ce\u9669", "status": "pass", "detail": "\u5f53\u524d\u6ca1\u6709\u660e\u663e\u7684\u6a21\u677f\u53e5\u3001\u7a7a\u6cdb\u5347\u534e\u6216\u5ba3\u4f20\u8154\u98ce\u9669\u3002"})

    fail_count = sum(1 for item in items if item.get("status") == "fail")
    warn_count = sum(1 for item in items if item.get("status") == "warn")
    overall_status = "fail" if fail_count else "warn" if warn_count else "pass"
    return {
        "overall_status": overall_status,
        "fail_count": fail_count,
        "warn_count": warn_count,
        "items": items,
        "h2_count": h2_count,
        "expected_h2_range": expected_h2_range,
        "intro_paragraph_count": len(intro_paragraphs),
        "title_count": title_count,
    }


def detect_auto_retry_issues(
    article_text,
    explicit_title_candidates=None,
    highlighted_article="",
    require_highlight=False,
    target_words=None,
):
    report = build_publish_quality_gate_report(
        article_text,
        title_candidates=explicit_title_candidates,
        highlighted_article=highlighted_article,
        term_scan_summary={},
        target_words=target_words,
    )
    issue_keys = []
    for item in report.get("items", []):
        if item.get("key") in {"titles", "intro", "h2_count"} and item.get("status") == "fail":
            issue_keys.append(item.get("key"))
        if item.get("key") == "highlight" and require_highlight and item.get("status") != "pass":
            issue_keys.append("highlight")
    return issue_keys


def build_auto_retry_instruction(issue_keys, target_words=None, require_highlight=False):
    issue_set = set(issue_keys or [])
    min_h2, max_h2 = get_expected_h2_range(target_words)
    instructions = ["【自动补跑修正】请只修复下面这些结构问题，再重新完整输出一次。"]
    if "titles" in issue_set:
        instructions.append("- 重新补齐备选标题，必须保留 3-5 个标题。")
    if "intro" in issue_set:
        instructions.append("- 开头先补 1-2 段导语，先交代对象、事件和分析缘由，再进入正文分析。")
    if "h2_count" in issue_set:
        instructions.append(f"- 正文补足 `##` 二级标题，目标范围是 {min_h2}-{max_h2} 个，不要再写成一整堵长段。")
    if "highlight" in issue_set and require_highlight:
        instructions.append("- 请同时输出完整的高亮阅读版，不要遗漏“高亮阅读版”区块。")
    return "\n".join(instructions)


def build_auto_retry_notice(stage_name, issue_keys):
    stage_map = {
        "draft_generation": "初稿生成",
        "modification_generation": "修改稿生成",
        "de_ai_generation": "去 AI 定稿",
    }
    label_map = {
        "titles": "标题组",
        "intro": "开头导语",
        "h2_count": "二级标题结构",
        "highlight": "高亮阅读版",
    }
    issue_labels = [label_map.get(key, key) for key in issue_keys or []]
    if not issue_labels:
        return ""
    return f"自动补跑已触发：{stage_map.get(stage_name, stage_name or '当前阶段')}缺少“{'、'.join(issue_labels)}”，系统已追加修正要求并重跑 1 次。"


def resolve_active_term_rules(scope=None, respect_enabled=True):
    global_banned_terms = parse_banned_terms_text(st.session_state.get("banned_terms_text", ""))

    legacy_global_replacement_text = st.session_state.get("replacement_terms_text", "")
    legacy_article_replacement_text = st.session_state.get("article_replacement_terms_text", "")
    global_default_replacement_text = st.session_state.get("default_replacement_terms_text", legacy_global_replacement_text)
    article_default_replacement_text = st.session_state.get("article_default_replacement_terms_text", legacy_article_replacement_text)
    global_suggested_replacement_text = st.session_state.get("suggested_replacement_terms_text", "")
    article_suggested_replacement_text = st.session_state.get("article_suggested_replacement_terms_text", "")

    global_default_replacements, global_default_invalid_lines = parse_replacement_terms_text(global_default_replacement_text)
    article_banned_terms = parse_banned_terms_text(st.session_state.get("article_banned_terms_text", ""))
    article_default_replacements, article_default_invalid_lines = parse_replacement_terms_text(article_default_replacement_text)
    global_suggested_replacements, global_suggested_invalid_lines = parse_replacement_terms_text(global_suggested_replacement_text)
    article_suggested_replacements, article_suggested_invalid_lines = parse_replacement_terms_text(article_suggested_replacement_text)

    merged_banned_terms, merged_default_replacements = merge_term_rules(
        global_banned_terms,
        global_default_replacements,
        article_banned_terms,
        article_default_replacements,
    )
    merged_suggested_replacements = merge_replacement_terms(
        global_suggested_replacements,
        article_suggested_replacements,
    )
    invalid_lines = (
        global_default_invalid_lines
        + article_default_invalid_lines
        + global_suggested_invalid_lines
        + article_suggested_invalid_lines
    )
    scopes = st.session_state.get("term_rules_scope", TERM_RULE_SCOPE_DEFAULT)
    if not isinstance(scopes, list):
        scopes = list(TERM_RULE_SCOPE_DEFAULT)
    enabled = bool(st.session_state.get("term_rules_enabled", False))
    if respect_enabled and (not enabled or (scope and scope not in scopes)):
        return [], {}, {}, invalid_lines
    return merged_banned_terms, merged_default_replacements, merged_suggested_replacements, invalid_lines


OBSIDIAN_CATEGORY_WEIGHTS = {
    "00_overviews": 60,
    "05_analyses": 50,
    "03_concepts": 40,
    "01_sources": 30,
    "02_entities": 24,
    "04_topics": 22,
    "06_queries": 18,
}

OBSIDIAN_CATEGORY_LABELS = {
    "00_overviews": "综述",
    "01_sources": "来源",
    "02_entities": "实体",
    "03_concepts": "概念",
    "04_topics": "主题",
    "05_analyses": "分析",
    "06_queries": "查询",
}

OBSIDIAN_CATEGORY_ORDER = [
    "00_overviews",
    "05_analyses",
    "03_concepts",
    "01_sources",
    "02_entities",
    "04_topics",
    "06_queries",
]

QUERY_STOPWORDS_EN = {
    "about", "after", "again", "also", "among", "an", "and", "any", "are", "article",
    "because", "before", "being", "between", "block", "brief", "can", "could", "content",
    "does", "for", "from", "game", "games", "gaming", "have", "how", "into", "its",
    "just", "more", "most", "news", "over", "report", "said", "same", "sort", "than",
    "that", "the", "their", "there", "these", "they", "this", "those", "through", "under",
    "using", "video", "were", "what", "when", "where", "which", "while", "will", "with",
    "would", "your", "you",
}

QUERY_STOPWORDS_ZH = {
    "\u6e38\u620f", "\u884c\u4e1a", "\u6587\u7ae0", "\u7d20\u6750", "\u5185\u5bb9", "\u89c6\u9891", "\u4fe1\u606f", "\u76f8\u5173", "\u4eca\u5929", "\u672c\u6b21", "\u8fd9\u6b21",
    "\u8fd9\u7bc7", "\u8fd9\u4e2a", "\u90a3\u4e2a", "\u6211\u4eec", "\u4ed6\u4eec", "\u4ee5\u53ca", "\u56e0\u4e3a", "\u6240\u4ee5", "\u53ef\u4ee5", "\u5982\u679c", "\u5bf9\u4e8e",
    "\u5173\u4e8e", "\u901a\u8fc7", "\u8fdb\u884c", "\u8868\u793a", "\u8ba4\u4e3a", "\u663e\u793a", "\u5176\u4e2d", "\u5df2\u7ecf", "\u53ef\u80fd",
}

ENGLISH_SIGNAL_HINTS = {
    "abtest", "arpdau", "battlepass", "battle-pass", "casual", "cpi", "cpp", "creative",
    "event", "fail", "gacha", "hybridcasual", "hybrid-casual", "iap", "idle", "liveops",
    "ltv", "market", "match3", "merge", "meta", "midcore", "monetization", "puzzle",
    "retention", "revenue", "revive", "roi", "rpg", "season", "shop", "slg", "subgenre",
    "ua", "webshop", "web-shop",
}

LOW_SIGNAL_DOC_FILENAMES = {"index.md", "log.md"}
LOW_SIGNAL_DIRNAMES = {"07_attachments", "raw", "schema", "templates"}
LOW_SIGNAL_SECTION_MARKERS = (
    "suggested pages", "related pages", "related page", "related notes", "related sources", "see also",
    "\u63a8\u8350\u9605\u8bfb", "\u76f8\u5173\u9875\u9762", "\u76f8\u5173\u9875", "\u76f8\u5173\u9605\u8bfb", "\u76f8\u5173\u6765\u6e90", "\u53c2\u89c1",
)

DEFINITION_SECTION_MARKERS = (
    "definition", "overview", "summary", "core", "thesis", "framework", "signals", "metrics",
    "\u5b9a\u4e49", "\u6982\u5ff5", "\u662f\u4ec0\u4e48", "\u6838\u5fc3", "\u7279\u5f81", "\u673a\u5236", "\u73a9\u6cd5", "\u6846\u67b6",
    "\u7814\u7a76\u4e0e\u6307\u6807", "\u5173\u952e\u4e3b\u5f20", "\u603b\u89c8", "\u6458\u8981",
)

def reset_obsidian_context():
    st.session_state.obsidian_hits = []
    st.session_state.obsidian_research_brief = ""
    st.session_state.obsidian_retrieval_error = ""
    st.session_state.obsidian_query_terms = []
    st.session_state.obsidian_wiki_root = ""
    st.session_state.obsidian_last_indexed_at = ""
    st.session_state.obsidian_retrieval_signature = ""
    st.session_state.obsidian_influence_map = []
    st.session_state.obsidian_influence_summary = ""
    st.session_state.obsidian_influence_signature = ""
    st.session_state.evidence_map = []
    st.session_state.evidence_summary = ""
    st.session_state.evidence_signature = ""


def resolve_obsidian_wiki_root(vault_path):
    raw_path = (vault_path or "").strip()
    if not raw_path:
        return None, ""

    base_path = Path(os.path.expandvars(raw_path)).expanduser()
    candidate_paths = [
        base_path / "LLM Wiki" / "wiki",
        base_path / "wiki",
        base_path,
    ]
    for candidate in candidate_paths:
        try:
            resolved = candidate.resolve()
        except Exception:
            resolved = candidate
        if not resolved.exists() or not resolved.is_dir():
            continue
        if resolved.name.lower() == "wiki":
            return str(resolved), ""
        if (resolved / "00_overviews").exists() or (resolved / "01_sources").exists():
            return str(resolved), ""
    return None, "未找到有效的 Obsidian wiki 目录。请确认 LLM Wiki 或 LLM Wiki/wiki 路径存在。"

def read_markdown_file(path_obj):
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            with open(path_obj, "r", encoding=encoding) as f:
                return f.read(), ""
        except Exception:
            continue
    return "", f"Unsupported file encoding: {path_obj.name}"


def strip_markdown_frontmatter(text):
    if not isinstance(text, str):
        return ""
    if text.startswith("---"):
        match = re.match(r"^---\s*\n.*?\n---\s*\n?", text, flags=re.DOTALL)
        if match:
            return text[match.end():]
    return text


def split_markdown_blocks(text):
    return [block.strip() for block in re.split(r"\n\s*\n", text or "") if block.strip()]

def strip_wikilink_markup(text):
    return re.sub(r"\[\[([^\]|]+)(?:\|([^\]]+))?\]\]", lambda match: match.group(2) or match.group(1), text or "")

def is_low_signal_heading(line):
    cleaned = re.sub(r"^[#>\-*\s]+", "", (line or "")).strip().casefold()
    return any(marker in cleaned for marker in LOW_SIGNAL_SECTION_MARKERS)

def is_pure_wikilink_block(block):
    lines = [line.strip() for line in (block or "").splitlines() if line.strip()]
    if not lines:
        return True
    joined = "\n".join(lines)
    without_links = re.sub(r"\[\[[^\]]+\]\]", " ", joined)
    without_markup = re.sub(r"[`*_>#\-|/:,\uFF0C\u3002.!?\uFF1F\uFF01\s\[\]\(\)]", "", without_links)
    return len(without_markup) <= 12

def clean_obsidian_content(text, *, preserve_glossary=False):
    blocks = split_markdown_blocks(text)
    if not blocks:
        return ""
    cleaned_blocks = []
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        first_line = lines[0]
        if is_low_signal_heading(first_line):
            continue
        if is_pure_wikilink_block(block):
            continue
        if not preserve_glossary and block.count("[[") >= 3 and len(lines) <= 6:
            continue
        cleaned_blocks.append(block.strip())
    if cleaned_blocks:
        return "\n\n".join(cleaned_blocks)
    return "\n\n".join(blocks[:2])

def parse_obsidian_doc(path_obj, wiki_root):
    raw_text, read_error = read_markdown_file(path_obj)
    if read_error:
        return None, read_error
    clean_text = strip_markdown_frontmatter(raw_text).strip()
    if not clean_text:
        return None, ""
    title = path_obj.stem
    title_match = re.search(r"^\s*#\s+(.+?)\s*$", clean_text, flags=re.MULTILINE)
    if title_match:
        title = title_match.group(1).strip()
    relative_path = path_obj.resolve().relative_to(Path(wiki_root).resolve()).as_posix()
    parts = relative_path.split("/")
    category = parts[0] if parts else ""
    wikilinks = re.findall(r"\[\[([^\]]+)\]\]", clean_text)
    rel_path_lower = relative_path.casefold()
    is_glossary = rel_path_lower.endswith("00_overviews/terms-glossary.md")
    clean_content = clean_obsidian_content(clean_text, preserve_glossary=is_glossary).strip() or clean_text
    return {
        "path": str(path_obj.resolve()),
        "relative_path": relative_path,
        "title": title,
        "category": category,
        "content": clean_text,
        "clean_content": clean_content,
        "wikilinks": wikilinks,
        "is_low_signal_doc": path_obj.name.casefold() in LOW_SIGNAL_DOC_FILENAMES,
        "modified_at": datetime.fromtimestamp(path_obj.stat().st_mtime).strftime("%Y-%m-%d %H:%M"),
    }, ""

def load_obsidian_documents(wiki_root):
    root_path = Path(wiki_root)
    documents = []
    read_errors = []
    for path_obj in sorted(root_path.rglob("*.md")):
        lower_parts = {part.casefold() for part in path_obj.parts}
        if not path_obj.is_file() or any(part in LOW_SIGNAL_DIRNAMES for part in lower_parts):
            continue
        if path_obj.name.casefold() in LOW_SIGNAL_DOC_FILENAMES:
            continue
        doc, read_error = parse_obsidian_doc(path_obj, wiki_root)
        if read_error:
            read_errors.append(read_error)
            continue
        if doc and not doc.get("is_low_signal_doc"):
            documents.append(doc)
    return documents, read_errors

def normalize_query_token(token):
    cleaned = token.strip().strip("()[]{}<>.,!?;:'\"`")
    return cleaned.casefold()

def extract_query_context(source_content):
    text = (source_content or "").strip()
    english_tokens = re.findall(r"[A-Za-z][A-Za-z0-9\-\+]{2,}", text)
    chinese_tokens = re.findall("[\u4E00-\u9FFF]{2,12}", text)
    score_map = {}
    order_map = {}

    def bump(term, weight, order_idx):
        score_map[term] = score_map.get(term, 0) + weight
        order_map.setdefault(term, order_idx)

    for idx, token in enumerate(english_tokens):
        normalized = normalize_query_token(token)
        if not normalized or normalized in QUERY_STOPWORDS_EN:
            continue
        if len(normalized) < 4 and normalized not in ENGLISH_SIGNAL_HINTS:
            continue
        weight = 3 if normalized in ENGLISH_SIGNAL_HINTS or any(ch.isdigit() for ch in normalized) or "-" in normalized or "+" in normalized else 1
        bump(normalized, weight, idx)

    offset = len(english_tokens)
    for idx, token in enumerate(chinese_tokens, start=offset):
        normalized = normalize_query_token(token)
        if not normalized or normalized in QUERY_STOPWORDS_ZH:
            continue
        weight = 2 if len(normalized) >= 3 else 1
        bump(normalized, weight, idx)

    ranked_terms = sorted(score_map.items(), key=lambda item: (-item[1], order_map.get(item[0], 0), item[0]))
    return {"keywords": [term for term, _ in ranked_terms[:24]], "raw_text": text}

def load_terms_glossary(documents):
    glossary_map = {}
    glossary_doc = None
    for doc in documents:
        rel_path = doc.get("relative_path", "").casefold()
        title = doc.get("title", "").casefold()
        if rel_path.endswith("00_overviews/terms-glossary.md") or "terms-glossary" in title:
            glossary_doc = doc
            break
    if not glossary_doc:
        return glossary_map

    for line in glossary_doc.get("content", "").splitlines():
        cleaned = re.sub(r"[`*_>#\-|]", " ", line).strip()
        if not cleaned:
            continue
        paren_match = re.search(r"(.+?)[\(（]([^\)）]{2,80})[\)）]", cleaned)
        if paren_match:
            left = normalize_query_token(paren_match.group(1))
            right = normalize_query_token(paren_match.group(2))
            if left and right and left != right:
                glossary_map.setdefault(left, set()).add(right)
                glossary_map.setdefault(right, set()).add(left)
        if ":" in cleaned or "：" in cleaned:
            left, right = re.split(r"[:：]", cleaned, maxsplit=1)
            left_token = normalize_query_token(left)
            right_tokens = [normalize_query_token(part) for part in re.split(r"[,，/、]", right)]
            if left_token:
                for token in right_tokens:
                    if token and token != left_token:
                        glossary_map.setdefault(left_token, set()).add(token)
                        glossary_map.setdefault(token, set()).add(left_token)
    return glossary_map


def expand_query_terms(query_terms, glossary_map):
    expanded_terms = []
    seen_terms = set()
    def push_term(term):
        normalized = normalize_query_token(term)
        if not normalized or normalized in seen_terms:
            return
        seen_terms.add(normalized)
        expanded_terms.append(normalized)
    for term in query_terms:
        push_term(term)
        for alias in glossary_map.get(normalize_query_token(term), set()):
            push_term(alias)
    return expanded_terms


def get_obsidian_category_label(category):
    return OBSIDIAN_CATEGORY_LABELS.get(category, category or "未分类")

def count_wikilinks(text):
    return len(re.findall(r"\[\[[^\]]+\]\]", text or ""))


def is_definition_like_block(first_line, block, category):
    first_line_clean = re.sub(r"^[#>\-*\s]+", "", (first_line or "")).strip().casefold()
    block_clean = strip_wikilink_markup(block or "")
    if any(marker in first_line_clean for marker in DEFINITION_SECTION_MARKERS):
        return True
    if category == "03_concepts":
        if not first_line.startswith("#") and len(block_clean) >= 80 and count_wikilinks(block) <= 1:
            return True
        if re.search(r"\b(is|means|refers to|describes)\b", block_clean.casefold()):
            return True
        if any(marker in block_clean for marker in ("是指", "指的是", "通常指", "本质上", "核心在于")):
            return True
    return False


def score_excerpt_block(block, matched_terms, category):
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    if not lines:
        return -999
    first_line = lines[0]
    if is_low_signal_heading(first_line) or is_pure_wikilink_block(block):
        return -999

    lowered = block.casefold()
    block_score = sum(14 for term in matched_terms[:6] if term in lowered)
    wikilink_count = count_wikilinks(block)
    prose_block = strip_wikilink_markup(block)

    if wikilink_count == 0:
        block_score += 4
    elif wikilink_count == 1:
        block_score += 1
    else:
        block_score -= min(6, wikilink_count)

    if re.search(r"[。！？.!?]", prose_block):
        block_score += 2
    if first_line.startswith("#"):
        block_score -= 2
    if category == "03_concepts":
        if is_definition_like_block(first_line, block, category):
            block_score += 12
        elif wikilink_count >= 2:
            block_score -= 8
    elif category == "00_overviews" and is_definition_like_block(first_line, block, category):
        block_score += 6
    elif category == "01_sources" and any(marker in first_line.casefold() for marker in ("key claims", "summary", "takeaways", "结论", "要点", "摘要", "关键主张")):
        block_score += 8

    return block_score


def extract_best_excerpt(content, matched_terms, max_chars, *, category="", title=""):
    blocks = split_markdown_blocks(content) or [(content or "").strip()]
    chosen_block = ""
    chosen_score = -1
    fallback_block = ""
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        block_score = score_excerpt_block(block, matched_terms, category)
        if block_score <= -999:
            continue
        if not fallback_block:
            fallback_block = block
        if block_score > chosen_score:
            chosen_score = block_score
            chosen_block = block

    chosen_block = chosen_block or fallback_block
    if not chosen_block:
        return ""
    chosen_block = strip_wikilink_markup(chosen_block)
    chosen_block = re.sub(r"\n{3,}", "\n\n", chosen_block).strip()
    if title and chosen_block.casefold().startswith(title.casefold()):
        chosen_block = chosen_block[len(title):].lstrip(" :\uFF1A-\n") or chosen_block
    if len(chosen_block) > max_chars:
        chosen_block = chosen_block[:max_chars].rstrip() + "..."
    return chosen_block


def score_obsidian_docs(query_context, documents, glossary_map):
    base_terms = [normalize_query_token(term) for term in query_context.get("keywords", []) if normalize_query_token(term)]
    base_term_set = set(base_terms)
    query_terms = expand_query_terms(base_terms, glossary_map)
    scored_docs = []
    for doc in documents:
        if doc.get("is_low_signal_doc"):
            continue
        title_lower = doc.get("title", "").casefold()
        content_lower = doc.get("clean_content", doc.get("content", "")).casefold()
        path_lower = doc.get("relative_path", "").casefold()
        category = doc.get("category", "")
        score = OBSIDIAN_CATEGORY_WEIGHTS.get(category, 12)
        title_matches = []
        path_matches = []
        content_matches = []
        for term in query_terms:
            is_base_term = term in base_term_set
            if term in title_lower:
                score += 18 if is_base_term else 12
                title_matches.append(term)
            elif term in path_lower:
                score += 9 if is_base_term else 6
                path_matches.append(term)
            elif term in content_lower:
                score += 5 if is_base_term else 3
                content_matches.append(term)
        rel_path = doc.get("relative_path", "")
        rel_path_lower = rel_path.casefold()
        matched_terms = []
        seen_terms = set()
        for term in title_matches + path_matches + content_matches:
            if term in seen_terms:
                continue
            seen_terms.add(term)
            matched_terms.append(term)
        if not matched_terms:
            continue
        if not title_matches and not path_matches and len(set(content_matches)) < 2:
            continue
        if rel_path_lower.endswith("00_overviews/mobile-game-trends-2025.md") and matched_terms:
            score += 8
        if rel_path_lower.endswith("00_overviews/terms-glossary.md") and matched_terms:
            score -= 28
            if title_matches or path_matches:
                score += 2
            if len(set(content_matches)) < 3 and not title_matches and not path_matches:
                continue
            if any(term in ENGLISH_SIGNAL_HINTS for term in matched_terms[:6]):
                score += 1
        if category == "03_concepts" and title_matches:
            score += 6
        if title_matches:
            score += 4
        scored_docs.append({
            "doc": doc,
            "score": score,
            "matched_terms": matched_terms,
            "title_matches": title_matches,
            "path_matches": path_matches,
            "content_matches": content_matches,
        })
    scored_docs.sort(key=lambda item: (-item["score"], OBSIDIAN_CATEGORY_ORDER.index(item["doc"].get("category")) if item["doc"].get("category") in OBSIDIAN_CATEGORY_ORDER else 999, item["doc"].get("title", "")))
    return scored_docs

def build_obsidian_hits(scored_docs, max_hits, max_chars_per_hit):
    hits = []
    for item in scored_docs:
        if len(hits) >= max_hits:
            break
        doc = item["doc"]
        matched_terms = item.get("matched_terms", [])
        excerpt = extract_best_excerpt(
            doc.get("clean_content", doc.get("content", "")),
            matched_terms,
            max_chars_per_hit,
            category=doc.get("category", ""),
            title=doc.get("title", ""),
        )
        first_excerpt_line = excerpt.splitlines()[0] if excerpt.splitlines() else ""
        if not excerpt or is_pure_wikilink_block(excerpt) or is_low_signal_heading(first_excerpt_line):
            continue
        hits.append({
            "title": doc.get("title", ""),
            "path": doc.get("relative_path", ""),
            "category": doc.get("category", ""),
            "category_label": get_obsidian_category_label(doc.get("category", "")),
            "score": item.get("score", 0),
            "excerpt": excerpt,
            "reason": "命中词：" + "、".join(matched_terms[:5]),
            "matched_terms": matched_terms,
            "modified_at": doc.get("modified_at", ""),
        })
    return hits


def build_research_brief(hits):
    if not hits:
        return ""

    def is_glossary_hit(hit):
        return hit.get("path", "").casefold().endswith("00_overviews/terms-glossary.md")

    grouped_hits = {}
    for hit in hits:
        excerpt = hit.get("excerpt", "")
        if not excerpt or is_pure_wikilink_block(excerpt):
            continue
        first_line = excerpt.splitlines()[0] if excerpt.splitlines() else ""
        if is_low_signal_heading(first_line):
            continue
        grouped_hits.setdefault(hit.get("category", ""), []).append(hit)

    brief_parts = [
        "[知识使用说明] 以下内容来自本地 Obsidian 知识库，仅可用于背景补充、概念定义、历史案例与分析框架。如果与当前素材包冲突，以当前素材包为准。"
    ]
    heading_map = {
        "00_overviews": "[相关综述]",
        "05_analyses": "[可复用分析框架]",
        "03_concepts": "[相关概念定义]",
        "01_sources": "[相关案例与来源摘要]",
    }
    for category in ("00_overviews", "05_analyses", "03_concepts", "01_sources"):
        category_hits = grouped_hits.get(category, [])
        if not category_hits:
            continue
        if category == "00_overviews":
            category_hits = sorted(category_hits, key=lambda hit: (is_glossary_hit(hit), -hit.get("score", 0), hit.get("title", "")))
        else:
            category_hits = sorted(category_hits, key=lambda hit: (-hit.get("score", 0), hit.get("title", "")))
        brief_parts.append(heading_map[category])
        for hit in category_hits[:2]:
            line = f"- {hit.get('title', '')}：{hit.get('excerpt', '')}"
            if hit.get("matched_terms"):
                line += f"（命中词：{'、'.join(hit['matched_terms'][:4])}）"
            brief_parts.append(line)
    brief = "\n".join(brief_parts).strip()
    return brief[:3200].rstrip() + ("..." if len(brief) > 3200 else "")


def build_editor_user_content(source_content, research_brief, use_images=False):
    lead = "Below is a merged source packet. Use the attached reference images together with the text for deep analysis and synthesis:" if use_images else "Below is a merged source packet. Use the text-only source packet for deep analysis and synthesis:"
    parts = [f"{lead}\n\n{(source_content or '').strip()}"]
    if research_brief:
        parts.append(
            "Below is a research brief from your local Obsidian knowledge base. It may only be used for background, concept definitions, historical cases, and analysis frameworks. If anything conflicts with the current source packet, the current source packet wins.\n\n" + research_brief.strip()
        )
    return "\n\n================\n\n".join(part for part in parts if part.strip())


def normalize_title_candidates(title_candidates):
    if isinstance(title_candidates, str):
        raw_lines = title_candidates.splitlines()
    else:
        raw_lines = []
        for item in title_candidates or []:
            raw_lines.extend(str(item).splitlines())

    normalized = []
    seen = set()
    ignored_markers = {
        ARTICLE_TITLE_MARKER,
        ARTICLE_BODY_MARKER,
        PURE_TITLE_MARKER,
        PURE_BODY_MARKER,
        HIGHLIGHT_MARKER,
    }
    for line in raw_lines:
        cleaned = str(line).strip()
        cleaned = re.sub(r"^\s*(?:[-*#]|\d+[\.)])\s*", "", cleaned)
        cleaned = cleaned.strip(" \t\"'`[]")
        if not cleaned or cleaned in ignored_markers:
            continue
        key = cleaned.casefold()
        if key in seen:
            continue
        seen.add(key)
        normalized.append(cleaned)
    return normalized[:5]


def parse_title_candidates_block(block_text):
    return normalize_title_candidates(block_text)


def split_structured_article_sections(article_text):
    clean_text = (article_text or "").strip()
    if not clean_text:
        return [], ""

    if ARTICLE_TITLE_MARKER in clean_text:
        after_title = clean_text.split(ARTICLE_TITLE_MARKER, 1)[1]
        if ARTICLE_BODY_MARKER in after_title:
            title_block, body_block = after_title.split(ARTICLE_BODY_MARKER, 1)
            return parse_title_candidates_block(title_block), body_block.strip()

    if ARTICLE_BODY_MARKER in clean_text:
        return [], clean_text.split(ARTICLE_BODY_MARKER, 1)[1].strip()

    return [], clean_text


def build_title_candidates_block(title_candidates, marker=ARTICLE_TITLE_MARKER):
    titles = normalize_title_candidates(title_candidates)
    if not titles:
        return ""
    lines = [marker]
    lines.extend(f"{idx}. {title}" for idx, title in enumerate(titles, start=1))
    return "\n".join(lines)


def build_structured_article_text(title_candidates, article_body):
    title_block = build_title_candidates_block(title_candidates)
    clean_body = (article_body or "").strip()
    parts = []
    if title_block:
        parts.append(title_block)
    if clean_body:
        parts.append(f"{ARTICLE_BODY_MARKER}\n{clean_body}")
    return "\n\n".join(parts).strip()


def build_display_article_text(article_text, title_candidates=None):
    titles, body = split_structured_article_sections(article_text)
    resolved_titles = titles or normalize_title_candidates(title_candidates)
    clean_body = (body or article_text or "").strip()
    combined = build_structured_article_text(resolved_titles, clean_body)
    return combined or clean_body


def parse_article_generation_response(response_text, fallback_titles=None):
    titles, article_body = split_structured_article_sections(response_text)
    resolved_titles = titles or normalize_title_candidates(fallback_titles)
    clean_body = (article_body or "").strip()
    if not clean_body:
        clean_body = (response_text or "").strip()
    combined_text = build_structured_article_text(resolved_titles, clean_body)
    return resolved_titles, clean_body, combined_text or clean_body


def get_article_body_text(article_text):
    _, article_body = split_structured_article_sections(article_text)
    return (article_body or article_text or "").strip()


def build_reviewer_user_content(source_content, draft_article, research_brief="", title_candidates=None):
    resolved_titles = normalize_title_candidates(title_candidates)
    if not resolved_titles:
        resolved_titles, _ = split_structured_article_sections(draft_article)

    title_block = build_title_candidates_block(resolved_titles)
    if not title_block:
        title_block = ARTICLE_TITLE_MARKER + "\n(Missing title candidates; treat this as a structural issue.)"

    draft_body = get_article_body_text(draft_article)
    parts = [
        "Below is the merged source text (this is the only factual source for review):\n"
        + (source_content or "").strip()
        + "\n\n================\n\n"
        + "Below are the current title candidates:\n"
        + title_block
        + "\n\n================\n\n"
        + "Below is the draft body:\n"
        + draft_body
    ]
    if research_brief:
        parts.append(
            "Below is a knowledge-base research brief. It may only help you judge whether the article aligns with prior analysis frameworks. It must not be used for new fact checking or to override the source packet.\n\n"
            + research_brief.strip()
        )
    return "\n\n".join(parts)


def parse_review_actions(review_feedback):
    clean_feedback = (review_feedback or "").strip()
    if not clean_feedback:
        return []

    section_match = re.search(
        "\u3010\u4fee\u6539\u4efb\u52a1\u6e05\u5355\u3011\\s*(.*?)(?=\\n\\s*\u3010[^\u3011]+\u3011|\\Z)",
        clean_feedback,
        flags=re.S,
    )
    if not section_match:
        return []

    task_section = section_match.group(1).strip()
    if not task_section:
        return []

    def build_actions_from_blocks(block_matches, strip_pattern):
        actions = []
        for idx, match in enumerate(block_matches, start=1):
            start_idx = match.start()
            end_idx = block_matches[idx].start() if idx < len(block_matches) else len(task_section)
            block_text = task_section[start_idx:end_idx].strip()
            if not block_text:
                continue
            first_line = block_text.splitlines()[0].strip()
            title = re.sub(strip_pattern, "", first_line).strip() or f"review action {len(actions) + 1}"
            actions.append({
                "id": f"review_action_{len(actions) + 1}",
                "title": title,
                "summary": title,
                "body": block_text,
            })
        return actions

    bullet_matches = list(re.finditer(r"(?m)^(?:[-*+])\s+.+$", task_section))
    if bullet_matches:
        actions = build_actions_from_blocks(bullet_matches, r"^(?:[-*+])\s*")
        if actions:
            return actions

    numbered_matches = list(re.finditer(r"(?m)^\d+[\.)]\s+.+$", task_section))
    if numbered_matches:
        actions = build_actions_from_blocks(numbered_matches, r"^\d+[\.)]\s*")
        if actions:
            return actions

    return []


def build_selected_review_feedback(review_actions, accepted_review_items):
    accepted_ids = set(accepted_review_items or [])
    selected_blocks = [
        (action.get("body") or "").strip()
        for action in review_actions or []
        if action.get("id") in accepted_ids and (action.get("body") or "").strip()
    ]
    if not selected_blocks:
        return ""
    return "\u3010\u4fee\u6539\u4efb\u52a1\u6e05\u5355\u3011\n" + "\n\n".join(selected_blocks)


def hydrate_review_action_state(review_feedback, reset_selection=False):
    review_actions = parse_review_actions(review_feedback)
    st.session_state.review_actions = review_actions

    if reset_selection:
        for key in list(st.session_state.keys()):
            if isinstance(key, str) and key.startswith("review_action_pick_"):
                del st.session_state[key]

    action_ids = [action.get("id") for action in review_actions if action.get("id")]
    current_selection = st.session_state.get("accepted_review_items", [])
    if reset_selection or not isinstance(current_selection, list):
        st.session_state.accepted_review_items = action_ids
        return

    st.session_state.accepted_review_items = [
        action_id for action_id in current_selection
        if action_id in action_ids
    ]


def build_modification_user_content(review_feedback, draft_article, title_candidates=None):
    resolved_titles = normalize_title_candidates(title_candidates)
    if not resolved_titles:
        resolved_titles, _ = split_structured_article_sections(draft_article)

    title_block = build_title_candidates_block(resolved_titles)
    if not title_block:
        title_block = ARTICLE_TITLE_MARKER + "\n1. ...\n2. ...\n3. ..."

    article_body = get_article_body_text(draft_article)
    return (
        "[Current title candidates]\n"
        + title_block
        + "\n\n================\n\n"
        + "[Review feedback to apply]\n"
        + (review_feedback or "").strip()
        + "\n\n================\n\n"
        + "[Current article body]\n"
        + article_body
        + "\n\n[Editing requirements]\n"
        + "Only apply the review feedback shown above. Do not absorb review suggestions that are not included in that block.\n"
        + "Keep the title-candidates block alive. If the direction has not changed, preserve the current title skeleton and only improve wording, precision, and communication power."
    ).strip()


def build_chat_knowledge_context():
    research_brief = (st.session_state.get("obsidian_research_brief", "") or "").strip()
    hits = st.session_state.get("obsidian_hits", []) or []
    if not research_brief and not hits:
        return ""
    parts = []
    if research_brief:
        parts.append(f"[Obsidian 研究摘要]\n{research_brief}")
    if hits:
        hit_lines = [f"- {hit.get('title', '')} | {hit.get('category_label', '')} | {hit.get('path', '')}" for hit in hits[:8]]
        parts.append("[命中笔记清单]\n" + "\n".join(hit_lines))
    return "\n\n".join(parts)


    return "\n\n".join(parts)


def split_article_paragraphs(text):
    return [part.strip() for part in re.split(r"\n\s*\n", text or "") if part.strip()]


def extract_obsidian_signal_terms(text):
    terms = []
    seen = set()
    english_tokens = re.findall(r"[A-Za-z][A-Za-z0-9\-\+]{2,}", text or "")
    chinese_tokens = re.findall("[\u4E00-\u9FFF]{2,12}", text or "")
    for token in english_tokens:
        normalized = normalize_query_token(token)
        if not normalized or normalized in QUERY_STOPWORDS_EN:
            continue
        if len(normalized) < 4 and normalized not in ENGLISH_SIGNAL_HINTS:
            continue
        if normalized not in seen:
            seen.add(normalized)
            terms.append(normalized)
    for token in chinese_tokens:
        normalized = normalize_query_token(token)
        if not normalized or normalized in QUERY_STOPWORDS_ZH:
            continue
        if normalized not in seen:
            seen.add(normalized)
            terms.append(normalized)
    return terms[:24]


def score_paragraph_against_obsidian_hit(paragraph, hit):
    paragraph_terms = set(extract_obsidian_signal_terms(paragraph))
    if not paragraph_terms:
        return 0, []

    candidate_terms = []
    for term in hit.get("matched_terms", [])[:8]:
        normalized = normalize_query_token(term)
        if normalized:
            candidate_terms.append(normalized)
    candidate_terms.extend(extract_obsidian_signal_terms(hit.get("title", ""))[:6])
    candidate_terms.extend(extract_obsidian_signal_terms(hit.get("excerpt", ""))[:10])

    deduped_terms = []
    seen = set()
    for term in candidate_terms:
        if term and term not in seen:
            seen.add(term)
            deduped_terms.append(term)

    matched_terms = [term for term in deduped_terms if term in paragraph_terms]
    score = len(matched_terms) * 3
    if hit.get("title") and any(term in paragraph_terms for term in extract_obsidian_signal_terms(hit.get("title", ""))[:4]):
        score += 2
    if any(term in paragraph.casefold() for term in hit.get("matched_terms", [])[:4]):
        score += 2
    return score, matched_terms[:6]


def build_obsidian_influence_map(final_article, hits):
    paragraphs = split_article_paragraphs(final_article)
    influence_items = []
    for idx, paragraph in enumerate(paragraphs, start=1):
        scored_hits = []
        for hit in hits or []:
            score, matched_terms = score_paragraph_against_obsidian_hit(paragraph, hit)
            if score < 5:
                continue
            scored_hits.append({
                "title": hit.get("title", ""),
                "score": score,
                "matched_terms": matched_terms,
            })
        if not scored_hits:
            continue
        scored_hits.sort(key=lambda item: (-item["score"], item["title"]))
        top_hits = scored_hits[:3]
        top_score = top_hits[0]["score"]
        preview = paragraph if len(paragraph) <= 120 else paragraph[:120].rstrip() + "..."
        merged_terms = []
        seen_terms = set()
        for item in top_hits:
            for term in item.get("matched_terms", []):
                if term not in seen_terms:
                    seen_terms.add(term)
                    merged_terms.append(term)
        influence_items.append({
            "paragraph_index": idx,
            "paragraph_preview": preview,
            "influence_level": "高" if top_score >= 10 else "中",
            "score": top_score,
            "note_titles": [item.get("title", "") for item in top_hits if item.get("title")],
            "matched_terms": merged_terms[:6],
        })
    return influence_items


def parse_source_packet_segments(source_content):
    if not (source_content or "").strip():
        return []

    type_meta = {
        "\u6587\u7ae0\u7d20\u6750": ("article", "\u6765\u6e90\u4e8e"),
        "\u89c6\u9891\u7d20\u6750": ("video", "\u6765\u6e90\u4e8e"),
        "\u4e0a\u4f20\u6587\u4ef6\u7d20\u6750": ("upload_text", "\u6587\u4ef6\u540d"),
        "\u4e0a\u4f20\u56fe\u7247\u7d20\u6750": ("upload_image", "\u6587\u4ef6\u540d"),
    }
    segments = []
    blocks = [block.strip() for block in re.split(r"\n\s*=+\s*\n", source_content or "") if block.strip()]
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        header = lines[0]
        match = re.match(
            r"^\u3010(?P<label_type>\u6587\u7ae0\u7d20\u6750|\u89c6\u9891\u7d20\u6750|\u4e0a\u4f20\u6587\u4ef6\u7d20\u6750|\u4e0a\u4f20\u56fe\u7247\u7d20\u6750)\s*(?P<index>\d+)\u3011\s*(?P<meta_key>\u6765\u6e90\u4e8e|\u6587\u4ef6\u540d)\s*:\s*(?P<meta_value>.+)$",
            header,
        )
        if not match:
            continue
        label_type = match.group("label_type")
        source_type, expected_meta_key = type_meta[label_type]
        meta_key = match.group("meta_key")
        source_locator = match.group("meta_value").strip()
        if meta_key != expected_meta_key:
            continue
        body_text = "\n".join(lines[1:]).strip()
        paragraphs = split_article_paragraphs(body_text)
        segments.append({
            "source_type": source_type,
            "source_label": f"{label_type} {match.group('index')}",
            "source_locator": source_locator,
            "raw_block": block,
            "paragraphs": paragraphs,
        })
    return segments


def extract_claim_sentences(paragraph):
    paragraph = (paragraph or "").strip()
    if not paragraph:
        return []
    sentences = [segment.strip() for segment in re.split(r"(?<=[\u3002\uff01\uff1f!?\uff1b;])\s*|\n+", paragraph) if segment.strip()]
    if not sentences:
        return []

    markers = (
        "\u662f", "\u5c06", "\u610f\u5473\u7740", "\u8868\u660e", "\u663e\u793a", "\u8bf4\u660e", "\u6210\u4e3a", "\u63a8\u52a8", "\u5e26\u6765", "\u5f15\u53d1",
        "not only", "rather than", "compared", "suggests", "shows", "indicates",
    )
    comparators = ("\u66f4", "\u8f83", "\u6bd4", "\u4f4e\u4e8e", "\u9ad8\u4e8e", "\u8d85\u8fc7", "\u81f3\u5c11", "\u6700\u591a", "\u9996\u6b21", "\u7ee7\u7eed", "\u5347\u7ea7", "\u4e0b\u964d")
    scored = []
    for idx, sentence in enumerate(sentences):
        score = 0
        if re.search(r"\d", sentence):
            score += 4
        if any(token in sentence for token in comparators):
            score += 3
        if any(token in sentence.lower() for token in markers):
            score += 2
        length = len(sentence)
        if 10 <= length <= 70:
            score += 2
        elif length > 70:
            score += 1
        if idx == 0:
            score += 1
        scored.append((score, idx, sentence))
    scored.sort(key=lambda item: (-item[0], item[1]))

    selected = []
    seen = set()
    for _, _, sentence in scored:
        if sentence not in seen:
            seen.add(sentence)
            selected.append(sentence)
        if len(selected) >= 2:
            break
    return selected or [sentences[0]]


def score_paragraph_against_source_segment(paragraph, segment):
    if not paragraph or not isinstance(segment, dict):
        return 0, [], ""
    if segment.get("source_type") == "upload_image":
        return 0, [], ""

    paragraph_terms = set(extract_obsidian_signal_terms(paragraph))
    if not paragraph_terms:
        return 0, [], ""

    candidate_terms = []
    candidate_terms.extend(extract_obsidian_signal_terms(segment.get("source_label", ""))[:4])
    candidate_terms.extend(extract_obsidian_signal_terms(segment.get("source_locator", ""))[:4])
    for source_paragraph in segment.get("paragraphs", [])[:8]:
        candidate_terms.extend(extract_obsidian_signal_terms(source_paragraph)[:10])

    deduped_terms = []
    seen_terms = set()
    for term in candidate_terms:
        if term and term not in seen_terms:
            seen_terms.add(term)
            deduped_terms.append(term)
    matched_terms = [term for term in deduped_terms if term in paragraph_terms]

    claim_sentences = extract_claim_sentences(paragraph)
    paragraph_numbers = set(re.findall(r"\d+(?:\.\d+)?", paragraph))
    best_excerpt = ""
    best_overlap = 0
    for source_paragraph in segment.get("paragraphs", []) or [segment.get("raw_block", "")]:
        overlap = 0
        normalized_source = source_paragraph.casefold()
        for claim in claim_sentences:
            claim_lower = claim.casefold()
            if claim_lower and claim_lower in normalized_source:
                overlap += max(6, min(len(claim) // 5, 10))
        source_terms = set(extract_obsidian_signal_terms(source_paragraph))
        overlap += len(paragraph_terms & source_terms) * 2
        number_overlap = paragraph_numbers & set(re.findall(r"\d+(?:\.\d+)?", source_paragraph))
        overlap += len(number_overlap) * 3
        if overlap > best_overlap:
            best_overlap = overlap
            best_excerpt = source_paragraph.strip()

    score = len(matched_terms) * 3 + best_overlap
    source_locator = (segment.get("source_locator", "") or "").casefold()
    if any(term in source_locator for term in matched_terms[:3]):
        score += 1
    return score, matched_terms[:8], best_excerpt


def build_article_evidence_map(final_article, source_segments, obsidian_hits):
    paragraphs = split_article_paragraphs(final_article)
    evidence_items = []
    for idx, paragraph in enumerate(paragraphs, start=1):
        claim_sentences = extract_claim_sentences(paragraph)
        support_items = []

        for segment in source_segments or []:
            score, matched_terms, excerpt = score_paragraph_against_source_segment(paragraph, segment)
            if score < 8:
                continue
            support_items.append({
                "source_type": "source_packet",
                "source_label": segment.get("source_label", ""),
                "source_locator": segment.get("source_locator", ""),
                "matched_excerpt": excerpt[:180],
                "matched_terms": matched_terms[:6],
                "score": score,
                "confidence": "\u9ad8" if score >= 14 else "\u4e2d" if score >= 10 else "\u4f4e",
            })

        for hit in obsidian_hits or []:
            score, matched_terms = score_paragraph_against_obsidian_hit(paragraph, hit)
            if score < 5:
                continue
            support_items.append({
                "source_type": "obsidian",
                "source_label": hit.get("title", ""),
                "source_locator": hit.get("path", ""),
                "matched_excerpt": (hit.get("excerpt", "") or "")[:180],
                "matched_terms": matched_terms[:6],
                "score": score,
                "confidence": "\u9ad8" if score >= 10 else "\u4e2d" if score >= 7 else "\u4f4e",
            })

        if not support_items:
            continue
        support_items.sort(key=lambda item: (-item.get("score", 0), item.get("source_label", "")))
        if support_items[0].get("score", 0) < 6:
            continue
        top_items = support_items[:3]
        preview = paragraph if len(paragraph) <= 120 else paragraph[:120].rstrip() + "..."
        top_score = top_items[0].get("score", 0)
        source_types = []
        for item in top_items:
            source_type = item.get("source_type")
            if source_type and source_type not in source_types:
                source_types.append(source_type)
        evidence_items.append({
            "paragraph_index": idx,
            "paragraph_preview": preview,
            "claim_sentences": claim_sentences[:2],
            "support_items": top_items,
            "coverage_level": "\u9ad8" if top_score >= 14 else "\u4e2d" if top_score >= 9 else "\u4f4e",
            "source_types": source_types,
        })
    return evidence_items


def refresh_evidence_map(force=False):
    final_article = get_article_body_text((st.session_state.get("final_article", "") or "").strip())
    source_content = (st.session_state.get("source_content", "") or "").strip()
    hits = st.session_state.get("obsidian_hits", []) or []
    if not final_article or not source_content:
        st.session_state.evidence_map = []
        st.session_state.evidence_summary = ""
        st.session_state.evidence_signature = ""
        return

    signature_base = final_article + "\n" + source_content + "\n" + json.dumps(
        [
            {
                "title": hit.get("title", ""),
                "path": hit.get("path", ""),
                "excerpt": hit.get("excerpt", ""),
                "matched_terms": hit.get("matched_terms", []),
            }
            for hit in hits
        ],
        ensure_ascii=False,
        sort_keys=True,
    )
    signature = hashlib.sha1(signature_base.encode("utf-8")).hexdigest()
    if not force and st.session_state.get("evidence_signature") == signature:
        return

    source_segments = parse_source_packet_segments(source_content)
    evidence_map = build_article_evidence_map(final_article, source_segments, hits)
    high_coverage_count = sum(1 for item in evidence_map if item.get("coverage_level") == "\u9ad8")
    obsidian_only_count = sum(1 for item in evidence_map if item.get("source_types") == ["obsidian"])
    st.session_state.evidence_map = evidence_map
    st.session_state.evidence_summary = (
        f"\u547d\u4e2d\u6bb5\u843d {len(evidence_map)} \u4e2a\uff1b\u9ad8\u8986\u76d6 {high_coverage_count} \u4e2a\uff1b\u4ec5\u6709 Obsidian \u8865\u5145 {obsidian_only_count} \u4e2a\u3002"
        if evidence_map else
        "\u5f53\u524d\u5b9a\u7a3f\u91cc\u8fd8\u6ca1\u6709\u627e\u5230\u8fbe\u5230\u9608\u503c\u7684\u8bba\u636e\u5361\u7247\u3002"
    )
    st.session_state.evidence_signature = signature


def normalize_evidence_excerpt(text, max_chars=220):
    raw_text = (text or "").strip()
    if not raw_text:
        return ""

    normalized_lines = []
    for line in raw_text.splitlines():
        clean_line = line.strip()
        if not clean_line:
            continue
        clean_line = re.sub(r"^\s{0,3}#{1,6}\s*", "", clean_line)
        clean_line = re.sub(r"^\s*>\s*", "", clean_line)
        bullet_match = re.match(r"^\s*[-*+]\s+(.*)$", clean_line)
        numbered_match = re.match(r"^\s*\d+[.)]\s+(.*)$", clean_line)
        if bullet_match:
            clean_line = "- " + bullet_match.group(1).strip()
        elif numbered_match:
            clean_line = "- " + numbered_match.group(1).strip()
        clean_line = re.sub(r"\s+", " ", clean_line).strip()
        if clean_line:
            normalized_lines.append(clean_line)

    if not normalized_lines:
        return ""
    excerpt = "\n".join(normalized_lines[:4])
    if len(excerpt) > max_chars:
        excerpt = excerpt[:max_chars].rstrip() + "..."
    return excerpt


def render_evidence_support_cards(items, empty_text, tone="source"):
    if not items:
        st.caption(empty_text)
        return

    accent = "#0f766e" if tone == "source" else "#64748b"
    surface = "#f4fbf9" if tone == "source" else "#f7f8fb"
    pill_bg = "rgba(15, 118, 110, 0.12)" if tone == "source" else "rgba(100, 116, 139, 0.12)"
    pill_text = accent

    for support in items:
        title = html_lib.escape((support.get("source_label", "") or "").strip() or "\u672a\u547d\u540d\u6765\u6e90")
        confidence = html_lib.escape((support.get("confidence", "\u4f4e") or "\u4f4e").strip())
        locator = html_lib.escape((support.get("source_locator", "") or "").strip())
        excerpt = normalize_evidence_excerpt(support.get("matched_excerpt", ""))
        excerpt = html_lib.escape(excerpt or "\u5df2\u547d\u4e2d\u8be5\u6761\u8bba\u636e\uff0c\u4f46\u5f53\u524d\u6ca1\u6709\u53ef\u5c55\u793a\u7684\u6458\u5f55\u3002")
        matched_terms = [term for term in support.get("matched_terms", [])[:6] if (term or "").strip()]
        terms_html = ""
        if matched_terms:
            pills = "".join(
                f'<span style="display:inline-block;margin:0 0.35rem 0.35rem 0;padding:0.18rem 0.52rem;border-radius:999px;background:{pill_bg};color:{pill_text};font-size:0.78rem;line-height:1.2;">{html_lib.escape(str(term))}</span>'
                for term in matched_terms
            )
            terms_html = (
                '<div style="margin-top:0.65rem;">'
                '<div style="font-size:0.78rem;color:#6b7280;margin-bottom:0.35rem;">\u547d\u4e2d\u8bcd</div>'
                f'{pills}'
                '</div>'
            )

        locator_html = ""
        if locator:
            locator_html = (
                '<div style="margin-top:0.5rem;font-family:ui-monospace,SFMono-Regular,Consolas,monospace;'
                'font-size:0.76rem;line-height:1.5;color:#64748b;word-break:break-all;">'
                f'{locator}'
                '</div>'
            )

        card_html = f"""<div style=\"margin:0 0 0.85rem 0;padding:0.9rem 0.95rem;border:1px solid #e6ebf0;border-radius:16px;background:{surface};\">
<div style=\"display:flex;justify-content:space-between;gap:0.8rem;align-items:flex-start;\">
  <div style=\"font-size:0.98rem;font-weight:700;line-height:1.45;color:#0f172a;\">{title}</div>
  <div style=\"flex-shrink:0;padding:0.18rem 0.48rem;border-radius:999px;background:{pill_bg};color:{accent};font-size:0.75rem;font-weight:600;\">{confidence}\u7f6e\u4fe1</div>
</div>
{locator_html}
{terms_html}
<div style=\"margin-top:0.7rem;padding:0.75rem 0.8rem;border-radius:12px;background:#ffffff;border-left:3px solid {accent};font-size:0.92rem;line-height:1.72;color:#475569;white-space:pre-wrap;\">{excerpt}</div>
</div>"""
        st.markdown(card_html, unsafe_allow_html=True)


def render_evidence_map_panel(evidence_map, summary):
    render_section_intro("\u8bba\u636e\u6765\u6e90", "\u628a\u6b63\u6587\u91cc\u7684\u5173\u952e\u5224\u65ad\u6620\u5c04\u56de\u7d20\u6750\u4f9d\u636e\u548c Obsidian \u8865\u5145\uff0c\u5e2e\u52a9\u5feb\u901f\u56de\u770b\u51fa\u5178\u3002", "\u8bba\u636e")
    if summary:
        st.caption(summary)
    if not evidence_map:
        st.info("\u5f53\u524d\u5b9a\u7a3f\u91cc\u8fd8\u6ca1\u6709\u627e\u5230\u8fbe\u5230\u9608\u503c\u7684\u8bba\u636e\u5361\u7247\u3002")
        return

    for item in evidence_map:
        label = f"\u7b2c {item.get('paragraph_index', 0)} \u6bb5 \u00b7 {item.get('coverage_level', '\\u4f4e')}\u8986\u76d6"
        support_items = item.get("support_items", [])
        source_packet_items = [support for support in support_items if support.get("source_type") == "source_packet"]
        obsidian_items = [support for support in support_items if support.get("source_type") == "obsidian"]
        preview = html_lib.escape((item.get("paragraph_preview", "") or "").strip())
        with st.expander(label, expanded=(item.get("paragraph_index", 0) <= 2)):
            if preview:
                preview_html = f'<div style="padding:0.75rem 0.82rem;border-radius:12px;background:#f8fafc;border:1px solid #e7edf3;color:#475569;font-size:0.93rem;line-height:1.72;">{preview}</div>'
                st.markdown(preview_html, unsafe_allow_html=True)
            claims = item.get("claim_sentences", [])
            if claims:
                st.markdown("**\u5173\u952e\u5224\u65ad**")
                for claim in claims:
                    st.markdown(f"- {(claim or '').strip()}")
            st.markdown("**\u7d20\u6750\u4f9d\u636e**")
            render_evidence_support_cards(
                source_packet_items,
                "\u8fd9\u4e00\u6bb5\u6682\u65f6\u6ca1\u6709\u627e\u5230\u8db3\u591f\u5f3a\u7684\u539f\u59cb\u7d20\u6750\u4f9d\u636e\u3002",
                tone="source",
            )

            st.markdown("**Obsidian \u8865\u5145**")
            st.caption("\u80cc\u666f\u8865\u5145\uff0c\u4e0d\u7b49\u540c\u4e8e\u539f\u59cb\u4e8b\u5b9e\u6765\u6e90\u3002")
            render_evidence_support_cards(
                obsidian_items,
                "\u8fd9\u4e00\u6bb5\u6ca1\u6709\u989d\u5916\u547d\u4e2d Obsidian \u8865\u5145\u3002",
                tone="obsidian",
            )


def refresh_obsidian_influence_map(force=False):
    final_article = get_article_body_text((st.session_state.get("final_article", "") or "").strip())
    hits = st.session_state.get("obsidian_hits", []) or []
    if not st.session_state.get("obsidian_enabled") or not final_article or not hits:
        st.session_state.obsidian_influence_map = []
        st.session_state.obsidian_influence_summary = ""
        st.session_state.obsidian_influence_signature = ""
        return
    signature_base = final_article + "\n" + json.dumps(
        [{"title": hit.get("title", ""), "excerpt": hit.get("excerpt", ""), "matched_terms": hit.get("matched_terms", [])} for hit in hits],
        ensure_ascii=False,
        sort_keys=True,
    )
    signature = hashlib.sha1(signature_base.encode("utf-8")).hexdigest()
    if not force and st.session_state.get("obsidian_influence_signature") == signature:
        return
    influence_map = build_obsidian_influence_map(final_article, hits)
    st.session_state.obsidian_influence_map = influence_map
    st.session_state.obsidian_influence_summary = (
        f"共发现 {len(influence_map)} 个段落明显受到 Obsidian 补充内容影响。"
        if influence_map else
        "当前定稿中暂无达到阈值的 Obsidian 影响段落。"
    )
    st.session_state.obsidian_influence_signature = signature


def render_obsidian_influence_panel(influence_map, summary):
    render_section_intro("Obsidian 影响地图", "按段落查看当前定稿中哪些内容明显受本地知识库补充影响。", "影响")
    if summary:
        st.caption(summary)
    if not influence_map:
        st.info("当前定稿里暂无达到阈值的 Obsidian 影响段落。")
        return
    for item in influence_map:
        note_titles = item.get("note_titles", [])
        label = f"第 {item.get('paragraph_index', 0)} 段 · {item.get('influence_level', '中')}影响"
        with st.expander(label, expanded=(item.get("paragraph_index", 0) <= 2)):
            if note_titles:
                st.caption("相关笔记：" + " / ".join(note_titles[:3]))
            if item.get("matched_terms"):
                st.caption("命中关键词：" + "、".join(item.get("matched_terms", [])[:6]))
            st.code(item.get("paragraph_preview", ""), language="markdown")


def run_obsidian_retrieval(force=False):
    if not st.session_state.get("obsidian_enabled"):
        reset_obsidian_context()
        refresh_evidence_map(force=True)
        return
    source_content = (st.session_state.get("source_content", "") or "").strip()
    if not source_content:
        reset_obsidian_context()
        refresh_evidence_map(force=True)
        return
    wiki_root, resolve_error = resolve_obsidian_wiki_root(st.session_state.get("obsidian_vault_path", ""))
    if resolve_error:
        reset_obsidian_context()
        st.session_state.obsidian_retrieval_error = resolve_error
        refresh_evidence_map(force=True)
        save_draft()
        return
    max_hits = int(st.session_state.get("obsidian_max_hits", 6) or 6)
    signature_base = f"{wiki_root}|{max_hits}|{hashlib.sha1(source_content.encode('utf-8')).hexdigest()}"
    signature = hashlib.sha1(signature_base.encode("utf-8")).hexdigest()
    has_cached_result = bool(st.session_state.get("obsidian_hits") or st.session_state.get("obsidian_research_brief") or st.session_state.get("obsidian_retrieval_error"))
    if not force and has_cached_result and st.session_state.get("obsidian_retrieval_signature") == signature:
        return
    documents, read_errors = load_obsidian_documents(wiki_root)
    query_context = extract_query_context(source_content)
    glossary_map = load_terms_glossary(documents)
    scored_docs = score_obsidian_docs(query_context, documents, glossary_map)
    hits = build_obsidian_hits(scored_docs, max_hits=max_hits, max_chars_per_hit=280)
    st.session_state.obsidian_hits = hits
    st.session_state.obsidian_research_brief = build_research_brief(hits)
    st.session_state.obsidian_query_terms = query_context.get("keywords", [])[:12]
    st.session_state.obsidian_wiki_root = wiki_root
    st.session_state.obsidian_last_indexed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.session_state.obsidian_retrieval_signature = signature
    st.session_state.obsidian_retrieval_error = "部分 Obsidian 笔记读取失败，已跳过异常文件。" if read_errors else ""
    refresh_obsidian_influence_map(force=True)
    refresh_evidence_map(force=True)
    save_draft()

def build_openai_timeout(base_url, model_name):
    base_url = (base_url or "").strip().lower()
    model_name = (model_name or "").strip().lower()

    # Yunwu Qwen models have shown billed-but-disconnected responses, so use a wider timeout window.
    if "yunwu.ai/v1" in base_url and model_name in {"qwen3.5-plus", "qwen3.6-plus"}:
        return httpx.Timeout(connect=30.0, read=1800.0, write=120.0, pool=120.0)

    return httpx.Timeout(connect=20.0, read=600.0, write=60.0, pool=60.0)


def build_openai_http_client(timeout_config):
    return httpx.Client(
        timeout=timeout_config,
        trust_env=False,
        http2=False,
        follow_redirects=True,
        limits=httpx.Limits(max_connections=20, max_keepalive_connections=0),
    )


def should_retry_with_requests_fallback(exc):
    error_text = format_llm_exception(exc).lower()
    retry_markers = [
        "apiconnectionerror",
        "remoteprotocolerror",
        "server disconnected without sending a response",
        "remote end closed connection without response",
        "connection aborted",
        "connection reset",
        "connection refused",
        "connection error",
        "timed out",
        "timeout",
        "temporarily unavailable",
        "502",
        "503",
        "504",
    ]
    return any(marker in error_text for marker in retry_markers)


def build_requests_timeout(timeout_config):
    connect_timeout = getattr(timeout_config, "connect", None) or 20.0
    read_timeout = getattr(timeout_config, "read", None) or 600.0
    return (connect_timeout, read_timeout)


def extract_llm_response_content(response_data, empty_choices_message, empty_body_message):
    choices = response_data.get("choices") if isinstance(response_data, dict) else None
    if not choices:
        raise ValueError(empty_choices_message)

    first_choice = choices[0] if isinstance(choices[0], dict) else {}
    message = first_choice.get("message", {}) if isinstance(first_choice, dict) else {}
    content = normalize_llm_response_content(message.get("content"))
    if not content.strip():
        raise ValueError(empty_body_message)
    return content


def call_llm_via_requests(api_key, base_url, request_payload, timeout_config):
    endpoint = (base_url or "").rstrip("/") + "/chat/completions"
    session = requests.Session()
    session.trust_env = False
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "User-Agent": "Article-Transcription-Assistant/1.0",
    }

    try:
        response = session.post(
            endpoint,
            headers=headers,
            json=request_payload,
            timeout=build_requests_timeout(timeout_config),
        )
        response.raise_for_status()
        return extract_llm_response_content(
            response.json(),
            "Fallback request returned no choices.",
            "Fallback request returned an empty message body.",
        )
    finally:
        session.close()


def call_llm_via_httpx_raw(api_key, base_url, request_payload, timeout_config):
    endpoint = (base_url or "").rstrip("/") + "/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "User-Agent": "Article-Transcription-Assistant/1.0",
    }

    with build_openai_http_client(timeout_config) as http_client:
        response = http_client.post(endpoint, headers=headers, json=request_payload)
        response.raise_for_status()
        return extract_llm_response_content(
            response.json(),
            "Raw HTTP fallback returned no choices.",
            "Raw HTTP fallback returned an empty message body.",
        )


def execute_llm_request_once(api_key, base_url, model_name, request_kwargs, timeout_config, use_stream):
    transport_errors = []

    try:
        with build_openai_http_client(timeout_config) as raw_http_client:
            client = OpenAI(
                api_key=api_key,
                base_url=base_url,
                max_retries=0,
                timeout=timeout_config,
                http_client=raw_http_client,
            )

            if use_stream:
                parts = []
                with client.chat.completions.create(stream=True, **request_kwargs) as stream:
                    for chunk in stream:
                        if not getattr(chunk, "choices", None):
                            continue
                        for choice in chunk.choices:
                            delta = getattr(choice, "delta", None)
                            delta_content = getattr(delta, "content", None) if delta else None
                            if delta_content:
                                parts.append(delta_content)
                content = "".join(parts).strip()
            else:
                response = client.chat.completions.create(**request_kwargs)
                if not getattr(response, "choices", None):
                    raise ValueError("Model returned no choices.")
                message = response.choices[0].message if response.choices else None
                content = normalize_llm_response_content(getattr(message, "content", None))

        if content.strip():
            return content
        raise ValueError("Model returned an empty message body.")
    except Exception as sdk_error:
        transport_errors.append(f"sdk={format_llm_exception(sdk_error)}")
        if not should_retry_with_requests_fallback(sdk_error):
            raise

    for transport_name, transport_fn in (
        ("requests", call_llm_via_requests),
        ("httpx_raw", call_llm_via_httpx_raw),
    ):
        try:
            return transport_fn(api_key, base_url, request_kwargs, timeout_config)
        except Exception as transport_error:
            transport_errors.append(f"{transport_name}={format_llm_exception(transport_error)}")
            if not should_retry_with_requests_fallback(transport_error):
                raise

    raise RuntimeError("All transports failed: " + " | ".join(transport_errors))


def should_stream_llm_request(base_url, model_name):
    base_url = (base_url or "").strip().lower()
    model_name = (model_name or "").strip().lower()
    return "yunwu.ai/v1" in base_url and model_name in {"qwen3.5-plus", "qwen3.6-plus"}


BLTCY_MODEL_OPTIONS = [
    "qwen3.5-plus",
    "kimi-k2.5",
    "gpt-5.4",
    "claude-opus-4-6",
    "gpt-5.4-mini-2026-03-17",
    "gpt-5.4-nano",
    "gemini-3.1-pro-preview",
    "claude-sonnet-4-6-thinking",
    "claude-opus-4-6-thinking",
    "claude-opus-4-5-20251101-thinking",
    "gemini-3.1-pro-preview-thinking-high",
    "gemini-3.1-flash-lite-preview-thinking-high",
    "MiniMax-M2.7",
    "MiniMax-M2.7-highspeed"
]
DEERAPI_MODEL_OPTIONS = [
    "gpt-5.4",
    "qwen3.5-27b",
    "qwen3.5-flash",
    "gpt-5.4-nano",
    "gemini-3.1-pro-preview",
    "gemini-3.1-pro-preview-thinking",
    "gemini-3.1-flash-lite",
    "gemini-3.1-flash-lite-preview-thinking"
]
YUNWU_MODEL_OPTIONS = [
    "qwen3.6-plus",
    "qwen3.5-plus",
    "deepseek-v4-flash",
    "deepseek-v4-pro",
    "kimi-k2.5",
    "gpt-5.4",
    "claude-opus-4-6",
    "doubao-seed-2-0-lite-260215",
    "gpt-5.4-mini-2026-03-17",
    "gpt-5.4-nano",
    "gemini-3.1-pro-preview",
    "claude-sonnet-4-6-thinking",
    "claude-opus-4-6-thinking",
    "claude-opus-4-5-20251101-thinking",
    "gemini-3.1-pro-preview-thinking-high",
    "gemini-3.1-flash-lite-preview-thinking-high",
    "MiniMax-M2.7",
    "MiniMax-M2.7-highspeed"
]
DE_AI_MODEL_MIGRATION = {
    "deepseek-v3-1-terminus": "deepseek-v3.1",
    "deepseek-v3-2-exp": "deepseek-v3.2",
}
DE_AI_MODELS = ["deepseek-v3.1", "deepseek-v3.2", "deepseek-v4-flash", "deepseek-v4-pro", "qwen3.5-plus", "glm-5"]
DE_AI_VARIANTS = ["普通版", "社区文章去AI版", "自然唠嗑版", "Humanizer-zh 版"]
DE_AI_VARIANT_DEFAULT = DE_AI_VARIANTS[0]
DE_AI_VARIANT_COMMUNITY = DE_AI_VARIANTS[1]
DE_AI_VARIANT_CHAT = DE_AI_VARIANTS[2]
DE_AI_VARIANT_HUMANIZER = DE_AI_VARIANTS[3]

TERM_RULE_SCOPE_LABELS = {
    "modification": "修改稿",
    "de_ai": "去AI",
    "final_check": "最终检查",
}
TERM_RULE_SCOPE_DEFAULT = ["modification", "de_ai", "final_check"]
TERM_RULE_DEFAULT_ENABLED = True
TERM_RULE_DEFAULT_BANNED_TERMS_TEXT = "\n".join([
    "综上所述",
    "总而言之",
    "毋庸置疑",
    "毫无疑问",
    "不可否认",
    "显而易见",
    "诚然",
    "值得一提的是",
    "众所周知",
    "从某种意义上说",
    "令人惊叹",
    "史无前例",
    "前所未有",
    "完美融合",
    "淋漓尽致",
    "叹为观止",
    "不可或缺",
    "深刻印象",
    "旨在",
    "赋能",
    "织就了一幅",
    "踏上",
    "展露无遗",
    "焕发新机",
    "引领潮流",
    "势必会",
    "进一步来说",
    "需要指出的是",
    "有理由相信",
    "可以预见的是",
    "归根结底",
    "归根到底",
    "在这一过程中",
    "不止如此",
])
TERM_RULE_DEFAULT_REPLACEMENT_TERMS_TEXT = "\n".join([
    "综上所述 => 整体来看",
    "总而言之 => 整体来看",
    "深入探讨 => 拆解",
    "深入分析 => 拆解",
    "旨在 => 核心是",
    "毋庸置疑 => 现实情况是",
    "毫无疑问 => 现实情况是",
    "不可否认 => 客观来说",
    "势必会 => 大概率会",
    "值得一提的是 => 更值得细看的是",
    "值得关注 => 更值得细看的是",
    "诚然 => 确实",
    "从某种意义上说 => 换个更直接的说法",
    "进一步来说 => 往下看",
    "归根结底 => 说到底",
    "归根到底 => 说到底",
    "在这一过程中 => 在这个过程里",
    "史无前例 => 很少见",
    "前所未有 => 和过去不太一样",
    "需要指出的是 => 这里还有个关键点",
])
TERM_RULE_SUGGESTED_REPLACEMENT_TERMS_TEXT = "\n".join([
    "不可或缺 => 少不了的一环",
    "具有重大意义 => 算一个关键变化",
    "面临着挑战 => 难点在于",
    "面临困境 => 遇到了明显瓶颈",
    "创新性地 => 新的地方在于",
    "核心优势在于 => 真正有优势的地方在于",
    "引起了广泛关注 => 引发了不少讨论",
    "优化用户体验 => 减轻负担",
    "显而易见 => 这点其实很清楚",
    "众所周知 => 这在圈内基本是常识",
    "令人惊叹 => 确实很亮眼",
    "令人瞩目 => 表现很突出",
    "完美融合 => 结合得比较顺",
    "打破传统 => 跳出了原来的做法",
    "提供了强有力的支持 => 提供了很强的支撑",
    "深刻印象 => 让人记得住",
    "赋能 => 直接提高了",
    "引领潮流 => 成了风向标之一",
    "有理由相信 => 更稳妥的说法是",
    "可以预见的是 => 后面大概率会出现",
    "不止如此 => 还不止这些",
])
ROLE_AUDIENCE_MAP = {
    "\u53d1\u884c\u4e3b\u7f16": "\u6e38\u620f\u884c\u4e1a\u4ece\u4e1a\u8005",
    "\u7814\u53d1\u4e3b\u7f16": "\u6e38\u620f\u5708\u540c\u884c\u548c\u786c\u6838\u73a9\u5bb6",
    "\u6e38\u620f\u5feb\u8baf\u7f16\u8f91": "\u884c\u4e1a\u4ece\u4e1a\u8005",
    "\u5ba2\u89c2\u8f6c\u5f55\u7f16\u8f91": "\u516c\u4f17\u8bfb\u8005",
    "\u6e38\u620f\u884c\u4e1a\u8bc4\u8bba\u5458": "\u6e38\u620f\u884c\u4e1a\u4ece\u4e1a\u8005",
}
ROLE_JARGON_MAP = {
    "\u53d1\u884c\u4e3b\u7f16": "ROI, LTV, \u4e70\u91cf, \u53d8\u73b0\u6548\u7387",
    "\u7814\u53d1\u4e3b\u7f16": "\u6838\u5fc3\u5faa\u73af, \u7cfb\u7edf\u8bbe\u8ba1, \u6280\u672f\u503a, \u5de5\u4e1a\u5316\u7ba1\u7ebf",
    "\u6e38\u620f\u5feb\u8baf\u7f16\u8f91": "\u7248\u53f7, \u4e0a\u7ebf\u6863\u671f, \u53d1\u884c\u8282\u594f, \u5e02\u573a\u53cd\u9988",
    "\u5ba2\u89c2\u8f6c\u5f55\u7f16\u8f91": "\u7559\u5b58, \u53d8\u73b0, \u672c\u5730\u5316, \u8fd0\u8425\u8282\u594f",
    "\u6e38\u620f\u884c\u4e1a\u8bc4\u8bba\u5458": "ROI, \u6d17\u91cf, \u8dd1\u91cf, \u5546\u4e1a\u5316\u6548\u7387",
}
ROLE_TONE_MAP = {
    "\u53d1\u884c\u4e3b\u7f16": "\u51b7\u9759\u3001\u950b\u5229\u3001\u5224\u65ad\u660e\u786e",
    "\u7814\u53d1\u4e3b\u7f16": "\u4e13\u4e1a\u3001\u76f4\u63a5\u3001\u5e26\u4e00\u70b9\u6bd2\u820c",
    "\u6e38\u620f\u5feb\u8baf\u7f16\u8f91": "\u514b\u5236\u3001\u6e05\u695a\u3001\u4fe1\u606f\u5bc6\u5ea6\u9ad8",
    "\u5ba2\u89c2\u8f6c\u5f55\u7f16\u8f91": "\u5e73\u5b9e\u3001\u5e72\u51c0\u3001\u7ed3\u6784\u6e05\u6670",
    "\u6e38\u620f\u884c\u4e1a\u8bc4\u8bba\u5458": "\u4e00\u9488\u89c1\u8840\u3001\u5f3a\u89c2\u70b9\u4f46\u4e0d\u6d6e\u5938",
}


def infer_role_persona(role_name, editor_prompt):
    prompt_text = (editor_prompt or "").strip()
    first_line = prompt_text.splitlines()[0].strip() if prompt_text else ""
    if first_line.startswith("# Role:"):
        persona = first_line.split(":", 1)[1].strip()
        if persona:
            return persona
    return role_name or "\u8d44\u6df1\u884c\u4e1a\u4f5c\u8005"


def infer_article_topic(source_content):

    clean_text = " ".join((source_content or "").split())
    if not clean_text:
        return "当前游戏行业主题"
    return clean_text[:40] + ("..." if len(clean_text) > 40 else "")


def build_de_ai_prompt_template(role_name, editor_prompt, source_content, variant=DE_AI_VARIANT_DEFAULT, term_rules_instruction=""):
    persona = infer_role_persona(role_name, editor_prompt)
    topic = infer_article_topic(source_content)
    audience = ROLE_AUDIENCE_MAP.get(role_name, "行业读者")
    jargon = ROLE_JARGON_MAP.get(role_name, "ROI, LTV, retention, monetization")
    tone = ROLE_TONE_MAP.get(role_name, "冷静, 尖锐, 有判断")
    target_words = get_target_article_words()
    try:
        target_words = int(target_words)
    except Exception:
        target_words = 1500
    if target_words < 1200:
        heading_instruction = "短稿默认使用 3 个 `##` 二级标题。"
    elif target_words <= 2500:
        heading_instruction = "中等篇幅默认使用 3-5 个 `##` 二级标题，让正文中段有清晰层级。"
    else:
        heading_instruction = "中长稿默认使用 4-5 个 `##` 二级标题，只有某一节明显复杂时才补少量 `###`。"
    structure_instruction = "\n".join([
        "【文章结构协议｜最高优先级】",
        "正文必须使用 Markdown 输出。",
        "开头先用 1-2 段导语交代对象和背景，再进入正文分析。",
        heading_instruction,
        "每个 `##` 小节通常保持 2-4 个自然段，不要把正文写成一整堵长墙。",
        "小节标题必须具体、信息明确，禁止使用“背景介绍”“总结一下”“最后说说”这类空标题。",
        "结构要清楚，但不要把每一节写成机械等长的模板段落。",
    ])
    community_instruction = ""
    chatty_instruction = ""
    humanizer_instruction = ""
    if variant == DE_AI_VARIANT_COMMUNITY:
        community_instruction = """
【社区文章去AI版附加要求】
请在现有去 AI 约束上，再做一层面向玩家社区的轻口语化改写。
- 用更自然的玩家向表达，减少媒体稿腔、说明书腔和行业报告腔。
- 观点要清楚，但句子节奏要更像玩家社区里的高质量长帖。
- 多用短句和自然停顿，让段落推进更顺，不要一段话拧得太满。
- 可以适度加入玩家视角的共鸣感，但不要写成情绪宣泄或粉黑大战。
- 把过硬、过媒体化、过行业黑话化的表达，改成普通玩家也能顺着读下去的话。
- 减少术语堆叠。必要术语可以保留，但要顺手讲成人话。
- 结尾要更明确地落到“这对玩家意味着什么”“接下来社区会怎么讨论”。
- 不要过度互动化：避免频繁反问、硬玩梗、贴吧口癖和夸张情绪词泛滥。
"""
    elif variant == DE_AI_VARIANT_CHAT:
        chatty_instruction = """
【自然唠嗑版附加要求】
请在现有去 AI 约束上，再做一层中强度口语化改写，让文章像一个有经验、有判断的人在认真把事情讲明白。
- 保留分析的锋利度和信息量，但整体表达要像真人当面在讲，不要像在交报告。
- 允许出现“说实话”“其实”“换句话说”“问题就在这儿”这类自然转折，但要放在自然的位置，不要刻意堆叠。
- 让句子长短更有变化。该短的时候就短，不要每一句都写得一样整齐。
- 可以有少量场景感、轻微自我补充，或者一句克制的括号插话，但不要喧宾夺主。
- 把媒体稿腔、报告腔、说明书腔换成更像人话的表达，同时保留对读者的尊重。
- 专业术语在必要时可以保留，但要顺手解释，不要把术语一层层往上堆。
- 结尾要落到“这件事到底意味着什么”，而不是只把分析收住。
- 文章仍然要像高质量长文，不要写成低质水贴、表情包文风、贴吧口癖、密集反问或自我陶醉式絮叨。
- 第一人称存在感可以有，但不能压过主题本身。核心始终是对象、事实和论证。
- 标题结构、开头背景和分析骨架都必须保留。唠嗑感不等于松散，也不等于想到哪写到哪。
"""
    elif variant == DE_AI_VARIANT_HUMANIZER:
        humanizer_instruction = """
【Humanizer-zh版附加要求】
请在现有去 AI 约束上，再做一轮面向「去模板化」的精细重写。
- 删掉 AI 高频填充词、连接词堆叠和听起来很像「标准答案」的过渡句。
- 主动打破「不是……而是……」、「不仅……而且……」、「这不仅仅是……」这类公式化句式。
- 减少三项并列、金句感、空泛升华和过度结论化的写法。
- 去掉宣传腔、报告腔、媒体腔和伪深刻措辞，能用人话说清的地方就不要继续套话化。
- 遇到模糊归因或虚空权威化表达，请改写得更具体、更像真人在讲话。
- 保留事实和结构，但让句子节奏更自然，判断更像有人在思考过后写下来的文字。
- 不要把文章写成故意凹人设的「新文风」，目标是减少机器纹理，不是再造另一种固定腔调。
"""
    return f"""# 角色
你现在的身份是：{persona}

# 背景
我现在有一篇关于【{topic}】的文章草稿。它的信息量和分析骨架是有价值的，但语言质感还偏 AI，不够像真人写的成熟中文长文。
请把它改写成一篇会让目标读者【{audience}】相信是真人作者写出来的自然简体中文文章。

# 任务
保留草稿中的关键事实、数据点、观点、例子和推理链条，只改表达方式、句子节奏、段落推进和语言质感。
如果原稿开头太突兀，请先补 1-2 段导语，让读者一上来就知道这篇文章到底在讲哪款游戏、哪个公司、哪个事件或哪场争议，以及为什么这件事现在值得分析。

# 核心约束
1. 全部输出必须使用简体中文。
2. 不得删除备选标题组。
3. 不得把文章重新抹平成一整堵大长段。
4. 事实顺序和分析方向必须保持稳定，不要擅自改动论证重心。
5. 需要用到术语时可以自然使用，但要符合中文语境：{jargon}。
6. 整体语气优先靠近：{tone}。

{term_rules_instruction}
{community_instruction}
{chatty_instruction}
{humanizer_instruction}
# 结构保留要求
{structure_instruction}
- 如果原稿现有标题层级是合理的，可以润色标题措辞，但不能把层级删掉。
- 最终必须保留备选标题组，并维持 3-5 个备选标题；如果文章方向没有变，优先保留原标题骨架，只优化措辞和传播力。
- 如果开头背景不够，就在正文分析前补出简洁的 1-2 段导语。
- 去 AI 的本质是重写表达，不是推翻原稿重新起草。

# 输出格式要求
严格按以下顺序输出 3 个区块，不要添加任何额外解释：

{PURE_TITLE_MARKER}
1. ...
2. ...
3. ...
这里保留 3-5 个备选标题。

{PURE_BODY_MARKER}
这里只输出纯净定稿正文，不要重复标题组，不要输出 HTML。

{HIGHLIGHT_MARKER}
这里输出同一篇正文的高亮阅读版 HTML。
允许使用的标签只有：<h2>、<h3>、<p>、<strong>、<span class="highlight-positive">、<span class="highlight-risk">。
如果需要分节标题，优先使用 <h2>/<h3>，不要继续输出 Markdown 的 `##` 标题。
不要改动事实，不要凭空增删内容，只高亮真正重要的认知点、判断点或风险提示。

# 原始草稿
[在这里粘贴完整草稿]"""


def parse_de_ai_raw_title_candidates(response_text):
    clean_text = (response_text or "").strip()
    if not clean_text:
        return []
    if PURE_TITLE_MARKER in clean_text and PURE_BODY_MARKER in clean_text:
        title_part = clean_text.split(PURE_TITLE_MARKER, 1)[1].split(PURE_BODY_MARKER, 1)[0]
        return parse_title_candidates_block(title_part)
    if ARTICLE_TITLE_MARKER in clean_text:
        title_part = clean_text.split(ARTICLE_TITLE_MARKER, 1)[1]
        if ARTICLE_BODY_MARKER in title_part:
            title_part = title_part.split(ARTICLE_BODY_MARKER, 1)[0]
        return parse_title_candidates_block(title_part)
    return []


def extract_de_ai_raw_title_candidates(response_text):
    return parse_de_ai_raw_title_candidates(response_text)


def ensure_highlighted_article_context(highlighted_text, title_candidates, article_body):
    clean_highlight = (highlighted_text or "").strip()
    if not clean_highlight:
        return ""

    def normalize_plain_text(text):
        normalized = (text or "").strip()
        if not normalized:
            return ""
        normalized = re.sub(r"<\s*br\s*/?>", "\n", normalized, flags=re.IGNORECASE)
        normalized = re.sub(r"</\s*(p|div|h1|h2|h3|li)\s*>", "\n\n", normalized, flags=re.IGNORECASE)
        normalized = re.sub(r"<\s*li\b[^>]*>", "- ", normalized, flags=re.IGNORECASE)
        normalized = re.sub(r"<[^>]+>", "", normalized)
        normalized = (
            normalized
            .replace("&nbsp;", " ")
            .replace("&amp;", "&")
            .replace("&lt;", "<")
            .replace("&gt;", ">")
            .replace("&quot;", '"')
            .replace("&#39;", "'")
        )
        return re.sub(r"\s+", " ", normalized).strip()

    prefix_parts = []
    normalized_titles = normalize_title_candidates(title_candidates)
    intro_source = re.split(r"^\s*##\s+.+$", article_body or "", maxsplit=1, flags=re.MULTILINE)[0].strip()
    intro_paragraphs = [paragraph.strip() for paragraph in re.split(r"\n\s*\n", intro_source) if paragraph.strip()]

    if normalized_titles and ARTICLE_TITLE_MARKER not in clean_highlight and PURE_TITLE_MARKER not in clean_highlight:
        title_block = build_title_candidates_block(normalized_titles)
        if title_block:
            prefix_parts.append(title_block)

    highlight_starts_with_heading = bool(re.match(r"^(?:<h[1-3][^>]*>|##\s+|###\s+)", clean_highlight))
    missing_intro = ""
    if intro_source and ARTICLE_BODY_MARKER not in clean_highlight and PURE_BODY_MARKER not in clean_highlight:
        highlight_plain = normalize_plain_text(clean_highlight)
        matched_intro = False
        for start_idx in range(len(intro_paragraphs)):
            intro_tail = "\n\n".join(intro_paragraphs[start_idx:])
            intro_tail_plain = normalize_plain_text(intro_tail)
            if intro_tail_plain and highlight_plain.startswith(intro_tail_plain):
                missing_intro = "\n\n".join(intro_paragraphs[:start_idx]).strip()
                matched_intro = True
                break
        if not matched_intro:
            missing_intro = intro_source

    if missing_intro and (highlight_starts_with_heading or prefix_parts):
        prefix_parts.append(f"{ARTICLE_BODY_MARKER}\n{missing_intro}")

    if not prefix_parts:
        return clean_highlight
    return "\n\n".join(prefix_parts + [clean_highlight]).strip()


def parse_de_ai_dual_output(response_text, fallback_titles=None):
    clean_text = (response_text or "").strip()
    if not clean_text:
        return normalize_title_candidates(fallback_titles), "", ""

    resolved_titles = normalize_title_candidates(fallback_titles)
    pure_part = clean_text
    if PURE_TITLE_MARKER in clean_text and PURE_BODY_MARKER in clean_text:
        title_part = clean_text.split(PURE_TITLE_MARKER, 1)[1].split(PURE_BODY_MARKER, 1)[0]
        parsed_titles = parse_title_candidates_block(title_part)
        if parsed_titles:
            resolved_titles = parsed_titles
        pure_part = clean_text.split(PURE_BODY_MARKER, 1)[1]
    elif PURE_BODY_MARKER in clean_text:
        pure_part = clean_text.split(PURE_BODY_MARKER, 1)[1]
    else:
        parsed_titles, pure_body = split_structured_article_sections(clean_text)
        return parsed_titles or resolved_titles, pure_body or clean_text, ""

    highlighted_text = ""
    if HIGHLIGHT_MARKER in pure_part:
        pure_text, highlighted_text = pure_part.split(HIGHLIGHT_MARKER, 1)
    else:
        pure_text = pure_part

    _, pure_body = split_structured_article_sections(pure_text)
    clean_body = (pure_body or pure_text or "").strip()
    clean_highlighted = ensure_highlighted_article_context(highlighted_text, resolved_titles, clean_body)
    return resolved_titles, clean_body, clean_highlighted


def sanitize_highlighted_article(html_text):
    clean_text = (html_text or "").strip()
    if not clean_text:
        return ""

    clean_text = clean_text.replace("<script", "&lt;script")
    clean_text = clean_text.replace("</script>", "&lt;/script&gt;")
    clean_text = clean_text.replace("<style", "&lt;style")
    clean_text = clean_text.replace("</style>", "&lt;/style&gt;")
    clean_text = clean_text.replace("onerror=", "data-onerror=")
    clean_text = clean_text.replace("onclick=", "data-onclick=")
    clean_text = clean_text.replace("onload=", "data-onload=")
    clean_text = clean_text.replace("&#x1F517;", "")
    clean_text = clean_text.replace("\U0001F517", "")

    normalized_lines = []
    for raw_line in clean_text.splitlines():
        stripped = raw_line.strip()
        heading_match = re.match(r"^(#{2,3})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).replace("\U0001F517", "")
            heading_text = re.sub(r"\s*#+\s*$", "", heading_text).strip()
            normalized_lines.append(f"<h{level}>{heading_text}</h{level}>")
        else:
            normalized_lines.append(raw_line)
    clean_text = "\n".join(normalized_lines)

    clean_text = re.sub(
        r'<a[^>]*?(?:headerlink|anchor|fragment-link)[^>]*>.*?</a>',
        "",
        clean_text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    clean_text = re.sub(
        r'<a[^>]*?href="#.*?"[^>]*>.*?</a>',
        "",
        clean_text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return clean_text


def render_highlighted_article_panel(html_text):


    if not (html_text or "").strip():
        st.info("当前版本未生成高亮阅读视图。")
        return

    st.markdown(
        """
        <style>
        .highlight-article {
            padding: 1.15rem 1.2rem;
            border: 1px solid rgba(22, 33, 28, 0.08);
            border-radius: 18px;
            background: rgba(255, 255, 255, 0.74);
            box-shadow: 0 18px 40px rgba(22, 33, 28, 0.06);
            line-height: 1.9;
            color: #16211c;
        }
        .highlight-article h2,
        .highlight-article h3 {
            margin: 0 0 1rem;
            color: #16211c;
            font-weight: 700;
            line-height: 1.3;
        }
        .highlight-article h2 {
            font-size: 2.1rem;
        }
        .highlight-article h3 {
            font-size: 1.45rem;
        }
        .highlight-article p {
            margin: 0 0 1rem;
        }
        .highlight-article h1 a,
        .highlight-article h2 a,
        .highlight-article h3 a,
        .highlight-article a[href^="#"] {
            display: none !important;
        }
        .highlight-article .highlight-positive {
            display: inline;
            padding: 0.05rem 0.32rem;
            border-radius: 0.4rem;
            background: rgba(54, 111, 214, 0.15);
            color: #1f57b8;
            font-weight: 700;
        }
        .highlight-article .highlight-risk {
            display: inline;
            padding: 0.05rem 0.32rem;
            border-radius: 0.4rem;
            background: rgba(214, 76, 76, 0.14);
            color: #b3261e;
            font-weight: 700;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    sanitized_html = sanitize_highlighted_article(html_text)
    render_rich_html_copy_button(sanitized_html, "highlighted_article_step6")
    st.markdown(f'<div class="highlight-article">{sanitized_html}</div>', unsafe_allow_html=True)

def generate_script_for_current_article(api_key, base_url, model_name, script_duration):
    final_article_body = get_article_body_text(st.session_state.get("final_article", ""))
    if not final_article_body:
        st.session_state.spoken_script = ""
        return
    script_sys_prompt = get_script_sys_prompt(script_duration)
    st.session_state.spoken_script = call_llm(
        api_key=api_key,
        base_url=base_url,
        model_name=model_name,
        system_prompt=script_sys_prompt,
        user_content=f"\u3010\u8bf7\u5c06\u4ee5\u4e0b\u6df1\u5ea6\u6587\u7ae0\u8f6c\u5316\u4e3a\u4f9b\u526a\u6620AI\u89e3\u6790\u7684{script_duration}\u53e3\u64ad\u4e0e\u5206\u955c\u811a\u672c\uff1a\u3011\n\n{final_article_body}"
    )

# ==========================================


# 1. API 与外部推送函数
# ==========================================
def call_llm(api_key, base_url, model_name, system_prompt, user_content, image_urls=None, history=None, temperature=None):
    if not api_key:
        st.error("Please enter an API key in the sidebar first.")
        st.stop()

    if image_urls is None:
        image_urls = []
    if history is None:
        history = []

    current_stage = st.session_state.get("pending_ai_stage", "")
    retry_delays = (0.8, 1.6)
    attempt_count = 0

    try:
        timeout_config = build_openai_timeout(base_url, model_name)
        use_stream = should_stream_llm_request(base_url, model_name)

        messages = [{"role": "system", "content": system_prompt}]

        if history:
            messages.extend(history)

        if image_urls:
            message_content = [{"type": "text", "text": user_content}]
            for img_url in image_urls:
                message_content.append({
                    "type": "image_url",
                    "image_url": {"url": img_url}
                })
        else:
            message_content = user_content

        messages.append({"role": "user", "content": message_content})

        request_kwargs = {
            "model": model_name,
            "messages": messages,
            "temperature": 0.3 if temperature is None else temperature,
        }

        last_error = None
        for attempt_index in range(len(retry_delays) + 1):
            attempt_count = attempt_index + 1
            try:
                content = execute_llm_request_once(
                    api_key=api_key,
                    base_url=base_url,
                    model_name=model_name,
                    request_kwargs=request_kwargs,
                    timeout_config=timeout_config,
                    use_stream=use_stream,
                )
                st.session_state.last_ai_error = ""
                return content
            except Exception as attempt_error:
                last_error = attempt_error
                if attempt_index >= len(retry_delays) or not should_retry_with_requests_fallback(attempt_error):
                    raise
                time.sleep(retry_delays[attempt_index])

        raise last_error if last_error else RuntimeError("Unknown LLM request failure.")
    except Exception as e:
        formatted_error = format_llm_exception(e)
        st.session_state.last_ai_error = formatted_error
        st.session_state.pending_ai_stage = ""
        write_ai_diagnostic_log(
            current_stage,
            base_url,
            model_name,
            formatted_error,
            user_content,
            image_count=len(image_urls),
            history_count=len(history),
            extra_details={
                **build_proxy_diagnostic_snapshot(),
                "attempt_count": attempt_count,
            },
        )
        save_draft()
        play_ai_error_sound()
        st.error(f"API call failed: {formatted_error}")
        st.stop()


def push_to_feishu(article_text, script_text=None):
    webhook_url = "https://open.feishu.cn/open-apis/bot/v2/hook/a0f50778-0dd2-4963-a0b2-0c7b68e113d8"
    headers = {"Content-Type": "application/json"}
    
    if script_text:
        text_content = f"📣 【公众号文章定稿通知】\n\n{article_text}\n\n================\n\n🎬【短视频 AI 分镜脚本】\n\n{script_text}"
    else:
        text_content = f"📣 【公众号文章定稿通知】\n\n{article_text}"
        
    payload = {
        "msg_type": "text",
        "content": {
            "text": text_content
        }
    }
    
    try:
        response = requests.post(webhook_url, headers=headers, data=json.dumps(payload))
        if response.status_code == 200:
            resp_json = response.json()
            if resp_json.get("code") == 0:
                return True, "推送成功"
            else:
                return False, f"飞书返回错误: {resp_json.get('msg')}"
        else:
            return False, f"HTTP 请求失败，状态码: {response.status_code}"
    except Exception as e:
        return False, f"请求发生异常: {str(e)}"

# ==========================================
# 2. 核心抓取函数
# ==========================================
# (抓取函数保持不变)
def extract_youtube_transcript(url):
    try:
        parsed_url = urlparse(url)
        if 'youtube.com' in parsed_url.netloc:
            video_id = parse_qs(parsed_url.query).get('v', [None])[0]
        elif 'youtu.be' in parsed_url.netloc:
            video_id = parsed_url.path.lstrip('/')
        else:
            return None, [], "未找到有效的 YouTube Video ID"

        if not video_id:
            return None, [], "无法解析 YouTube 链接"

        target_langs = ['zh-Hans', 'zh-Hant', 'zh-CN', 'zh-TW', 'zh', 'en']
        try:
            if hasattr(YouTubeTranscriptApi, 'get_transcript'):
                transcript_fetched = YouTubeTranscriptApi.get_transcript(video_id, languages=target_langs)
            else:
                ytt_api = YouTubeTranscriptApi()
                if hasattr(ytt_api, 'list'):
                    transcript_list = ytt_api.list(video_id)
                elif hasattr(ytt_api, 'list_transcripts'):
                    transcript_list = ytt_api.list_transcripts(video_id)
                else:
                    transcript_fetched = ytt_api.fetch(video_id)
                    transcript_list = None
                    
                if transcript_list is not None:
                    try:
                        transcript = transcript_list.find_transcript(target_langs)
                    except Exception:
                        transcript = list(transcript_list)[0]
                    transcript_fetched = transcript.fetch()

            texts = []
            for item in transcript_fetched:
                if isinstance(item, dict) and 'text' in item:
                    texts.append(item['text'])
                elif hasattr(item, 'text'):
                    texts.append(item.text)
                else:
                    texts.append(str(item))
                    
            text = "\n".join(texts)
            return text, [], None
            
        except Exception as inner_e:
            error_str = str(inner_e).lower()
            if "block" in error_str or "proxy" in error_str or "could not retrieve a transcript" in error_str:
                try:
                    jina_url = f"https://r.jina.ai/{url}"
                    headers = {"Accept": "text/plain"}
                    response = requests.get(jina_url, headers=headers, timeout=20)
                    response.raise_for_status()
                    
                    text = response.text
                    if text and len(text) > 50:
                        return text, [], None
                    else:
                        return None, [], "原生API被封锁，且备用穿透解析器未返回有效文本。"
                except Exception as jina_e:
                    return None, [], f"YouTube 提取失败: 原生报错({str(inner_e)[:30]}...) | 备用报错({str(jina_e)[:30]}...)"
            else:
                raise inner_e

    except Exception as e:
        return None, [], f"YouTube 字幕抓取失败: {str(e)}"

def extract_article_content(url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        text = trafilatura.extract(response.text)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        noise_tags = ['aside', 'nav', 'footer', 'header']
        for noise in soup.find_all(noise_tags):
            noise.decompose()
            
        noise_keywords = ['author', 'related', 'comment', 'share', 'widget', 'sidebar', 'ad-container', 'popup', 'newsletter']
        for noise in soup.find_all(attrs={'class': lambda c: c and any(k in str(c).lower() for k in noise_keywords)}):
            noise.decompose()
        for noise in soup.find_all(attrs={'id': lambda i: i and any(k in str(i).lower() for k in noise_keywords)}):
            noise.decompose()
            
        main_content = soup.find('article') or soup.find('main') or soup.find(class_=lambda c: c and 'content' in str(c).lower()) or soup

        images = []
        for img in main_content.find_all('img'):
            try:
                html_w = int(img.get('width', 0))
                html_h = int(img.get('height', 0))
            except:
                html_w, html_h = 0, 0
                
            src = None
            srcset = img.get('data-srcset') or img.get('srcset')
            if srcset:
                sources = []
                for s in srcset.split(','):
                    parts = s.strip().split()
                    if len(parts) == 2 and parts[1].endswith('w') and parts[1][:-1].isdigit():
                        sources.append((parts[0], int(parts[1][:-1])))
                if sources:
                    sources.sort(key=lambda x: x[1], reverse=True)
                    src = sources[0][0]
                    
            if not src:
                for attr in ['data-original', 'data-lazy-src', 'data-src', 'src']:
                    val = img.get(attr)
                    if val:
                        if isinstance(val, list): val = val[0]
                        val = str(val).strip()
                        if not val.startswith('data:image'):
                            src = val
                            break
                            
            if not src: continue
                
            if src.startswith('//'): src = 'https:' + src
            elif src.startswith('/'):
                parsed_url = urlparse(url)
                src = f"{parsed_url.scheme}://{parsed_url.netloc}{src}"
            elif not src.startswith('http'): continue
                
            src_lower = src.lower()
            junk_keywords = ['icon', 'spinner', 'svg', 'gif', 'button', 'tracker', 'avatar']
            if any(junk in src_lower for junk in junk_keywords):
                continue
                
            if html_w >= 300 or html_h >= 300:
                pass 
            else:
                match = re.search(r'-(\d{2,3})x(\d{2,3})\.(jpg|jpeg|png|webp)', src_lower)
                if match:
                    mw, mh = int(match.group(1)), int(match.group(2))
                    if mw <= 300 or mh <= 300:
                        continue
                        
            if src not in images:
                images.append(src)
                if len(images) >= 8:
                    break
        
        if not text:
            paragraphs = soup.find_all('p')
            text = '\n'.join([p.get_text() for p in paragraphs])
            
        if text and text.strip():
            return text, images, None
        else:
            return None, [], "未能提取到有效纯文本。"
    except Exception as e:
        return None, [], f"文章抓取失败: {str(e)}"

def extract_article_content_with_fallback(url):
    def build_article_request_headers():
        return {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/135.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            "Cache-Control": "no-cache",
            "Pragma": "no-cache",
            "Upgrade-Insecure-Requests": "1",
        }

    def extract_article_text_and_images_from_html(html_text, page_url):
        text = trafilatura.extract(html_text)
        soup = BeautifulSoup(html_text, 'html.parser')

        noise_tags = ['aside', 'nav', 'footer', 'header']
        for noise in soup.find_all(noise_tags):
            noise.decompose()

        noise_keywords = ['author', 'related', 'comment', 'share', 'widget', 'sidebar', 'ad-container', 'popup', 'newsletter']
        for noise in soup.find_all(attrs={'class': lambda c: c and any(k in str(c).lower() for k in noise_keywords)}):
            noise.decompose()
        for noise in soup.find_all(attrs={'id': lambda i: i and any(k in str(i).lower() for k in noise_keywords)}):
            noise.decompose()

        main_content = soup.find('article') or soup.find('main') or soup.find(class_=lambda c: c and 'content' in str(c).lower()) or soup

        images = []
        for img in main_content.find_all('img'):
            try:
                html_w = int(img.get('width', 0))
                html_h = int(img.get('height', 0))
            except Exception:
                html_w, html_h = 0, 0

            src = None
            srcset = img.get('data-srcset') or img.get('srcset')
            if srcset:
                sources = []
                for s in srcset.split(','):
                    parts = s.strip().split()
                    if len(parts) == 2 and parts[1].endswith('w') and parts[1][:-1].isdigit():
                        sources.append((parts[0], int(parts[1][:-1])))
                if sources:
                    sources.sort(key=lambda x: x[1], reverse=True)
                    src = sources[0][0]

            if not src:
                for attr in ['data-original', 'data-lazy-src', 'data-src', 'src']:
                    val = img.get(attr)
                    if val:
                        if isinstance(val, list):
                            val = val[0]
                        val = str(val).strip()
                        if not val.startswith('data:image'):
                            src = val
                            break

            if not src:
                continue

            if src.startswith('//'):
                src = 'https:' + src
            elif src.startswith('/'):
                parsed_url = urlparse(page_url)
                src = f"{parsed_url.scheme}://{parsed_url.netloc}{src}"
            elif not src.startswith('http'):
                continue

            src_lower = src.lower()
            junk_keywords = ['icon', 'spinner', 'svg', 'gif', 'button', 'tracker', 'avatar']
            if any(junk in src_lower for junk in junk_keywords):
                continue

            if html_w < 300 and html_h < 300:
                match = re.search(r'-(\d{2,3})x(\d{2,3})\.(jpg|jpeg|png|webp)', src_lower)
                if match:
                    mw, mh = int(match.group(1)), int(match.group(2))
                    if mw <= 300 or mh <= 300:
                        continue

            if src not in images:
                images.append(src)
                if len(images) >= 8:
                    break

        if not text:
            paragraphs = soup.find_all('p')
            text = '\n'.join([p.get_text() for p in paragraphs])

        return (text or '').strip(), images

    def fetch_article_via_jina_reader(page_url):
        jina_url = f"https://r.jina.ai/{page_url}"
        headers = {
            'Accept': 'text/plain',
            'User-Agent': build_article_request_headers()['User-Agent'],
        }
        response = requests.get(jina_url, headers=headers, timeout=20)
        response.raise_for_status()
        text = (response.text or '').strip()
        if len(text) < 50:
            raise ValueError('Fallback reader returned insufficient text.')
        return text

    direct_images = []
    direct_error = None
    try:
        response = requests.get(url, headers=build_article_request_headers(), timeout=15)
        response.raise_for_status()

        content_type = str(response.headers.get('Content-Type', '') or '').lower()
        if content_type and 'text/html' not in content_type and 'application/xhtml+xml' not in content_type:
            raise ValueError(f'Unsupported content type from origin: {content_type}')

        text, direct_images = extract_article_text_and_images_from_html(response.text, url)
        if text:
            return text, direct_images, None
        direct_error = ValueError('Page body was empty after extraction.')
    except Exception as exc:
        direct_error = exc

    try:
        fallback_text = fetch_article_via_jina_reader(url)
        return fallback_text, direct_images, None
    except Exception as fallback_exc:
        return None, [], f'Article fetch failed: origin({str(direct_error)}) | fallback({str(fallback_exc)})'


def get_content_from_url(url):
    if "youtube.com" in url or "youtu.be" in url:
        return extract_youtube_transcript(url)
    else:
        return extract_article_content_with_fallback(url)


MAX_UPLOAD_TEXT_CHARS = 18000
MAX_UPLOAD_IMAGE_BYTES = 6 * 1024 * 1024
SUPPORTED_UPLOAD_EXTENSIONS = [
    "png", "jpg", "jpeg", "webp", "bmp", "gif",
    "doc", "docx", "xlsx", "xls", "txt", "md", "csv", "log", "json"
]


def trim_uploaded_text(text, max_chars=MAX_UPLOAD_TEXT_CHARS):
    if not text:
        return ""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + f"\n\n[内容过长，已截断。仅保留前 {max_chars} 个字符用于分析。]"


def decode_text_file_bytes(file_bytes):
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return file_bytes.decode(encoding), None
        except Exception:
            continue
    return None, "文本文件编码无法识别，请尝试转为 UTF-8 后重传。"


def extract_text_from_docx_bytes(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        blocks = []
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                blocks.append(paragraph_text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))

        if not blocks:
            return None, "Word 文件没有可提取文本。"
        return "\n".join(blocks), None
    except Exception as exc:
        return None, f"Word 解析失败: {str(exc)}"


def extract_text_from_excel_bytes(file_bytes):
    if pd is None:
        return None, "当前环境未安装 pandas/openpyxl，暂时无法读取 Excel 文件。"

    try:
        workbook = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
    except Exception as exc:
        return None, f"Excel 解析失败: {str(exc)}"

    sheet_blocks = []
    for sheet_name, df in workbook.items():
        try:
            cleaned = df.dropna(how="all").fillna("")
        except Exception:
            cleaned = df

        if cleaned is None or getattr(cleaned, "empty", False):
            continue

        try:
            sheet_text = cleaned.astype(str).to_csv(index=False)
        except Exception:
            sheet_text = str(cleaned)

        if sheet_text.strip():
            sheet_blocks.append(f"[工作表: {sheet_name}]\n{sheet_text}")

    if not sheet_blocks:
        return None, "Excel 文件没有可提取内容。"
    return "\n\n".join(sheet_blocks), None


def image_bytes_to_data_url(file_bytes, mime_type):
    safe_mime = mime_type if mime_type and mime_type.startswith("image/") else "image/png"
    encoded = base64.b64encode(file_bytes).decode("utf-8")
    return f"data:{safe_mime};base64,{encoded}"


def get_image_preview_payload(image_ref):
    if isinstance(image_ref, str) and image_ref.startswith("data:image"):
        try:
            _, encoded = image_ref.split(",", 1)
            return base64.b64decode(encoded)
        except Exception:
            return None
    return image_ref


def extract_content_from_uploaded_files(uploaded_files):
    if not uploaded_files:
        return "", [], [], 0

    text_blocks = []
    image_refs = []
    errors = []
    success_count = 0

    image_suffixes = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}
    text_suffixes = {".txt", ".md", ".csv", ".log", ".json"}

    for idx, uploaded_file in enumerate(uploaded_files, start=1):
        file_name = uploaded_file.name or f"upload_{idx}"
        suffix = Path(file_name).suffix.lower()

        try:
            file_bytes = uploaded_file.getvalue()
        except Exception as exc:
            errors.append(f"文件 {file_name} 读取失败: {str(exc)}")
            continue

        if not file_bytes:
            errors.append(f"文件 {file_name} 为空，已跳过。")
            continue

        mime_type = getattr(uploaded_file, "type", "") or ""

        if suffix in image_suffixes or mime_type.startswith("image/"):
            if len(file_bytes) > MAX_UPLOAD_IMAGE_BYTES:
                errors.append(f"图片 {file_name} 体积过大（>{MAX_UPLOAD_IMAGE_BYTES // (1024 * 1024)}MB），已跳过。")
                continue

            image_refs.append(image_bytes_to_data_url(file_bytes, mime_type))
            text_blocks.append(f"【上传图片素材 {idx}】文件名: {file_name}\n该文件为图片素材，已作为视觉依据加入分析。")
            success_count += 1
            continue

        extracted_text = None
        extract_error = None

        if suffix in text_suffixes:
            extracted_text, extract_error = decode_text_file_bytes(file_bytes)
        elif suffix == ".docx":
            extracted_text, extract_error = extract_text_from_docx_bytes(file_bytes)
        elif suffix == ".doc":
            extract_error = "暂不支持 .doc 老格式，请另存为 .docx 后上传。"
        elif suffix in {".xlsx", ".xls"}:
            extracted_text, extract_error = extract_text_from_excel_bytes(file_bytes)
        else:
            extract_error = f"文件类型不支持: {file_name}"

        if extract_error:
            errors.append(f"文件 {file_name} 处理失败: {extract_error}")
            continue

        cleaned_text = trim_uploaded_text((extracted_text or "").strip())
        if not cleaned_text:
            errors.append(f"文件 {file_name} 未提取到有效内容。")
            continue

        text_blocks.append(f"【上传文件素材 {idx}】文件名: {file_name}\n{cleaned_text}")
        success_count += 1

    deduped_images = []
    seen_images = set()
    for item in image_refs:
        if item in seen_images:
            continue
        seen_images.add(item)
        deduped_images.append(item)

    combined_text = "\n\n================\n\n".join(text_blocks)
    return combined_text, deduped_images, errors, success_count


# ==========================================
# 3. 状态管理初始化
# ==========================================
def init_state():
    if 'current_step' not in st.session_state:
        st.session_state.current_step = 1
    if 'article_url' not in st.session_state:
        st.session_state.article_url = ""
    if 'video_url' not in st.session_state:
        st.session_state.video_url = ""
    if 'source_content' not in st.session_state:
        st.session_state.source_content = ""
    if 'source_images' not in st.session_state:
        st.session_state.source_images = []
    if 'source_images_all' not in st.session_state:
        st.session_state.source_images_all = list(st.session_state.source_images)
    if 'selected_source_image_ids' not in st.session_state:
        st.session_state.selected_source_image_ids = []
    if 'extraction_success' not in st.session_state:
        st.session_state.extraction_success = False
    if 'draft_article' not in st.session_state:
        st.session_state.draft_article = ""
    if 'review_feedback' not in st.session_state:
        st.session_state.review_feedback = ""
    if 'review_actions' not in st.session_state:
        st.session_state.review_actions = []
    if 'accepted_review_items' not in st.session_state:
        st.session_state.accepted_review_items = []
    if 'modified_article' not in st.session_state:
        st.session_state.modified_article = ""
    if 'final_article' not in st.session_state:
        st.session_state.final_article = ""
    if 'title_candidates' not in st.session_state:
        st.session_state.title_candidates = []
    if 'highlighted_article' not in st.session_state:
        st.session_state.highlighted_article = ""
    if 'spoken_script' not in st.session_state:
        st.session_state.spoken_script = ""
    if 'de_ai_model' not in st.session_state:
        st.session_state.de_ai_model = DE_AI_MODELS[0]
    else:
        st.session_state.de_ai_model = DE_AI_MODEL_MIGRATION.get(
            st.session_state.de_ai_model,
            st.session_state.de_ai_model,
        )
        if st.session_state.de_ai_model not in DE_AI_MODELS:
            st.session_state.de_ai_model = DE_AI_MODELS[0]
    if 'de_ai_variant' not in st.session_state:
        st.session_state.de_ai_variant = DE_AI_VARIANT_DEFAULT
    elif st.session_state.de_ai_variant not in DE_AI_VARIANTS:
        st.session_state.de_ai_variant = DE_AI_VARIANT_DEFAULT
    if 'de_ai_temperature' not in st.session_state:
        st.session_state.de_ai_temperature = 0.75
    if 'de_ai_prompt_template' not in st.session_state:
        st.session_state.de_ai_prompt_template = ""
    if 'term_rules_enabled' not in st.session_state:
        st.session_state.term_rules_enabled = TERM_RULE_DEFAULT_ENABLED
    if 'term_rules_scope' not in st.session_state or not isinstance(st.session_state.term_rules_scope, list):
        st.session_state.term_rules_scope = list(TERM_RULE_SCOPE_DEFAULT)
    else:
        normalized_term_scope = [scope for scope in st.session_state.term_rules_scope if scope in TERM_RULE_SCOPE_LABELS]
        st.session_state.term_rules_scope = normalized_term_scope or list(TERM_RULE_SCOPE_DEFAULT)
    if 'banned_terms_text' not in st.session_state or st.session_state.banned_terms_text is None:
        st.session_state.banned_terms_text = TERM_RULE_DEFAULT_BANNED_TERMS_TEXT
    legacy_global_replacement_text = st.session_state.get('replacement_terms_text', '')
    if 'default_replacement_terms_text' not in st.session_state or st.session_state.default_replacement_terms_text is None:
        st.session_state.default_replacement_terms_text = legacy_global_replacement_text or TERM_RULE_DEFAULT_REPLACEMENT_TERMS_TEXT
    if 'suggested_replacement_terms_text' not in st.session_state or st.session_state.suggested_replacement_terms_text is None:
        st.session_state.suggested_replacement_terms_text = TERM_RULE_SUGGESTED_REPLACEMENT_TERMS_TEXT
    st.session_state.replacement_terms_text = st.session_state.get('default_replacement_terms_text', '') or TERM_RULE_DEFAULT_REPLACEMENT_TERMS_TEXT
    if 'article_banned_terms_text' not in st.session_state or st.session_state.article_banned_terms_text is None:
        st.session_state.article_banned_terms_text = ""
    legacy_article_replacement_text = st.session_state.get('article_replacement_terms_text', '')
    if 'article_default_replacement_terms_text' not in st.session_state or st.session_state.article_default_replacement_terms_text is None:
        st.session_state.article_default_replacement_terms_text = legacy_article_replacement_text or ""
    if 'article_suggested_replacement_terms_text' not in st.session_state or st.session_state.article_suggested_replacement_terms_text is None:
        st.session_state.article_suggested_replacement_terms_text = ""
    st.session_state.article_replacement_terms_text = st.session_state.get('article_default_replacement_terms_text', '') or ""
    if 'term_scan_result' not in st.session_state or not isinstance(st.session_state.term_scan_result, list):
        st.session_state.term_scan_result = []
    if 'term_scan_summary' not in st.session_state or not isinstance(st.session_state.term_scan_summary, dict):
        st.session_state.term_scan_summary = {}
    st.session_state.title_candidates = normalize_title_candidates(st.session_state.get("title_candidates", []))
    if not st.session_state.title_candidates:
        for article_key in ("final_article", "modified_article", "draft_article"):
            recovered_titles, _ = split_structured_article_sections(st.session_state.get(article_key, ""))
            if recovered_titles:
                st.session_state.title_candidates = recovered_titles
                break
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'image_keywords' not in st.session_state:
        st.session_state.image_keywords = ""
    if 'podcast_enabled' not in st.session_state:
        st.session_state.podcast_enabled = False
    if 'podcast_duration' not in st.session_state:
        st.session_state.podcast_duration = "5分钟"
    if 'podcast_script_raw' not in st.session_state:
        st.session_state.podcast_script_raw = ""
    if 'podcast_script_segments' not in st.session_state:
        st.session_state.podcast_script_segments = []
    if 'podcast_audio_path' not in st.session_state:
        st.session_state.podcast_audio_path = ""
    if 'podcast_audio_manifest' not in st.session_state:
        st.session_state.podcast_audio_manifest = {}
    if 'podcast_last_error' not in st.session_state:
        st.session_state.podcast_last_error = ""
    if 'podcast_voice' not in st.session_state:
        st.session_state.podcast_voice = (
            st.session_state.get("podcast_voice_a")
            or st.session_state.get("podcast_voice_b")
            or DEFAULT_TTS_VOICE
        )
    legacy_podcast_segments = st.session_state.get("podcast_script_segments", [])
    if legacy_podcast_segments:
        try:
            st.session_state.podcast_script_segments = normalize_podcast_segments(legacy_podcast_segments)
        except ValueError:
            st.session_state.podcast_script_segments = []
    legacy_podcast_manifest = st.session_state.get("podcast_audio_manifest", {}) or {}
    if legacy_podcast_manifest.get("voice_map") and not legacy_podcast_manifest.get("voice"):
        st.session_state.podcast_audio_path = ""
        st.session_state.podcast_audio_manifest = {}
    if 'podcast_tts_provider' not in st.session_state:
        st.session_state.podcast_tts_provider = "DashScope Qwen-TTS"
    if 'podcast_tts_api_key_present' not in st.session_state:
        st.session_state.podcast_tts_api_key_present = False
    if 'pending_completion_sound' not in st.session_state:
        st.session_state.pending_completion_sound = False
    if 'target_article_words' not in st.session_state:
        st.session_state.target_article_words = 1500
    if 'selected_role_widget' not in st.session_state:
        st.session_state.selected_role_widget = st.session_state.get('selected_role', '')
    if 'target_article_words_slider' not in st.session_state:
        st.session_state.target_article_words_slider = get_target_article_words()
    if 'article_versions' not in st.session_state or not isinstance(st.session_state.article_versions, list):
        st.session_state.article_versions = []
    if 'active_article_version_id' not in st.session_state:
        st.session_state.active_article_version_id = None
    if 'pending_ai_stage' not in st.session_state:
        st.session_state.pending_ai_stage = ""
    if 'last_completed_ai_stage' not in st.session_state:
        st.session_state.last_completed_ai_stage = ""
    if 'last_completed_ai_target_step' not in st.session_state:
        st.session_state.last_completed_ai_target_step = 0
    if 'last_ai_error' not in st.session_state:
        st.session_state.last_ai_error = ""
    if 'recovered_ai_notice' not in st.session_state:
        st.session_state.recovered_ai_notice = ""
    if 'obsidian_enabled' not in st.session_state:
        st.session_state.obsidian_enabled = False
    if 'obsidian_vault_path' not in st.session_state:
        st.session_state.obsidian_vault_path = "E:\\Obsidian\\originvault\\1\u53f7\u4ed3\u5e93\\LLM Wiki\\wiki"
    if st.session_state.get("obsidian_vault_path") == r"E:\Obsidian\originvault\1???\LLM Wiki\wiki":
        st.session_state.obsidian_vault_path = "E:\\Obsidian\\originvault\\1\u53f7\u4ed3\u5e93\\LLM Wiki\\wiki"
    if 'obsidian_max_hits' not in st.session_state:
        st.session_state.obsidian_max_hits = 6
    if 'obsidian_show_hits' not in st.session_state:
        st.session_state.obsidian_show_hits = True
    if 'obsidian_hits' not in st.session_state:
        st.session_state.obsidian_hits = []
    if 'obsidian_research_brief' not in st.session_state:
        st.session_state.obsidian_research_brief = ""
    if 'obsidian_retrieval_error' not in st.session_state:
        st.session_state.obsidian_retrieval_error = ""
    if 'obsidian_query_terms' not in st.session_state:
        st.session_state.obsidian_query_terms = []
    if 'obsidian_wiki_root' not in st.session_state:
        st.session_state.obsidian_wiki_root = ""
    if 'obsidian_last_indexed_at' not in st.session_state:
        st.session_state.obsidian_last_indexed_at = ""
    if 'obsidian_retrieval_signature' not in st.session_state:
        st.session_state.obsidian_retrieval_signature = ""
    if 'obsidian_influence_map' not in st.session_state:
        st.session_state.obsidian_influence_map = []
    if 'obsidian_influence_summary' not in st.session_state:
        st.session_state.obsidian_influence_summary = ""
    if 'obsidian_influence_signature' not in st.session_state:
        st.session_state.obsidian_influence_signature = ""
    if 'evidence_map' not in st.session_state:
        st.session_state.evidence_map = []
    if 'evidence_summary' not in st.session_state:
        st.session_state.evidence_summary = ""
    if 'evidence_signature' not in st.session_state:
        st.session_state.evidence_signature = ""
    if 'task_queue' not in st.session_state or not isinstance(st.session_state.task_queue, list):
        st.session_state.task_queue = []
    if 'task_templates' not in st.session_state or not isinstance(st.session_state.task_templates, list):
        st.session_state.task_templates = []
    if 'active_task_id' not in st.session_state:
        st.session_state.active_task_id = ""
    if '_task_queue_loaded' not in st.session_state:
        st.session_state._task_queue_loaded = False
    if 'task_filter_status' not in st.session_state:
        st.session_state.task_filter_status = "??"
    if 'task_template_apply_targets' not in st.session_state:
        st.session_state.task_template_apply_targets = []
    st.session_state.target_article_words = get_target_article_words()
    st.session_state.target_article_words_slider = get_target_article_words()
    current_role = st.session_state.get('selected_role', '')
    if isinstance(current_role, str) and current_role.strip():
        st.session_state.selected_role_widget = current_role


apply_pending_draft_restore()
init_state()


def mark_ai_stage_started(stage_name):
    st.session_state.pending_ai_stage = stage_name
    st.session_state.last_ai_error = ""
    save_draft()


def checkpoint_ai_stage(stage_name, target_step=None):
    st.session_state.pending_ai_stage = ""
    st.session_state.last_completed_ai_stage = stage_name
    st.session_state.last_completed_ai_target_step = target_step or 0
    st.session_state.last_ai_error = ""
    st.session_state.recovered_ai_notice = ""
    save_draft()


def clear_ai_stage_checkpoint():
    st.session_state.pending_ai_stage = ""
    st.session_state.last_completed_ai_stage = ""
    st.session_state.last_completed_ai_target_step = 0
    save_draft()


def recover_ai_progress_if_needed():
    target_step = st.session_state.get("last_completed_ai_target_step", 0) or 0
    current_step = st.session_state.get("current_step", 1)
    last_stage = st.session_state.get("last_completed_ai_stage", "")
    if target_step and current_step < target_step:
        st.session_state.current_step = target_step
        st.session_state.recovered_ai_notice = f"AI 已完成「{format_ai_stage_name(last_stage)}」，已自动恢复到第 {target_step} 步。"
        clear_ai_stage_checkpoint()
        st.rerun()
    if target_step and current_step >= target_step:
        clear_ai_stage_checkpoint()


def normalize_llm_response_content(content):
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, str):
                parts.append(item)
            elif isinstance(item, dict):
                text_value = item.get("text")
                if isinstance(text_value, str):
                    parts.append(text_value)
        return "\n".join(part for part in parts if part).strip()
    return "" if content is None else str(content)


def format_ai_stage_name(stage_name):
    return {
        "draft_generation": "写出稿",
        "review_generation": "严格审稿",
        "modification_generation": "修改稿件",
        "de_ai_generation": "去 AI 味重写",
    }.get(stage_name, stage_name or "AI 任务")


def render_ai_progress_banner():
    draft_restore_notice = st.session_state.pop(DRAFT_RESTORE_NOTICE_KEY, "")
    if draft_restore_notice:
        st.success(draft_restore_notice)

    recovered_notice = st.session_state.get("recovered_ai_notice", "")
    if recovered_notice:
        st.success(recovered_notice)

    pending_stage = st.session_state.get("pending_ai_stage", "")
    if pending_stage:
        st.info(f"AI 正在执行「{format_ai_stage_name(pending_stage)}」。如果页面短暂重载，系统会自动尝试续上流程。")

    last_error = (st.session_state.get("last_ai_error", "") or "").strip()
    if last_error:
        st.warning(f"上一轮 AI 调用未完整收尾：{last_error}")


def is_task_action_blocked(pending_stage, confirmed_interrupt=False):
    return bool(str(pending_stage or "").strip()) and not bool(confirmed_interrupt)


def build_task_interrupt_notice(task_name, pending_stage):
    stage_labels = {
        "draft_generation": "\u5199\u51fa\u7a3f",
        "review_generation": "\u4e25\u683c\u5ba1\u7a3f",
        "modification_generation": "\u4fee\u6539\u7a3f\u4ef6",
        "de_ai_generation": "\u53bb AI \u5473\u91cd\u5199",
    }
    clean_task_name = str(task_name or "\u5f53\u524d\u4efb\u52a1").strip() or "\u5f53\u524d\u4efb\u52a1"
    clean_stage = str(pending_stage or "").strip()
    stage_label = stage_labels.get(clean_stage, clean_stage or "AI \u4efb\u52a1")
    return f"{clean_task_name} \u6b63\u5728\u6267\u884c\u300c{stage_label}\u300d\u3002\u5982\u679c\u73b0\u5728\u5207\u6362\u3001\u65b0\u5efa\u3001\u590d\u5236\u6216\u5957\u7528\u6a21\u677f\uff0c\u8fd9\u4e00\u8f6e\u540c\u6b65\u8c03\u7528\u4f1a\u88ab\u4e2d\u65ad\u3002"


def render_task_queue_panel():
    init_task_queue_state()
    tasks = st.session_state.get("task_queue", []) or []
    if not tasks:
        return

    queue_notice = st.session_state.pop(TASK_QUEUE_NOTICE_KEY, "")
    if queue_notice:
        st.success(queue_notice)

    active_task_id = st.session_state.get("active_task_id", "") or tasks[0].get("id", "")
    st.session_state.active_task_id = active_task_id
    active_task = get_task_by_id(active_task_id) or tasks[0]
    metrics = build_task_queue_metrics(tasks)
    pending_stage = str(st.session_state.get("pending_ai_stage", "") or "").strip()

    render_section_intro("任务队列", "将单篇工作流升级为可连续处理的本地生产队列。", "Queue")
    with st.container(border=True):
        metric_cols = st.columns(6)
        metric_values = [
            ("全部", metrics.get("total", 0)),
            ("待处理", metrics.get("pending", 0)),
            ("进行中", metrics.get("in_progress", 0)),
            ("待人工确认", metrics.get("needs_review", 0)),
            ("已完成", metrics.get("completed", 0)),
            ("失败", metrics.get("failed", 0)),
        ]
        for col, (label, value) in zip(metric_cols, metric_values):
            with col:
                st.metric(label, value)

        interrupt_confirmed = False
        if pending_stage:
            st.warning(build_task_interrupt_notice(active_task.get("name", active_task_id), pending_stage))
            interrupt_confirmed = st.checkbox(
                "我确认切换并中断当前任务",
                key="task_queue_interrupt_confirm",
            )
        else:
            st.session_state["task_queue_interrupt_confirm"] = False

        task_action_blocked = is_task_action_blocked(pending_stage, interrupt_confirmed)

        task_options = [task.get("id", "") for task in tasks]
        task_labels = {
            task.get("id", ""): f"[{TASK_STATUS_LABELS.get(task.get('status', 'pending'), task.get('status', 'pending'))}] {task.get('name', task.get('id', '??'))}"
            for task in tasks
        }
        current_index = task_options.index(active_task_id) if active_task_id in task_options else 0

        selected_task_id = st.selectbox(
            "切换任务",
            options=task_options,
            index=current_index,
            format_func=lambda task_id: task_labels.get(task_id, task_id),
            key="task_queue_selected_id",
        )

        st.markdown(
            "<p class='toolbar-note'>先定位任务，再决定是切换、派生新任务，还是沿当前步骤继续处理。</p>",
            unsafe_allow_html=True,
        )
        render_context_strip([
            f"当前任务：{active_task.get('name', active_task_id)}",
            f"状态：{TASK_STATUS_LABELS.get(active_task.get('status', 'pending'), active_task.get('status', 'pending'))}",
            f"Step {active_task.get('current_step', 1)}",
            f"最近更新：{active_task.get('updated_at', '')}",
        ])

        switch_disabled = selected_task_id == active_task_id or task_action_blocked
        action_cols = st.columns([1.15, 0.95, 0.95, 0.95, 0.92])
        with action_cols[0]:
            if st.button(
                "切换到所选任务",
                key="switch_selected_task",
                type="primary",
                use_container_width=True,
                disabled=switch_disabled,
            ):
                switch_to_task(selected_task_id)
                st.rerun()
        with action_cols[1]:
            if st.button("新建空白", key="create_blank_task", use_container_width=True, disabled=task_action_blocked):
                create_task_from_current_state(clone_current=False)
                st.rerun()
        with action_cols[2]:
            if st.button("复制当前", key="clone_current_task", use_container_width=True, disabled=task_action_blocked):
                create_task_from_current_state(clone_current=True)
                st.rerun()
        with action_cols[3]:
            resume_label = "继续处理"
            if active_task.get("status") == "failed":
                resume_label = "失败重试"
            if st.button(resume_label, key="resume_current_task", use_container_width=True):
                if resume_task(active_task_id):
                    st.rerun()
        with action_cols[4]:
            delete_confirmed = st.session_state.get("task_queue_delete_confirm", False)
            delete_disabled = task_action_blocked or len(tasks) <= 1 or not delete_confirmed
            if st.button("删除任务", key="delete_selected_task", use_container_width=True, disabled=delete_disabled):
                if delete_task(selected_task_id):
                    st.session_state["task_queue_delete_confirm"] = False
                    st.rerun()

        if len(tasks) <= 1:
            st.caption("队列里至少保留一个任务；如需重做当前条目，可以直接新建空白后再删除旧任务。")
        else:
            delete_target = task_labels.get(selected_task_id, selected_task_id)
            confirm_label = f"确认删除所选任务：{delete_target}"
            st.checkbox(confirm_label, key="task_queue_delete_confirm")

        if selected_task_id != active_task_id:
            if task_action_blocked:
                st.caption("运行中任务需要先勾选确认，才能执行会中断当前调用的操作。")
            else:
                st.caption("切换后会将当前编辑状态同步到所选任务。")
        else:
            st.caption("当前已定位到正在编辑的任务，可以直接继续处理或新建旁支任务。")

        with st.expander("任务模板", expanded=False):
            st.caption("把常用配置存为模板，可以快速复用到后续任务。")
            st.markdown("**保存为模板**")
            template_name_col, template_save_col = st.columns([2.4, 1])
            with template_name_col:
                template_name = st.text_input("模板名称", key="task_template_name_input", placeholder="例：标准游戏新闻流程")
            with template_save_col:
                st.caption(" ")
                if st.button("保存当前配置", key="save_task_template", use_container_width=True):
                    saved_template_id = save_current_config_as_template(template_name)
                    if saved_template_id:
                        st.session_state[TASK_QUEUE_NOTICE_KEY] = f"已保存模板：{template_name.strip()}"
                        st.rerun()

            templates = st.session_state.get("task_templates", []) or []
            if templates:
                st.markdown("**应用已有模板**")
                template_options = [item.get("id", "") for item in templates]
                render_context_strip([f"可用模板 {len(templates)} 个", "默认会勾选当前任务"])
                template_picker_col, apply_targets_col = st.columns([1.2, 1.8])
                with template_picker_col:
                    selected_template_id = st.selectbox(
                        "选择模板",
                        options=template_options,
                        format_func=lambda template_id: next((item.get("name", template_id) for item in templates if item.get("id") == template_id), template_id),
                        key="selected_task_template_id",
                    )
                with apply_targets_col:
                    apply_targets = st.multiselect(
                        "应用到任务",
                        options=task_options,
                        default=[active_task_id],
                        format_func=lambda task_id: task_labels.get(task_id, task_id),
                        key="task_template_apply_targets",
                    )

                apply_action_col, apply_hint_col = st.columns([1, 2.2])
                with apply_action_col:
                    if st.button("应用模板", key="apply_task_template_btn", use_container_width=True, disabled=task_action_blocked):
                        updated_count = apply_template_to_tasks(selected_template_id, apply_targets)
                        if updated_count:
                            st.session_state[TASK_QUEUE_NOTICE_KEY] = f"已将模板应用到 {updated_count} 个任务"
                            st.rerun()
                with apply_hint_col:
                    if task_action_blocked:
                        st.caption("运行中任务需先确认中断，才能批量套用模板。")
                    else:
                        st.caption("模板只会更新配置项，不会覆盖已生成的成果。")
            else:
                st.info("暂时还没有任务模板，可以先保存一个常用配置。")

        with st.expander("任务列表 / 交付台", expanded=False):
            filter_options = ["全部", "待处理", "进行中", "待人工确认", "已完成", "失败"]
            st.markdown("<p class='toolbar-note'>在这里快速筛选任务状态、检查产物完整度，并按完成项批量导出。</p>", unsafe_allow_html=True)
            selected_filter = st.selectbox("筛选状态", filter_options, key="task_filter_status")
            reverse_status_map = {label: key for key, label in TASK_STATUS_LABELS.items()}
            filtered_tasks = []
            for task in tasks:
                if selected_filter == "全部":
                    filtered_tasks.append(task)
                    continue
                if task.get("status") == reverse_status_map.get(selected_filter):
                    filtered_tasks.append(task)

            for task in filtered_tasks:
                metric_info = task.get("metrics", {}) or {}
                artifact_flags = []
                if metric_info.get("has_highlight"):
                    artifact_flags.append("高亮版")
                if metric_info.get("has_podcast"):
                    artifact_flags.append("播客")
                if metric_info.get("version_count"):
                    artifact_flags.append(f"版本 {metric_info.get('version_count')}")
                artifact_text = " / ".join(artifact_flags) if artifact_flags else "暂无产物"
                st.markdown(
                    f"**{task.get('name', task.get('id', '??'))}**  \n"
                    f"状态：{TASK_STATUS_LABELS.get(task.get('status', 'pending'), task.get('status', 'pending'))} | "
                    f"Step {task.get('current_step', 1)} | 字数：{metric_info.get('word_count', 0)} | {artifact_text} | 更新：{task.get('updated_at', '')}"
                )

            completed_tasks = [task for task in tasks if task.get("status") == "completed"]
            if completed_tasks:
                completed_ids = [task.get("id", "") for task in completed_tasks]
                selected_export_ids = st.multiselect(
                    "选择要批量导出的任务",
                    options=completed_ids,
                    format_func=lambda task_id: task_labels.get(task_id, task_id),
                    key="task_queue_export_ids",
                )
                export_targets = [task for task in completed_tasks if task.get("id") in selected_export_ids]
                export_text = build_batch_export_markdown(export_targets)
                st.download_button(
                    "导出 Markdown 合集",
                    data=export_text.encode("utf-8"),
                    file_name=f"article-task-batch-{datetime.now().strftime('%Y%m%d-%H%M%S')}.md",
                    mime="text/markdown",
                    disabled=not bool(export_targets),
                    use_container_width=True,
                    key="download_task_batch_export",
                )
            else:
                st.caption("还没有已完成的任务，暂时无法批量导出。")

def go_to_step(step):
    st.session_state.current_step = step
    save_draft()


def sync_selected_source_images():
    all_images = st.session_state.get("source_images_all", [])
    if not isinstance(all_images, list):
        all_images = []

    raw_selected = st.session_state.get("selected_source_image_ids", [])
    normalized_selected = []
    seen_ids = set()

    for item in raw_selected:
        try:
            idx = int(item)
        except Exception:
            continue
        if idx < 0 or idx >= len(all_images):
            continue
        if idx in seen_ids:
            continue
        seen_ids.add(idx)
        normalized_selected.append(idx)

    st.session_state.source_images_all = all_images
    st.session_state.selected_source_image_ids = normalized_selected
    st.session_state.source_images = [all_images[idx] for idx in normalized_selected]


sync_selected_source_images()


def apply_selected_source_images(selected_ids, reset_widget_state=False):
    all_images = st.session_state.get("source_images_all", [])
    if not isinstance(all_images, list):
        all_images = []

    normalized_selected = []
    seen_ids = set()
    for item in selected_ids:
        try:
            idx = int(item)
        except Exception:
            continue
        if idx < 0 or idx >= len(all_images):
            continue
        if idx in seen_ids:
            continue
        seen_ids.add(idx)
        normalized_selected.append(idx)

    st.session_state.selected_source_image_ids = normalized_selected
    st.session_state.source_images = [all_images[idx] for idx in normalized_selected]

    if reset_widget_state:
        for idx in range(len(all_images)):
            checkbox_key = f"source_image_pick_{idx}"
            if checkbox_key in st.session_state:
                del st.session_state[checkbox_key]


def ensure_article_version_state():
    if 'article_versions' not in st.session_state or not isinstance(st.session_state.article_versions, list):
        st.session_state.article_versions = []
    if 'active_article_version_id' not in st.session_state:
        st.session_state.active_article_version_id = None


def get_article_version_by_id(version_id):
    ensure_article_version_state()
    for version_item in st.session_state.article_versions:
        if version_item.get("id") == version_id:
            return version_item
    return None


def restore_article_version_to_session(version_item, fallback_titles=None):
    if not isinstance(version_item, dict):
        return False

    selected_titles, _, selected_article_text = parse_article_generation_response(
        version_item.get("content", ""),
        fallback_titles,
    )
    st.session_state.title_candidates = selected_titles
    st.session_state.final_article = selected_article_text
    st.session_state.highlighted_article = (version_item.get("highlighted_article") or "").strip()
    st.session_state.spoken_script = ""
    reset_podcast_outputs(delete_audio=True)
    if version_item.get("id"):
        st.session_state.active_article_version_id = version_item.get("id")
    refresh_obsidian_influence_map(force=True)
    refresh_evidence_map(force=True)
    return True


def append_article_version(content, stage, role=None, model=None, parent_id=None, highlighted_article=None):
    ensure_article_version_state()
    clean_content = (content or "").strip()
    if not clean_content:
        return None

    clean_highlighted_article = (highlighted_article or "").strip()
    versions = st.session_state.article_versions
    if versions:
        last_version = versions[-1]
        if last_version.get("stage") == stage and (last_version.get("content") or "").strip() == clean_content:
            if clean_highlighted_article or "highlighted_article" not in last_version:
                last_version["highlighted_article"] = clean_highlighted_article
            st.session_state.active_article_version_id = last_version.get("id")
            return last_version.get("id")

    version_id = f"V{len(versions) + 1:03d}"
    resolved_parent_id = parent_id if parent_id else st.session_state.get("active_article_version_id")
    version_item = {
        "id": version_id,
        "stage": stage,
        "content": clean_content,
        "role": role if role is not None else st.session_state.get("selected_role", ""),
        "model": model if model is not None else "",
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "word_count": len(clean_content),
        "parent_id": resolved_parent_id,
        "highlighted_article": clean_highlighted_article,
    }

    versions.append(version_item)
    if len(versions) > 30:
        versions = versions[-30:]
    st.session_state.article_versions = versions
    st.session_state.active_article_version_id = version_id
    return version_id


def bootstrap_article_versions():
    ensure_article_version_state()
    versions = st.session_state.article_versions

    if versions:
        active_id = st.session_state.get("active_article_version_id")
        if not active_id or not any(item.get("id") == active_id for item in versions):
            st.session_state.active_article_version_id = versions[-1].get("id")
        return

    final_text = (st.session_state.get("final_article") or "").strip()
    if final_text:
        append_article_version(
            final_text,
            "历史恢复定稿",
            model="",
            highlighted_article=st.session_state.get("highlighted_article", ""),
        )
        return

    draft_text = (st.session_state.get("draft_article") or "").strip()
    if draft_text:
        append_article_version(draft_text, "历史恢复初稿", model="")


bootstrap_article_versions()
ensure_task_queue_bootstrap()
recover_ai_progress_if_needed()
run_obsidian_retrieval()


def render_html_iframe(html_content, *, height=150, width=None, scrolling=False):
    iframe_html = "<!DOCTYPE html><html><head><meta charset='utf-8'></head><body style='margin:0;padding:0;'>" + html_content + "</body></html>"
    iframe_src = "data:text/html;base64," + base64.b64encode(iframe_html.encode("utf-8")).decode("ascii")
    iframe_renderer = getattr(components, "iframe", None)
    if callable(iframe_renderer):
        try:
            iframe_renderer(iframe_src, height=height, width=width, scrolling=scrolling)
            return
        except Exception:
            pass
    html_renderer = getattr(components, "html", None)
    if callable(html_renderer):
        try:
            html_renderer(iframe_html, height=height, width=width, scrolling=scrolling)
        except Exception:
            pass

def render_responsive_image(image_payload):
    try:
        st.image(image_payload, width="stretch")
    except TypeError:
        st.image(image_payload, use_container_width=True)


def play_step_completion_sound():
    render_html_iframe(
        """
        <script>
        (function () {
            try {
                const AudioContextRef = window.AudioContext || window.webkitAudioContext;
                if (!AudioContextRef) return;
                const ctx = new AudioContextRef();
                const playChime = (startOffset, freq, duration, peakGain) => {
                    const startAt = ctx.currentTime + startOffset;
                    const endAt = startAt + duration;
                    const mainOsc = ctx.createOscillator();
                    const harmonicOsc = ctx.createOscillator();
                    const mainGain = ctx.createGain();
                    const harmonicGain = ctx.createGain();
                    const masterGain = ctx.createGain();

                    mainOsc.type = "sine";
                    harmonicOsc.type = "sine";
                    mainOsc.frequency.setValueAtTime(freq, startAt);
                    harmonicOsc.frequency.setValueAtTime(freq * 2, startAt);

                    harmonicGain.gain.setValueAtTime(0.32, startAt);
                    masterGain.gain.setValueAtTime(0.0001, startAt);
                    masterGain.gain.exponentialRampToValueAtTime(peakGain, startAt + 0.012);
                    masterGain.gain.exponentialRampToValueAtTime(peakGain * 0.48, startAt + 0.06);
                    masterGain.gain.exponentialRampToValueAtTime(0.0001, endAt);

                    mainOsc.connect(mainGain);
                    harmonicOsc.connect(harmonicGain);
                    mainGain.connect(masterGain);
                    harmonicGain.connect(masterGain);
                    masterGain.connect(ctx.destination);

                    mainOsc.start(startAt);
                    harmonicOsc.start(startAt);
                    mainOsc.stop(endAt + 0.02);
                    harmonicOsc.stop(endAt + 0.02);
                };

                // iOS 风格：清脆、短促、偏高频的双音提示
                playChime(0.00, 1318.5, 0.18, 0.11);
                playChime(0.13, 1760.0, 0.24, 0.10);
                setTimeout(() => {
                    try { ctx.close(); } catch (e) {}
                }, 700);
            } catch (e) {}
        })();
        </script>
        """,
        height=0,
    )


def play_ai_error_sound():
    render_html_iframe(
        """
        <script>
        (function () {
            try {
                const AudioContextRef = window.AudioContext || window.webkitAudioContext;
                if (!AudioContextRef) return;
                const ctx = new AudioContextRef();
                const playAlert = (startOffset, startFreq, endFreq, duration, peakGain, type) => {
                    const startAt = ctx.currentTime + startOffset;
                    const endAt = startAt + duration;
                    const osc = ctx.createOscillator();
                    const gain = ctx.createGain();
                    osc.type = type || "sawtooth";
                    osc.frequency.setValueAtTime(startFreq, startAt);
                    osc.frequency.exponentialRampToValueAtTime(endFreq, endAt);
                    gain.gain.setValueAtTime(0.0001, startAt);
                    gain.gain.exponentialRampToValueAtTime(peakGain, startAt + 0.015);
                    gain.gain.exponentialRampToValueAtTime(0.0001, endAt);
                    osc.connect(gain);
                    gain.connect(ctx.destination);
                    osc.start(startAt);
                    osc.stop(endAt + 0.03);
                };

                playAlert(0.00, 880.0, 440.0, 0.22, 0.09, "square");
                playAlert(0.20, 880.0, 440.0, 0.22, 0.09, "square");
                playAlert(0.42, 660.0, 330.0, 0.32, 0.08, "sawtooth");
                setTimeout(() => {
                    try { ctx.close(); } catch (e) {}
                }, 1200);
            } catch (e) {}
        })();
        </script>
        """,
        height=0,
    )


def normalize_copy_text(text):
    normalized = (text or "")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\u2028", "\n").replace("\u2029", "\n")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _render_inline_markdown_for_clipboard(text):
    escaped = html_lib.escape(text)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"\*(.+?)\*", r"<em>\1</em>", escaped)
    return escaped.replace("\n", "<br/>")


def build_clipboard_html_document(fragment_html):
    fragment = fragment_html or "<div><br/></div>"
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>"
        "<!--StartFragment-->"
        f"{fragment}"
        "<!--EndFragment-->"
        "</body></html>"
    )


HTML_CLIPBOARD_FORMAT_NAME = "HTML Format"


def build_windows_html_clipboard_payload(document_html):
    html_text = (document_html or "").strip()
    if not html_text:
        html_text = build_clipboard_html_document("<div><br/></div>")

    html_bytes = html_text.encode("utf-8")
    start_fragment_marker = b"<!--StartFragment-->"
    end_fragment_marker = b"<!--EndFragment-->"
    start_fragment_index = html_bytes.find(start_fragment_marker)
    end_fragment_index = html_bytes.find(end_fragment_marker)
    if start_fragment_index == -1 or end_fragment_index == -1:
        raise ValueError("Clipboard HTML is missing fragment markers.")

    header_template = (
        "Version:0.9\r\n"
        "StartHTML:{start_html:010d}\r\n"
        "EndHTML:{end_html:010d}\r\n"
        "StartFragment:{start_fragment:010d}\r\n"
        "EndFragment:{end_fragment:010d}\r\n"
        "StartSelection:{start_fragment:010d}\r\n"
        "EndSelection:{end_fragment:010d}\r\n"
    )

    placeholder_header = header_template.format(
        start_html=0,
        end_html=0,
        start_fragment=0,
        end_fragment=0,
    )
    start_html = len(placeholder_header.encode("ascii"))
    start_fragment = start_html + start_fragment_index + len(start_fragment_marker)
    end_fragment = start_html + end_fragment_index
    end_html = start_html + len(html_bytes)
    header = header_template.format(
        start_html=start_html,
        end_html=end_html,
        start_fragment=start_fragment,
        end_fragment=end_fragment,
    )
    return header.encode("ascii") + html_bytes


def prepare_highlighted_html_for_clipboard(html_text):
    normalized = sanitize_highlighted_article(html_text)
    if not normalized:
        return ""

    def replace_highlight_span(match, *, tone):
        content = (match.group("content") or "").strip()
        if not content:
            return ""

        if tone == "positive":
            background = "#d8e7ff"
            text_color = "#1f57b8"
            border = "#9dbdff"
        else:
            background = "#fde1e1"
            text_color = "#b3261e"
            border = "#f3b2b2"

        return (
            f'<mark data-highlight-tone="{tone}" '
            f'style="background:{background};color:{text_color};padding:0.08rem 0.3rem;'
            f'border-radius:0.38rem;border:1px solid {border};font-weight:700;'
            f'text-decoration:none;box-decoration-break:clone;-webkit-box-decoration-break:clone;">'
            f'<strong style="color:{text_color};font-weight:700;">{content}</strong>'
            f'</mark>'
        )

    highlight_patterns = {
        "positive": re.compile(
            r"<span\b(?=[^>]*\bclass\s*=\s*['\"][^'\"]*\bhighlight-positive\b[^'\"]*['\"])[^>]*>(?P<content>.*?)</span>",
            flags=re.IGNORECASE | re.DOTALL,
        ),
        "risk": re.compile(
            r"<span\b(?=[^>]*\bclass\s*=\s*['\"][^'\"]*\bhighlight-risk\b[^'\"]*['\"])[^>]*>(?P<content>.*?)</span>",
            flags=re.IGNORECASE | re.DOTALL,
        ),
    }
    normalized = highlight_patterns["positive"].sub(lambda match: replace_highlight_span(match, tone="positive"), normalized)
    normalized = highlight_patterns["risk"].sub(lambda match: replace_highlight_span(match, tone="risk"), normalized)
    normalized = re.sub(r'<h2>', '<h2 style="font-size:28px;line-height:1.35;font-weight:800;margin:0 0 16px 0;color:#16211c;">', normalized, flags=re.IGNORECASE)
    normalized = re.sub(r'<h3>', '<h3 style="font-size:22px;line-height:1.4;font-weight:700;margin:0 0 14px 0;color:#16211c;">', normalized, flags=re.IGNORECASE)
    normalized = re.sub(r'<p>', '<p style="margin:0 0 16px 0;line-height:1.9;color:#16211c;">', normalized, flags=re.IGNORECASE)
    return build_clipboard_html_document(normalized)


def copy_rich_content_to_system_clipboard(plain_text, html_document=""):
    plain_text = normalize_copy_text(plain_text)
    html_document = (html_document or "").strip()
    if not plain_text:
        return False, "\u6ca1\u6709\u53ef\u590d\u5236\u7684\u5185\u5bb9\u3002"

    if os.name != "nt":
        return False, "\u5f53\u524d\u4ec5\u4e3a Windows \u672c\u673a\u526a\u8d34\u677f\u63d0\u4f9b\u4e00\u952e\u590d\u5236\u3002"

    user32 = ctypes.windll.user32
    kernel32 = ctypes.windll.kernel32
    gm_moveable = 0x0002
    cf_unicode_text = 13
    user32.RegisterClipboardFormatW.argtypes = [ctypes.c_wchar_p]
    user32.RegisterClipboardFormatW.restype = ctypes.c_uint
    user32.OpenClipboard.argtypes = [ctypes.c_void_p]
    user32.OpenClipboard.restype = ctypes.c_bool
    user32.EmptyClipboard.argtypes = []
    user32.EmptyClipboard.restype = ctypes.c_bool
    user32.CloseClipboard.argtypes = []
    user32.CloseClipboard.restype = ctypes.c_bool
    user32.SetClipboardData.argtypes = [ctypes.c_uint, ctypes.c_void_p]
    html_format = user32.RegisterClipboardFormatW(HTML_CLIPBOARD_FORMAT_NAME)

    kernel32.GlobalAlloc.argtypes = [ctypes.c_uint, ctypes.c_size_t]
    kernel32.GlobalAlloc.restype = ctypes.c_void_p
    kernel32.GlobalLock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalLock.restype = ctypes.c_void_p
    kernel32.GlobalUnlock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalUnlock.restype = ctypes.c_bool
    kernel32.GlobalFree.argtypes = [ctypes.c_void_p]
    kernel32.GlobalFree.restype = ctypes.c_void_p
    user32.SetClipboardData.restype = ctypes.c_void_p

    def set_clipboard_block(fmt, payload):
        handle = kernel32.GlobalAlloc(gm_moveable, len(payload))
        if not handle:
            raise OSError("GlobalAlloc failed")
        locked = kernel32.GlobalLock(handle)
        if not locked:
            kernel32.GlobalFree(handle)
            raise OSError("GlobalLock failed")
        ctypes.memmove(locked, payload, len(payload))
        kernel32.GlobalUnlock(handle)
        if not user32.SetClipboardData(fmt, handle):
            kernel32.GlobalFree(handle)
            raise OSError("SetClipboardData failed")

    plain_payload = plain_text.encode("utf-16le") + b"\\x00\\x00".decode("unicode_escape").encode("latin1")
    html_payload = build_windows_html_clipboard_payload(html_document) if html_document else b""

    if not user32.OpenClipboard(None):
        return False, "\u65e0\u6cd5\u6253\u5f00\u7cfb\u7edf\u526a\u8d34\u677f\u3002"
    try:
        if not user32.EmptyClipboard():
            return False, "\u65e0\u6cd5\u6253\u5f00\u7cfb\u7edf\u526a\u8d34\u677f\u3002"
        set_clipboard_block(cf_unicode_text, plain_payload)
        if html_payload and html_format:
            set_clipboard_block(html_format, html_payload)
    except Exception as exc:
        return False, f"\u590d\u5236\u5931\u8d25\uff1a{exc}"
    finally:
        user32.CloseClipboard()

    return True, "\u5df2\u590d\u5236\u5230\u7cfb\u7edf\u526a\u8d34\u677f\uff0c\u53ef\u76f4\u63a5\u7c98\u8d34\u5230\u5bcc\u6587\u672c\u7f16\u8f91\u5668\u3002"


def markdown_to_editor_html(markdown_text):
    normalized = normalize_copy_text(markdown_text)
    if not normalized:
        return build_clipboard_html_document("<div><br/></div>")

    paragraphs = [item.strip() for item in re.split(r"\n{2,}", normalized) if item.strip()]
    if not paragraphs:
        paragraphs = [normalized]

    html_parts = []
    for paragraph in paragraphs:
        rendered_paragraph = _render_inline_markdown_for_clipboard(paragraph)
        # Use block-level divs for better compatibility with community editors.
        html_parts.append(f"<div>{rendered_paragraph}</div>")

    fragment = "".join(html_parts)
    return build_clipboard_html_document(fragment)


def html_fragment_to_plain_text(html_text):
    normalized = (html_text or "").strip()
    if not normalized:
        return ""

    normalized = re.sub(r"<\s*br\s*/?>", "\n", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"</\s*(p|div|h1|h2|h3|li)\s*>", "\n\n", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"<\s*li\b[^>]*>", "- ", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"<[^>]+>", "", normalized)
    normalized = html_lib.unescape(normalized)
    normalized = normalize_copy_text(normalized)
    return normalized


def html_fragment_to_editor_html(html_text):
    normalized = (html_text or "").strip()
    if not normalized:
        return build_clipboard_html_document("<div><br/></div>")
    return build_clipboard_html_document(normalized)


def _build_copy_button_iframe(copy_key, label, plain_text, html_payload):
    if not plain_text and not html_payload:
        return

    safe_key = re.sub(r"[^a-zA-Z0-9_-]", "_", str(copy_key))
    safe_label = html_lib.escape(label)
    plain_text_b64 = base64.b64encode((plain_text or "").encode("utf-8")).decode("ascii")
    html_text_b64 = base64.b64encode((html_payload or "").encode("utf-8")).decode("ascii")

    html_renderer = getattr(components, "html", None)
    html_payload = f"""
        <div style="display:flex;align-items:center;gap:10px;margin:6px 0 0 0;">
          <button id="copy-btn-{safe_key}" onclick="window.__copyRichText_{safe_key} && window.__copyRichText_{safe_key}(); return false;" style="border:1px solid #d9d9df;background:#ffffff;padding:6px 12px;border-radius:8px;cursor:pointer;font-size:13px;">
            {safe_label}
          </button>
          <span id="copy-status-{safe_key}" style="font-size:12px;color:#667085;"></span>
        </div>
        <script>
        (function () {{
            const plainBase64 = "{plain_text_b64}";
            const htmlBase64 = "{html_text_b64}";
            const status = document.getElementById("copy-status-{safe_key}");

            function decodeBase64Utf8(input) {{
                const binary = window.atob(input);
                const bytes = new Uint8Array(binary.length);
                for (let i = 0; i < binary.length; i += 1) {{
                    bytes[i] = binary.charCodeAt(i);
                }}
                if (window.TextDecoder) {{
                    return new TextDecoder().decode(bytes);
                }}
                let encoded = "";
                for (let i = 0; i < bytes.length; i += 1) {{
                    encoded += "%" + bytes[i].toString(16).padStart(2, "0");
                }}
                return decodeURIComponent(encoded);
            }}

            function setStatus(message) {{
                if (status) {{
                    status.textContent = message;
                }}
            }}

            function getParentDocument() {{
                try {{
                    if (window.parent && window.parent.document) {{
                        return window.parent.document;
                    }}
                }} catch (err) {{}}
                return null;
            }}

            function getParentClipboard() {{
                try {{
                    if (window.parent && window.parent.navigator && window.parent.navigator.clipboard) {{
                        return window.parent.navigator.clipboard;
                    }}
                }} catch (err) {{}}
                return null;
            }}

            function copyViaExecCommand(targetDocument, htmlText, plainText) {{
                if (!targetDocument || !targetDocument.body || !targetDocument.execCommand) {{
                    return false;
                }}
                const listener = function (event) {{
                    event.preventDefault();
                    if (event.clipboardData) {{
                        event.clipboardData.setData("text/plain", plainText);
                        if (htmlText) {{
                            event.clipboardData.setData("text/html", htmlText);
                        }}
                    }}
                }};
                targetDocument.addEventListener("copy", listener);
                let copied = false;
                try {{
                    copied = targetDocument.execCommand("copy");
                }} catch (err) {{
                    copied = false;
                }}
                targetDocument.removeEventListener("copy", listener);
                return copied;
            }}

            async function copyViaClipboardApi(clipboard, htmlText, plainText) {{
                if (!clipboard) {{
                    return false;
                }}
                try {{
                    if (window.ClipboardItem && htmlText && clipboard.write) {{
                        const payload = new ClipboardItem({{
                            "text/plain": new Blob([plainText], {{ type: "text/plain" }}),
                            "text/html": new Blob([htmlText], {{ type: "text/html" }})
                        }});
                        await clipboard.write([payload]);
                        return true;
                    }}
                    if (clipboard.writeText) {{
                        await clipboard.writeText(plainText);
                        return true;
                    }}
                }} catch (err) {{}}
                return false;
            }}

            window.__copyRichText_{safe_key} = async function () {{
                const plainText = decodeBase64Utf8(plainBase64).replace(/\n/g, "\r\n");
                const htmlText = decodeBase64Utf8(htmlBase64);
                const parentDocument = getParentDocument();
                const parentClipboard = getParentClipboard();

                if (copyViaExecCommand(document, htmlText, plainText) || copyViaExecCommand(parentDocument, htmlText, plainText)) {{
                    setStatus("复制成功，已保留格式。");
                    return;
                }}
                if (await copyViaClipboardApi(parentClipboard, htmlText, plainText) || await copyViaClipboardApi(navigator.clipboard, htmlText, plainText)) {{
                    setStatus("复制成功，已保留格式。");
                    return;
                }}
                setStatus("复制失败，请手动 Ctrl+C。");
            }};
        }})();
        </script>
        """
    if callable(html_renderer):
        try:
            html_renderer(html_payload, height=56)
            return
        except TypeError:
            html_renderer(html_payload)
            return
    render_html_iframe(html_payload, height=56)


def render_editor_friendly_copy_button(text, copy_key, label="\U0001F4CB \u517c\u5bb9\u590d\u5236\uff08\u4fdd\u7559\u6bb5\u843d\uff09"):
    plain_text = normalize_copy_text(text)
    if not plain_text:
        return

    safe_key = re.sub(r"[^a-zA-Z0-9_-]", "_", str(copy_key))
    status_key = f"copy_status_{safe_key}"
    if st.button(label, key=f"copy_btn_{safe_key}"):
        success, message = copy_rich_content_to_system_clipboard(plain_text, markdown_to_editor_html(plain_text))
        st.session_state[status_key] = ("success" if success else "error", message)

    status = st.session_state.get(status_key)
    if status:
        level, message = status
        if level == "success":
            st.caption(message)
        else:
            st.caption(message)


def render_wrapped_article_text(text):
    normalized = text or ""
    escaped = html_lib.escape(normalized)
    st.markdown(
        f'<div class="article-text-view"><pre>{escaped}</pre></div>',
        unsafe_allow_html=True,
    )


def render_rich_html_copy_button(html_text, copy_key, label="\U0001F4CB \u590d\u5236\u9ad8\u4eae\u9605\u8bfb\u7248\uff08\u4fdd\u7559\u683c\u5f0f\uff09"):
    sanitized_html = sanitize_highlighted_article(html_text)
    plain_text = html_fragment_to_plain_text(sanitized_html)
    if not plain_text:
        return

    safe_key = re.sub(r"[^a-zA-Z0-9_-]", "_", str(copy_key))
    status_key = f"copy_status_{safe_key}"
    if st.button(label, key=f"copy_btn_{safe_key}"):
        success, message = copy_rich_content_to_system_clipboard(plain_text, prepare_highlighted_html_for_clipboard(sanitized_html))
        st.session_state[status_key] = ("success" if success else "error", message)

    status = st.session_state.get(status_key)
    if status:
        level, message = status
        if level == "success":
            st.caption(message)
        else:
            st.caption(message)


def notify_step_completed(defer_until_rerun=False):


    if defer_until_rerun:
        st.session_state.pending_completion_sound = True
    else:
        play_step_completion_sound()

def render_pending_completion_sound():
    if st.session_state.get("pending_completion_sound"):
        st.session_state.pending_completion_sound = False
        play_step_completion_sound()

def get_script_sys_prompt(duration_str):
    duration_map = {
        "1分钟": "200 - 250",
        "3分钟": "600 - 700",
        "5分钟": "1000 - 1200",
        "8分钟": "1600 - 1900"
    }
    target_words = duration_map.get(duration_str, "1000 - 1200")
    
    return f"""你是一位资深的AI视频流工业化编导。请将提供给你的长篇深度文章，浓缩提炼成一份可直接输入给【剪映AI】等工具解析的【{duration_str}口播与分镜脚本】。

【强制执行规则】：
1. 语速与篇幅：正常人中文语速约为220字/分钟，因此总的口播旁白字数必须严格控制在 {target_words} 字左右（对应约{duration_str}的短视频长度）。
2. 零废话输出：因为你的输出将被下一个自动化代码节点读取，请【绝对不要】在开头或结尾输出“好的，这是为您生成的脚本”等任何人类视角的寒暄语句，直接输出 Markdown 结构！
3. 适配剪映AI：画面Prompt必须全部使用【纯中文】描述，包含明确的画面主体、场景细节和镜头动作，以便国内视频AI引擎精准生图/视频。
4. 结构化呈现：必须严格按照以下 Markdown 列表格式输出每一个镜头（Scene），绝不能混用格式。

【标准输出格式范例】：
### Scene 01 (0s - 5s)
* 🗣️ **口播旁白**: "就在昨天，游戏圈又爆出了一个令人震惊的超级大瓜！" 
* 🎬 **画面Prompt**: "电影级中景镜头，一个震惊的年轻玩家坐在昏暗的房间里看着发光的电脑屏幕，屏幕上显示着数据代码，动态光影，逼真质感，平移运镜"
* ✨ **视觉特效/字幕**: 屏幕突然亮起，居中显示大字特效“超级大瓜”。
"""

def inject_ui_theme():
    st.markdown(
        """
        <style>
        :root {
            --panel: rgba(255, 255, 255, 0.78);
            --border: rgba(41, 59, 51, 0.12);
            --text: #16211c;
            --text-muted: #5e6d66;
            --brand: #1f6f5f;
            --brand-strong: #174f44;
            --shadow-soft: 0 18px 40px rgba(22, 33, 28, 0.08);
            --shadow-strong: 0 24px 60px rgba(22, 33, 28, 0.12);
            --radius-lg: 24px;
            --radius-md: 18px;
        }
        .stApp {
            background:
                radial-gradient(circle at top left, rgba(31, 111, 95, 0.12), transparent 32%),
                radial-gradient(circle at top right, rgba(201, 139, 93, 0.10), transparent 24%),
                linear-gradient(180deg, #f4f8f5 0%, #ecf1ee 100%);
            color: var(--text);
        }
        [data-testid="stAppViewContainer"] > .main { padding-top: 1.5rem; }
        [data-testid="stHeader"] { background: rgba(244, 248, 245, 0.65); }
        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, rgba(18, 32, 27, 0.96), rgba(26, 44, 37, 0.92));
            border-right: 1px solid rgba(255, 255, 255, 0.08);
        }
        [data-testid="stSidebar"] * { color: #eef5f1; }
        [data-testid="stSidebar"] .stTextInput input,
        [data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] > div,
        [data-testid="stSidebar"] .stTextArea textarea {
            background: rgba(255, 255, 255, 0.08) !important;
            border: 1px solid rgba(255, 255, 255, 0.10) !important;
            color: #eef5f1 !important;
        }
        .block-container {
            max-width: 1440px;
            padding-top: 0.5rem;
            padding-bottom: 4rem;
        }
        .app-hero {
            margin-bottom: 1.25rem;
            padding: 2rem 2.25rem;
            border-radius: var(--radius-lg);
            border: 1px solid rgba(255, 255, 255, 0.5);
            background: linear-gradient(135deg, rgba(16, 33, 29, 0.97), rgba(31, 111, 95, 0.90));
            box-shadow: var(--shadow-strong);
            color: #f5faf7;
        }
        .app-kicker {
            display: inline-flex;
            margin-bottom: 0.85rem;
            padding: 0.35rem 0.8rem;
            border-radius: 999px;
            background: rgba(255, 255, 255, 0.14);
            border: 1px solid rgba(255, 255, 255, 0.22);
            font-size: 0.78rem;
            font-weight: 600;
            letter-spacing: 0.04em;
        }
        .app-hero h1 { margin: 0; color: #f8fdfb; font-size: 2.3rem; }
        .app-hero p {
            max-width: 840px;
            margin: 0.7rem 0 1.25rem;
            color: rgba(248, 253, 251, 0.82);
            font-size: 1rem;
            line-height: 1.7;
        }
        .hero-metrics, .step-grid, .chip-row, .mode-grid {
            display: flex;
            flex-wrap: wrap;
            gap: 0.75rem;
        }
        .metric-card, .mode-card {
            color: var(--text);
            min-width: 160px;
            padding: 0.95rem 1rem;
            border-radius: 16px;
            border: 1px solid rgba(255, 255, 255, 0.14);
            background: rgba(255, 255, 255, 0.10);
        }
        .metric-card strong, .mode-card strong {
            display: block;
            margin-bottom: 0.3rem;
            color: var(--text);
            font-size: 1.05rem;
        }
        .metric-card span, .mode-card span {
            color: var(--text-muted);
            font-size: 0.86rem;
            line-height: 1.5;
        }
        .stepper {
            margin: 0.5rem 0 1rem;
            padding: 1rem 1.1rem;
            border-radius: var(--radius-md);
            border: 1px solid var(--border);
            background: rgba(255, 255, 255, 0.65);
            box-shadow: var(--shadow-soft);
        }
        .stepper-item {
            flex: 1 1 150px;
            min-width: 120px;
            padding: 0.9rem 1rem;
            border-radius: 16px;
            border: 1px solid var(--border);
            background: rgba(244, 248, 246, 0.88);
        }
        .stepper-item.active {
            background: linear-gradient(135deg, rgba(31, 111, 95, 0.12), rgba(201, 139, 93, 0.12));
            border-color: rgba(31, 111, 95, 0.28);
        }
        .stepper-item.done {
            background: rgba(31, 138, 85, 0.08);
            border-color: rgba(31, 138, 85, 0.20);
        }
        .step-index {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 28px;
            height: 28px;
            margin-bottom: 0.45rem;
            border-radius: 50%;
            background: rgba(22, 33, 28, 0.08);
            color: var(--brand-strong);
            font-size: 0.82rem;
            font-weight: 700;
        }
        .stepper-item.active .step-index { background: var(--brand); color: #ffffff; }
        .step-label { display: block; margin-bottom: 0.2rem; color: var(--text); font-weight: 700; }
        .step-desc { color: var(--text-muted); font-size: 0.82rem; line-height: 1.5; }
        .section-head { display: flex; align-items: start; justify-content: space-between; gap: 1rem; margin-bottom: 0.8rem; }
        .section-title { margin: 0; font-size: 1.05rem; font-weight: 700; color: var(--text); }
        .section-subtitle { margin: 0.25rem 0 0; color: var(--text-muted); font-size: 0.92rem; line-height: 1.6; }
        .eyebrow {
            display: inline-flex;
            margin-bottom: 0.35rem;
            color: var(--brand);
            font-size: 0.78rem;
            font-weight: 700;
            letter-spacing: 0.05em;
            text-transform: uppercase;
        }
        .chip {
            display: inline-flex;
            align-items: center;
            padding: 0.42rem 0.72rem;
            border-radius: 999px;
            border: 1px solid var(--border);
            background: rgba(255, 255, 255, 0.65);
            color: var(--text-muted);
            font-size: 0.8rem;
            font-weight: 600;
        }
        .chip.active {
            background: rgba(31, 111, 95, 0.10);
            border-color: rgba(31, 111, 95, 0.22);
            color: var(--brand-strong);
        }
        .context-strip {
            margin: 0.2rem 0 1rem;
            padding: 0.95rem 1.1rem;
            border-radius: 14px;
            border: 1px solid rgba(31, 111, 95, 0.12);
            background: rgba(31, 111, 95, 0.06);
        }
        .toolbar-note { color: var(--text-muted); font-size: 0.9rem; line-height: 1.6; }
        .stButton > button, .stDownloadButton > button {
            min-height: 2.85rem;
            border-radius: 14px;
            border: 1px solid rgba(31, 111, 95, 0.14);
            background: linear-gradient(180deg, #ffffff, #f5f8f6);
            color: var(--text);
            font-weight: 700;
            box-shadow: 0 8px 18px rgba(22, 33, 28, 0.06);
        }
        .stButton > button[kind="primary"] {
            background: linear-gradient(135deg, var(--brand), var(--brand-strong));
            color: #ffffff;
            border-color: rgba(31, 111, 95, 0.55);
            box-shadow: 0 14px 28px rgba(31, 111, 95, 0.24);
        }
        .stTextInput input, .stTextArea textarea, .stSelectbox [data-baseweb="select"] > div {
            border-radius: 14px !important;
            border: 1px solid var(--border) !important;
            background: rgba(255, 255, 255, 0.88) !important;
        }
        .stCodeBlock, [data-testid="stCodeBlock"] {
            border-radius: 16px !important;
            border: 1px solid rgba(22, 33, 28, 0.08);
        }
        .article-text-view {
            padding: 1rem 1.15rem;
            border-radius: 16px;
            border: 1px solid rgba(22, 33, 28, 0.08);
            background: rgba(255, 255, 255, 0.82);
            box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.4);
        }
        .article-text-view pre {
            margin: 0;
            white-space: pre-wrap;
            overflow-wrap: anywhere;
            word-break: break-word;
            font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
            font-size: 0.97rem;
            line-height: 1.9;
            color: var(--text);
        }
        </style>
        """,
        unsafe_allow_html=True
    )


def render_app_hero():
    st.markdown(
        """
        <section class="app-hero">
            <div class="app-kicker">Professional Media Workspace</div>
            <h1>公众号文章生成助手</h1>
            <p>围绕游戏行业内容生产打造的一站式工作台。聚合多源素材、智能分配角色、统一审稿与定稿，并延伸到脚本、配图和分发环节。</p>
            <div class="hero-metrics">
                <div class="metric-card">
                    <strong>多源聚合</strong>
                    <span>文章链接、YouTube 字幕、网页图片统一进入同一工作流</span>
                </div>
                <div class="metric-card">
                    <strong>多角色协作</strong>
                    <span>编辑、审稿、精修、分镜脚本均可按角色和规范驱动</span>
                </div>
                <div class="metric-card">
                    <strong>一站式定稿</strong>
                    <span>从初稿到导出、飞书推送和配图建议全部在同一界面完成</span>
                </div>
            </div>
        </section>
        """,
        unsafe_allow_html=True
    )


def render_stepper(current_step):
    step_meta = [
        ("素材输入", "汇聚文章、视频与图像素材"),
        ("初稿生成", "选择角色并产出首版文章"),
        ("严格审稿", "核查事实、逻辑与风格"),
        ("定稿修订", "接收意见并形成最终版本"),
        ("分发工作台", "脚本、配图、导出与精修")
    ]
    columns = st.columns(len(step_meta))
    for idx, ((label, desc), col) in enumerate(zip(step_meta, columns), start=1):
        if idx < current_step:
            class_name = "stepper-item done"
        elif idx == current_step:
            class_name = "stepper-item active"
        else:
            class_name = "stepper-item"
        with col:
            card_html = textwrap.dedent(f"""
                <div class="{class_name}">
                    <div class="step-index">{idx}</div>
                    <span class="step-label">{label}</span>
                    <span class="step-desc">{desc}</span>
                </div>
            """).strip()
            st.markdown(card_html, unsafe_allow_html=True)


def render_section_intro(title, subtitle=None, eyebrow=None):
    if eyebrow:
        st.caption(eyebrow)
    st.markdown(f"#### {title}")
    if subtitle:
        st.caption(subtitle)


def render_context_strip(items):
    chips = "".join([f'<span class="chip active">{item}</span>' for item in items if item])
    if chips:
        st.markdown(f'<div class="context-strip"><div class="chip-row">{chips}</div></div>', unsafe_allow_html=True)
# ==========================================
# 4. 页面与工作流渲染
# ==========================================
st.set_page_config(page_title="公众号文章生成助手", page_icon="🕹️", layout="wide")
inject_ui_theme()

prompts_data = load_prompts()
render_pending_completion_sound()

with st.sidebar:
    st.markdown("## 控制面板")
    st.caption("管理模型、脚本生成策略与写作 Prompt，所有改动都会直接作用到当前工作流。")
    st.caption(f"Prompt 文件：{PROMPTS_FILE}")
    if PROMPTS_LOAD_REPORT.startswith("fallback_default"):
        st.warning("未读取到有效 prompts 配置，当前使用默认角色。请确认部署目录中的 prompts.json。")
        st.caption(PROMPTS_LOAD_REPORT)
    st.header("⚙️ 引擎设置")
    api_provider = st.selectbox("🌐 选择 API 中转站", ["BLTCY (柏拉图次元)", "DeerAPI", "云雾API"])
    
    if api_provider == "BLTCY (柏拉图次元)":
        api_key = st.text_input("🔑 输入 BLTCY Key", type="password")
        current_base_url = "https://api.bltcy.ai/v1"
        available_models = BLTCY_MODEL_OPTIONS
    elif api_provider == "云雾API":
        api_key = st.text_input("🔑 输入云雾API Key", type="password")
        current_base_url = "https://yunwu.ai/v1"
        available_models = YUNWU_MODEL_OPTIONS
    else:
        api_key = st.text_input("🔑 输入 DeerAPI Key", type="password")
        current_base_url = "https://api.deerapi.com/v1"
        available_models = DEERAPI_MODEL_OPTIONS

    preferred_default_models = {
        "BLTCY (柏拉图次元)": "qwen3.5-plus",
        "DeerAPI": "gpt-5.4",
        "云雾API": "qwen3.6-plus",
    }
    preferred_default_model = preferred_default_models.get(api_provider, available_models[0])
    default_model_idx = available_models.index(preferred_default_model) if preferred_default_model in available_models else 0

    selected_model = st.selectbox("🧠 选择驱动模型", available_models, index=default_model_idx)
    
    st.markdown("---")
    st.header("Obsidian 知识库")
    st.toggle("启用本地知识增强", key="obsidian_enabled", on_change=save_draft)
    st.text_input("知识库路径 / wiki 路径", key="obsidian_vault_path", placeholder=r"D:\Obsidian Vault", on_change=save_draft)
    st.slider("知识命中数", min_value=3, max_value=10, step=1, key="obsidian_max_hits", on_change=save_draft)
    st.toggle("显示命中详情", key="obsidian_show_hits", on_change=save_draft)
    wiki_root_preview, wiki_root_error = resolve_obsidian_wiki_root(st.session_state.get("obsidian_vault_path", ""))
    if st.session_state.get("obsidian_enabled"):
        if wiki_root_preview:
            st.caption(f"检测到 wiki 根目录：{wiki_root_preview}")
        elif wiki_root_error:
            st.caption(wiki_root_error)

    st.markdown("---")
    st.header("🎬 视频分镜设置")
    enable_script = st.toggle("启用伴生【短视频分镜脚本】", value=False)
    script_duration = st.selectbox(
        "⏱️ 设定分镜脚本目标时长",
        ["1分钟", "3分钟", "5分钟", "8分钟"],
        index=2,
        disabled=not enable_script
    )

    st.markdown("---")
    st.header("播客语音设置")
    st.toggle("启用伴生【播客解说稿】", key="podcast_enabled", on_change=save_draft)
    st.selectbox(
        "播客时长",
        ["3分钟", "5分钟", "8分钟", "12分钟"],
        key="podcast_duration",
        on_change=save_draft,
        disabled=not st.session_state.get("podcast_enabled", False)
    )
    podcast_api_key = st.text_input("输入阿里云 DashScope API Key", type="password")
    st.session_state.podcast_tts_api_key_present = bool((podcast_api_key or "").strip())
    st.caption("当前播客语音合成使用单音色 Qwen-TTS。")
    st.selectbox(
        "主播音色",
        options=list(PODCAST_VOICE_LABELS.keys()),
        format_func=lambda key: PODCAST_VOICE_LABELS.get(key, key),
        key="podcast_voice",
        on_change=save_draft,
        disabled=not st.session_state.get("podcast_enabled", False)
    )

    st.markdown("---")
    st.header("🗂️ 提示词管理中心")
    with st.expander("📝 角色与全局人设配置", expanded=False):
        tab1, tab2, tab3 = st.tabs(["✍️ 编辑人设", "🧐 审稿人设", "🌍 全局去AI味规范"])
        
        with tab1:
            action = st.radio("操作类型", ["编辑现有角色", "新增角色", "删除角色"], horizontal=True)
            if action == "编辑现有角色":
                edit_role = st.selectbox("选择编辑", list(prompts_data["editors"].keys()))
                if edit_role:
                    new_prompt = st.text_area("系统 Prompt", value=prompts_data["editors"][edit_role], height=200)
                    if st.button("💾 保存修改", key="save_edit"):
                        prompts_data["editors"][edit_role] = new_prompt
                        save_prompts(prompts_data)
                        st.success(f"已保存【{edit_role}】的修改！")
                        st.rerun()
            elif action == "新增角色":
                new_role_name = st.text_input("新角色名称 (如：独立游戏分析师)")
                new_role_prompt = st.text_area("新角色 Prompt", height=200)
                if st.button("➕ 确认新增"):
                    if new_role_name and new_role_name not in prompts_data["editors"]:
                        prompts_data["editors"][new_role_name] = new_role_prompt
                        save_prompts(prompts_data)
                        st.success(f"已成功添加角色：【{new_role_name}】！")
                        st.rerun()
                    else:
                        st.error("角色名不能为空，或该角色已存在！")
            elif action == "删除角色":
                del_role = st.selectbox("选择要删除的编辑", list(prompts_data["editors"].keys()))
                if st.button("🗑️ 确认删除", type="primary"):
                    if len(prompts_data["editors"]) > 1:
                        del prompts_data["editors"][del_role]
                        save_prompts(prompts_data)
                        st.success(f"已删除【{del_role}】！")
                        st.rerun()
                    else:
                        st.error("操作失败：必须至少保留一个编辑角色！")
                        
        with tab2:
            new_reviewer_prompt = st.text_area("主编/审稿员系统指令", value=prompts_data["reviewer"], height=300)
            if st.button("💾 保存审稿员设置"):
                prompts_data["reviewer"] = new_reviewer_prompt
                save_prompts(prompts_data)
                st.success("已更新审稿员人设！")
                st.rerun()
                
        with tab3:
            st.info("💡 此处的规范将**强制追加**到所有【编辑】的提示词末尾，用于统一全局的写作风格。")
            new_global_prompt = st.text_area("全局强制规范", value=prompts_data.get("global_instruction", ""), height=300)
            if st.button("💾 保存全局规范"):
                prompts_data["global_instruction"] = new_global_prompt
                save_prompts(prompts_data)
                st.success("已成功更新全局规范！")
                st.rerun()

render_app_hero()
render_stepper(st.session_state.current_step)
render_ai_progress_banner()
render_task_queue_panel()

# --- Step 1 ---
if st.session_state.current_step == 1:
    st.caption("STEP 01")
    st.markdown("## 素材输入中枢")

    if should_offer_draft_restore():
        with st.container(border=True):
            render_section_intro("恢复草稿", "检测到当前任务存在可恢复的旧工作现场。", "Recovery")
            st.markdown("<p class='toolbar-note'>如果你上次处理中断，可以直接恢复当前草稿；如果想重新开始，也可以清空当前任务后再继续。</p>", unsafe_allow_html=True)
            col_draft1, col_draft2 = st.columns(2)
            with col_draft1:
                if st.button("恢复草稿", type="primary", use_container_width=True):
                    if load_draft():
                        st.rerun()
                    else:
                        st.error("未找到可恢复的草稿。")
            with col_draft2:
                if st.button("抛弃旧草稿并重置当前任务", use_container_width=True):
                    clear_draft()
                    if reset_active_task_to_blank():
                        st.rerun()


    input_col1, input_col2 = st.columns(2)
    with input_col1:
        with st.container(border=True):
            render_section_intro("文章链接池", "适合导入多篇新闻、报告或博客文章；每行一个链接。", "Articles")
            article_url_input = st.text_area(
                "📝 输入文章链接 (每行一个，支持批量)",
                value=st.session_state.article_url,
                height=180,
                placeholder="https://example.com/article-1\nhttps://example.com/article-2"
            )
    with input_col2:
        with st.container(border=True):
            render_section_intro("视频链接池", "适合把 YouTube 字幕与文章一起并入同一份素材上下文。", "Videos")
            video_url_input = st.text_area(
                "📺 输入 YouTube 视频链接 (每行一个，支持批量)",
                value=st.session_state.video_url,
                height=180,
                placeholder="https://youtu.be/...\nhttps://www.youtube.com/watch?v=..."
            )

    with st.container(border=True):
        render_section_intro("文件资料上传", "支持上传本地图片、Word、Excel、TXT 等文件作为分析依据。", "Files")
        uploaded_source_files = st.file_uploader(
            "上传资料文件（支持图片、Word、Excel、TXT，可多选）",
            type=SUPPORTED_UPLOAD_EXTENSIONS,
            accept_multiple_files=True,
            key="source_files_uploader"
        )
        if uploaded_source_files:
            file_names = [file.name for file in uploaded_source_files]
            preview_names = "、".join(file_names[:6])
            if len(file_names) > 6:
                preview_names += f" 等 {len(file_names)} 个文件"
            st.caption(f"已选择 {len(file_names)} 个文件：{preview_names}")

    with st.container(border=True):
        render_section_intro("开始提取", "系统会先抓取正文和图片，再将多源素材聚合成统一工作底稿。", "Actions")
        st.markdown("<p class='toolbar-note'>建议先把主题相近的文章和视频放在同一批次里，方便后续自动路由和统一改写。</p>", unsafe_allow_html=True)
        if st.button("开始批量提取内容", type="primary", use_container_width=True):
            article_urls = [url.strip() for url in article_url_input.split('\n') if url.strip()]
            video_urls = [url.strip() for url in video_url_input.split('\n') if url.strip()]
            uploaded_file_count = len(uploaded_source_files) if uploaded_source_files else 0

            if not article_urls and not video_urls and uploaded_file_count == 0:
                st.warning("请至少输入链接或上传一个文件！")
            else:
                total_sources = len(article_urls) + len(video_urls) + uploaded_file_count
                with st.spinner(f"启动全息解析引擎，正在批量获取 {total_sources} 个素材..."):
                    combined_content = ""
                    extracted_imgs = []
                    errors = []
                    success_count = 0

                    for idx, a_url in enumerate(article_urls):
                        art_content, art_imgs, art_err = get_content_from_url(a_url)
                        if art_content:
                            combined_content += f"【文章素材 {idx+1}】来源于: {a_url}\n{art_content}\n\n================\n\n"
                            extracted_imgs.extend(art_imgs)
                            success_count += 1
                        else:
                            errors.append(f"文章 {idx+1} 提取失败: {art_err}")

                    for idx, v_url in enumerate(video_urls):
                        vid_content, vid_imgs, vid_err = get_content_from_url(v_url)
                        if vid_content:
                            combined_content += f"【视频素材 {idx+1}】来源于: {v_url}\n{vid_content}\n\n================\n\n"
                            success_count += 1
                        else:
                            errors.append(f"视频 {idx+1} 提取失败: {vid_err}")

                    uploaded_content, uploaded_imgs, uploaded_errors, uploaded_success_count = extract_content_from_uploaded_files(uploaded_source_files)
                    if uploaded_content:
                        combined_content += uploaded_content + "\n\n================\n\n"
                    if uploaded_imgs:
                        extracted_imgs.extend(uploaded_imgs)
                    if uploaded_errors:
                        errors.extend(uploaded_errors)
                    success_count += uploaded_success_count

                    deduped_imgs = []
                    seen_imgs = set()
                    for image_item in extracted_imgs:
                        if image_item in seen_imgs:
                            continue
                        seen_imgs.add(image_item)
                        deduped_imgs.append(image_item)
                    extracted_imgs = deduped_imgs[:15]

                    if combined_content.strip():
                        st.session_state.article_url = article_url_input
                        st.session_state.video_url = video_url_input
                        st.session_state.source_content = combined_content
                        st.session_state.source_images_all = extracted_imgs
                        st.session_state.selected_source_image_ids = []
                        st.session_state.source_images = []
                        for stale_key in [k for k in list(st.session_state.keys()) if k.startswith("source_image_pick_")]:
                            del st.session_state[stale_key]
                        st.session_state.extraction_success = True
                        st.session_state.obsidian_retrieval_signature = ""
                        run_obsidian_retrieval(force=True)
                        notify_step_completed()

                        if errors:
                            st.warning(f"部分内容提取成功 ({success_count}/{total_sources})，但有以下错误：\n" + "\n".join(errors))
                        else:
                            if len(extracted_imgs) > 0:
                                st.success(f"批量提取成功！共融合了 {success_count} 个素材，并提取到 {len(extracted_imgs)} 张核心配图。")
                            else:
                                st.success(f"批量提取成功！共融合了 {success_count} 个素材。（无有效配图，走纯文本模式）")
                    else:
                        st.session_state.extraction_success = False
                        reset_obsidian_context()
                        st.error("所有素材提取均失败，请检查链接、上传文件内容或网络状态。\n" + "\n".join(errors))

    if st.session_state.extraction_success:
        with st.container(border=True):
            render_section_intro("聚合素材预览", "先快速检查抓取结果，再决定是走手动精调还是全自动驾驶。", "Preview")
            all_source_images = st.session_state.get("source_images_all", [])
            if all_source_images:
                st.markdown("#### 核心配图（勾选纳入 AI 分析）")
                toolbar_col1, toolbar_col2, _ = st.columns([1, 1, 2.4])
                with toolbar_col1:
                    if st.button("一键全选", key="select_all_source_images", use_container_width=True):
                        apply_selected_source_images(range(len(all_source_images)), reset_widget_state=True)
                        save_draft()
                        st.rerun()
                with toolbar_col2:
                    if st.button("清空勾选", key="clear_source_images", use_container_width=True):
                        apply_selected_source_images([], reset_widget_state=True)
                        save_draft()
                        st.rerun()

                previous_selected_ids = set(st.session_state.get("selected_source_image_ids", []))
                selected_ids = []
                for idx, img_url in enumerate(all_source_images):
                    checkbox_key = f"source_image_pick_{idx}"
                    if checkbox_key not in st.session_state:
                        st.session_state[checkbox_key] = idx in previous_selected_ids
                    with st.container(border=True):
                        action_col, thumb_col = st.columns([1.2, 2.35])
                        with action_col:
                            checked = st.checkbox(
                                f"图片 {idx + 1} 纳入分析",
                                key=checkbox_key
                            )
                            if checked:
                                selected_ids.append(idx)
                        with thumb_col:
                            preview_image = get_image_preview_payload(img_url)
                            if preview_image is not None:
                                st.image(preview_image, width=220)
                            else:
                                st.caption("该图片暂时无法预览。")

                if selected_ids != st.session_state.get("selected_source_image_ids", []):
                    apply_selected_source_images(selected_ids)
                    save_draft()

                st.caption(f"当前已勾选 {len(st.session_state.source_images)} / {len(all_source_images)} 张图片进入分析。")
                if len(st.session_state.source_images) == 0:
                    st.info("当前未勾选任何图片，后续将按纯文本模式继续。")
                st.divider()

            st.markdown("#### 合并后的文本正文")
            preview_text = st.session_state.source_content[:1800] + "\n\n......(已省略后续内容)" if len(st.session_state.source_content) > 1800 else st.session_state.source_content
            st.code(preview_text, language="markdown")


        if st.session_state.get("obsidian_enabled"):
            with st.container(border=True):
                render_section_intro("Obsidian 检索简报", "在选择写作模式前，先检查本地知识库命中与自动生成的研究摘要。", "知识库")
                if st.button("刷新 Obsidian 检索", key="refresh_obsidian_hits", use_container_width=True):
                    run_obsidian_retrieval(force=True)
                    st.rerun()

                retrieval_error = (st.session_state.get("obsidian_retrieval_error", "") or "").strip()
                if retrieval_error:
                    st.warning(retrieval_error)

                wiki_root_display = st.session_state.get("obsidian_wiki_root", "")
                if wiki_root_display:
                    st.caption(f"Wiki 根目录：{wiki_root_display}")
                indexed_at = st.session_state.get("obsidian_last_indexed_at", "")
                if indexed_at:
                    st.caption(f"上次索引时间：{indexed_at}")

                query_terms = st.session_state.get("obsidian_query_terms", []) or []
                if query_terms:
                    st.caption("检索词：" + "、".join(query_terms))

                obsidian_hits = st.session_state.get("obsidian_hits", []) or []
                if obsidian_hits and st.session_state.get("obsidian_show_hits", True):
                    for idx, hit in enumerate(obsidian_hits, start=1):
                        with st.expander(f"{idx}. {hit.get('title', '')} | {hit.get('category_label', '')}", expanded=(idx <= 2)):
                            st.caption(f"{hit.get('path', '')} | {hit.get('modified_at', '')}")
                            st.markdown(hit.get("reason", ""))
                            st.code(hit.get("excerpt", ""), language="markdown")
                elif st.session_state.get("obsidian_enabled") and not retrieval_error:
                    st.info("当前这批素材没有匹配到相关 Obsidian 笔记，后续将按纯素材写作继续。")

                if st.session_state.get("obsidian_research_brief"):
                    st.markdown("#### 研究摘要预览")
                    st.code(st.session_state.obsidian_research_brief, language="markdown")

        with st.container(border=True):
            render_section_intro("选择工作流模式", "手动精调适合逐步把关，全自动驾驶适合快速得到高完成度定稿。", "Workflow")
            st.caption(f"🧮 当前全局目标字数：约 {get_target_article_words()} 字")
            col_flow1, col_flow2 = st.columns(2)

            with col_flow1:
                st.markdown("<p class='toolbar-note'>逐步确认编辑角色、初稿、审稿意见与修改结果，适合需要人工把关的稿件。</p>", unsafe_allow_html=True)
                if st.button("👉 手动精调模式 (逐步确认)", use_container_width=True):
                    go_to_step(2)
                    st.rerun()

            with col_flow2:
                st.markdown("<p class='toolbar-note'>自动完成角色路由、写稿、审稿、改稿和可选脚本生成，适合快速交付。</p>", unsafe_allow_html=True)
                if st.button("🚀 一键全自动驾驶 (AI路由直达定稿)", type="primary", use_container_width=True):
                    if not api_key:
                        st.error("⚠️ 请先在左侧边栏输入 API Key！")
                        st.stop()

                    with st.status("🤖 全自动驾驶已启动，AI 正在接管工作流...", expanded=True) as status:
                        st.write("🔍 正在分析素材内容，为您匹配最佳编辑人设...")
                        editor_names = list(prompts_data["editors"].keys())
                        routing_prompt = f"""你是一个智能路由系统。请阅读以下素材，判断哪种身份最适合将其改写为深度文章。
                        请只输出角色的完整名称，绝不允许包含任何其他标点或解释废话！
                        可选角色：{', '.join(editor_names)}"""

                        chosen_editor_raw = call_llm(
                            api_key=api_key, base_url=current_base_url, model_name=selected_model,
                            system_prompt=routing_prompt, user_content=st.session_state.source_content[:5000]
                        )

                        chosen_editor = chosen_editor_raw.strip() if chosen_editor_raw else ""
                        if chosen_editor not in editor_names:
                            chosen_editor = editor_names[0]

                        st.session_state.selected_role = chosen_editor
                        st.write(f"✅ 意图识别完成，已自动指派：**【{chosen_editor}】**")

                        st.write("✍️ 编辑正在奋笔疾书，生成初稿中...")
                        editor_prompt = prompts_data["editors"][chosen_editor]
                        global_instruction = prompts_data.get("global_instruction", "")
                        final_editor_system_prompt = build_editor_system_prompt(editor_prompt, global_instruction)

                        draft_content = build_editor_user_content(
                            st.session_state.source_content,
                            st.session_state.get("obsidian_research_brief", ""),
                            use_images=bool(st.session_state.source_images)
                        )
                        draft_response = call_llm(
                            api_key=api_key, base_url=current_base_url, model_name=selected_model,
                            system_prompt=final_editor_system_prompt, user_content=draft_content, image_urls=st.session_state.source_images
                        )
                        draft_titles, _, draft_article_text = parse_article_generation_response(
                            draft_response,
                            st.session_state.get("title_candidates", []),
                        )
                        st.session_state.title_candidates = draft_titles
                        st.session_state.draft_article = draft_article_text
                        append_article_version(st.session_state.draft_article, "自动驾驶初稿", role=chosen_editor, model=selected_model)
                        save_draft()

                        st.write("🧐 审稿主编介入，正在极其严苛地核对原文与逻辑...")
                        reviewer_prompt = prompts_data["reviewer"]
                        anti_hallucination_instruction = "\n\n【⚠️ 强制系统级指令：严禁幻觉】：你在审查事实时，**必须且只能**基于下方提供给你的【原始素材文本】！绝对不允许使用自身知识库进行事实核对。"
                        final_reviewer_system_prompt = build_reviewer_system_prompt(reviewer_prompt, anti_hallucination_instruction)

                        combined_content = build_reviewer_user_content(
                            st.session_state.source_content,
                            st.session_state.draft_article,
                            st.session_state.get("obsidian_research_brief", ""),
                            title_candidates=st.session_state.get("title_candidates", []),
                        )
                        st.session_state.review_feedback = call_llm(
                            api_key=api_key, base_url=current_base_url, model_name=selected_model,
                            system_prompt=final_reviewer_system_prompt, user_content=combined_content, image_urls=st.session_state.source_images
                        )
                        hydrate_review_action_state(st.session_state.review_feedback, reset_selection=True)

                        save_draft()
                        st.write("✨ 接收修改意见，正在进行最终打磨...")
                        modification_banned_terms, modification_default_replacements, _, _ = resolve_active_term_rules("modification")
                        modification_term_rules_instruction = build_term_rules_instruction(
                            modification_banned_terms,
                            modification_default_replacements,
                        )
                        modification_prompt = build_modification_system_prompt(
                            global_instruction,
                            term_rules_instruction=modification_term_rules_instruction,
                        )
                        selected_review_feedback = build_selected_review_feedback(
                            st.session_state.get("review_actions", []),
                            st.session_state.get("accepted_review_items", []),
                        )
                        content_to_modify = build_modification_user_content(
                            selected_review_feedback or st.session_state.review_feedback,
                            st.session_state.draft_article,
                            title_candidates=st.session_state.get("title_candidates", []),
                        )

                        final_response = call_llm(
                            api_key=api_key, base_url=current_base_url, model_name=selected_model,
                            system_prompt=modification_prompt, user_content=content_to_modify
                        )

                        final_titles, _, final_article_text = parse_article_generation_response(

                            final_response,

                            st.session_state.get("title_candidates", []),

                        )

                        st.session_state.title_candidates = final_titles

                        st.session_state.final_article = final_article_text
                        append_article_version(st.session_state.final_article, "自动驾驶定稿", role=chosen_editor, model=selected_model)
                        save_draft()

                        if enable_script:
                            st.write("🎬 正在同步生成口播与纯中文分镜脚本...")
                            script_sys_prompt = get_script_sys_prompt(script_duration)
                            st.session_state.spoken_script = call_llm(
                                api_key=api_key, base_url=current_base_url, model_name=selected_model,
                                system_prompt=script_sys_prompt,
                                user_content=f"【请将以下深度文章转化为供剪映AI解析的{script_duration}口播与分镜脚本】：\n\n{st.session_state.final_article}"
                            )
                        else:
                            st.session_state.spoken_script = ""

                        status.update(label="🎉 全自动驾驶完成！即将跳转定稿页。", state="complete", expanded=False)

                    notify_step_completed(defer_until_rerun=True)
                    go_to_step(6)
                    st.rerun()
# --- Step 2 (手动模式) ---
elif st.session_state.current_step == 2:
    render_section_intro("初稿生成", "选择合适的编辑角色，确认当前模型与写作规范，然后输出首版文章。", "Step 02")
    context_strip_placeholder = st.empty()
    
    editor_options = list(prompts_data["editors"].keys())
    if 'selected_role' not in st.session_state or st.session_state.selected_role not in editor_options:
        st.session_state.selected_role = editor_options[0]
        
    if st.session_state.get("selected_role_widget") not in editor_options:
        st.session_state.selected_role_widget = st.session_state.selected_role
    
    st.selectbox(
        "选择【编辑】视角",
        editor_options,
        key="selected_role_widget",
        on_change=sync_selected_role
    )
    editor_role = st.session_state.selected_role
    if st.session_state.get("target_article_words_slider") != get_target_article_words():
        st.session_state.target_article_words_slider = get_target_article_words()
    st.slider(
        "🧮 全局目标字数（200-5000）",
        min_value=200,
        max_value=5000,
        step=100,
        key="target_article_words_slider",
        on_change=sync_target_article_words,
        help="本轮稿件统一使用该字数目标；若角色 Prompt 里有固定字数要求，会自动被全局目标覆盖。"
    )

    render_context_strip([f"当前模型：{selected_model}", f"编辑角色：{st.session_state.selected_role if 'selected_role' in st.session_state else '未选择'}", f"目标字数：约 {get_target_article_words()} 字", f"分镜脚本：{'开启' if enable_script else '关闭'}"])
    
    editor_prompt = st.text_area(
        "✍️ 编辑 Prompt (支持临时微调)", 
        value=prompts_data["editors"][editor_role], 
        height=250,
        key=f"prompt_text_{editor_role}"
    )
    
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("🔙 返回上一步"):
            go_to_step(1)
            st.rerun()
    with col2:
        if st.button(f"🚀 使用 {selected_model} 生成文章初稿"):
            mark_ai_stage_started("draft_generation")
            with st.spinner("编辑正在分析所有素材并奋笔疾书，请耐心等待..."):
                global_instruction = prompts_data.get("global_instruction", "")
                final_editor_system_prompt = build_editor_system_prompt(editor_prompt, global_instruction)
                editor_user_content = build_editor_user_content(
                    st.session_state.source_content,
                    st.session_state.get("obsidian_research_brief", ""),
                    use_images=bool(st.session_state.source_images)
                )

                draft_response = call_llm(
                    api_key=api_key, 
                    base_url=current_base_url,
                    model_name=selected_model, 
                    system_prompt=final_editor_system_prompt,
                    user_content=editor_user_content,
                    image_urls=st.session_state.source_images
                )

                draft_titles, _, draft_article_text = parse_article_generation_response(

                    draft_response,

                    st.session_state.get("title_candidates", []),

                )

                st.session_state.title_candidates = draft_titles

                st.session_state.draft_article = draft_article_text
                append_article_version(st.session_state.draft_article, "手动初稿", role=editor_role, model=selected_model)
                checkpoint_ai_stage("draft_generation", target_step=3)
                save_draft()
                notify_step_completed(defer_until_rerun=True)
                go_to_step(3)
                st.rerun()

# --- Step 3 (手动模式) ---
elif st.session_state.current_step == 3:
    render_section_intro("严格审稿", "主编从事实、逻辑和风格三个维度核查初稿，确保对外可发布。", "Step 03")
    render_context_strip([f"当前模型：{selected_model}", f"编辑角色：{st.session_state.selected_role if 'selected_role' in st.session_state else '未选择'}", f"目标字数：约 {get_target_article_words()} 字", f"分镜脚本：{'开启' if enable_script else '关闭'}"])
    
    with st.expander("📝 查看当前初稿内容 (鼠标移至右上角可一键复制)", expanded=True):
        render_wrapped_article_text(st.session_state.draft_article)
        render_editor_friendly_copy_button(st.session_state.draft_article, "draft_article_step3")
    
    st.divider()
    
    reviewer_prompt = st.text_area("🧐 审稿员 Prompt (支持临时微调)", value=prompts_data["reviewer"], height=200)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("🔄 感觉不对，重写初稿"):
            go_to_step(2)
            st.rerun()
    with col2:
        if st.button("⏭️ 完美，跳过审查进入去AI味"):
            with st.spinner("正在跳过审查，并将初稿送入去 AI 味步骤..."):
                st.session_state.modified_article = st.session_state.draft_article
                st.session_state.final_article = ""
                st.session_state.spoken_script = ""
                append_article_version(st.session_state.modified_article, "跳过审查修改稿", role=st.session_state.get("selected_role", ""), model=selected_model)
                notify_step_completed(defer_until_rerun=True)
                go_to_step(5)
                st.rerun()
    with col3:
        if st.button(f"🔍 使用 {selected_model} 开始严格审查"):
            mark_ai_stage_started("review_generation")
            with st.spinner("主编正在核对原文素材..."):
                if st.session_state.source_images:
                    anti_hallucination_instruction = """\n\n【⚠️ 强制系统级指令：严禁幻觉】：
                    你在审查事实时，**必须且只能**基于下方提供给你的【原始素材文本】以及你所看到的【参考配图】！绝对不允许使用自身知识库进行事实核对。"""
                    combined_content = build_reviewer_user_content(
                        st.session_state.source_content,
                        st.session_state.draft_article,
                        st.session_state.get("obsidian_research_brief", ""),
                        title_candidates=st.session_state.get("title_candidates", []),
                    )
                else:
                    anti_hallucination_instruction = """\n\n【⚠️ 强制系统级指令：严禁幻觉】：
                    你在审查事实时，**必须且只能**基于下方提供给你的【原始素材文本】！绝对不允许使用自身知识库进行事实核对。"""
                    combined_content = build_reviewer_user_content(
                        st.session_state.source_content,
                        st.session_state.draft_article,
                        st.session_state.get("obsidian_research_brief", ""),
                        title_candidates=st.session_state.get("title_candidates", []),
                    )
                
                final_reviewer_system_prompt = build_reviewer_system_prompt(reviewer_prompt, anti_hallucination_instruction)
                
                st.session_state.review_feedback = call_llm(
                    api_key=api_key, 
                    base_url=current_base_url,
                    model_name=selected_model, 
                    system_prompt=final_reviewer_system_prompt, 
                    user_content=combined_content,
                    image_urls=st.session_state.source_images
                )
                hydrate_review_action_state(st.session_state.review_feedback, reset_selection=True)
                checkpoint_ai_stage("review_generation", target_step=4)
                save_draft()
                notify_step_completed(defer_until_rerun=True)
                go_to_step(4)
                st.rerun()

# --- Step 4 (手动模式) ---
elif st.session_state.current_step == 4:
    render_section_intro("\u4fee\u6539\u7a3f\u786e\u8ba4", "\u6839\u636e\u4e3b\u7f16\u53cd\u9988\u5b8c\u6210\u6700\u540e\u4e00\u8f6e\u4fee\u6539\uff0c\u5148\u4ea7\u51fa\u4fee\u6539\u7a3f\uff0c\u518d\u51b3\u5b9a\u662f\u5426\u8fdb\u5165\u53bb AI \u5473\u6b65\u9aa4\u3002", "Step 04")
    render_context_strip([f"\u5f53\u524d\u6a21\u578b\uff1a{selected_model}", f"\u7f16\u8f91\u89d2\u8272\uff1a{st.session_state.selected_role if 'selected_role' in st.session_state else '\u672a\u9009\u62e9'}", f"\u76ee\u6807\u5b57\u6570\uff1a\u7ea6 {get_target_article_words()} \u5b57", f"\u5206\u955c\u811a\u672c\uff1a{'\u5f00\u542f' if enable_script else '\u5173\u95ed'}"])

    if st.session_state.review_feedback and not st.session_state.get("review_actions"):
        hydrate_review_action_state(st.session_state.review_feedback, reset_selection=True)
        save_draft()

    review_actions = st.session_state.get("review_actions", []) or []
    accepted_lookup = set(st.session_state.get("accepted_review_items", []))

    st.info("**\u4e3b\u7f16\u5ba1\u7a3f\u610f\u89c1\uff1a**")
    if review_actions:
        action_ids = [action.get("id") for action in review_actions if action.get("id")]
        accepted_count = sum(1 for action_id in action_ids if action_id in accepted_lookup)
        st.caption(f"\u5df2\u91c7\u7eb3 {accepted_count} / {len(action_ids)} \u6761\u4fee\u6539\u4efb\u52a1\u3002")

        quick_col1, quick_col2, quick_col3 = st.columns(3)
        with quick_col1:
            if st.button("\u5168\u9009\u5ba1\u7a3f\u4efb\u52a1"):
                st.session_state.accepted_review_items = list(action_ids)
                for action_id in action_ids:
                    st.session_state[f"review_action_pick_{action_id}"] = True
                save_draft()
                st.rerun()
        with quick_col2:
            if st.button("\u6e05\u7a7a\u5df2\u91c7\u7eb3\u4efb\u52a1"):
                st.session_state.accepted_review_items = []
                for action_id in action_ids:
                    st.session_state[f"review_action_pick_{action_id}"] = False
                save_draft()
                st.rerun()
        with quick_col3:
            if st.button("\u6062\u590d\u9ed8\u8ba4\u5168\u9009"):
                st.session_state.accepted_review_items = list(action_ids)
                for action_id in action_ids:
                    st.session_state[f"review_action_pick_{action_id}"] = True
                save_draft()
                st.rerun()

        selected_ids = []
        for action in review_actions:
            action_id = action.get("id")
            checkbox_key = f"review_action_pick_{action_id}"
            if checkbox_key not in st.session_state:
                st.session_state[checkbox_key] = action_id in accepted_lookup
            with st.container(border=True):
                checked = st.checkbox(action.get("summary") or action.get("title") or action_id, key=checkbox_key)
                if checked:
                    selected_ids.append(action_id)
                if action.get("body"):
                    st.markdown(action["body"])

        if selected_ids != st.session_state.get("accepted_review_items", []):
            st.session_state.accepted_review_items = selected_ids
            save_draft()

        with st.expander("\u67e5\u770b\u5b8c\u6574\u5ba1\u7a3f\u62a5\u544a"):
            st.code(st.session_state.review_feedback, language="markdown")
            render_editor_friendly_copy_button(st.session_state.review_feedback, "review_feedback_step4")
    else:
        st.caption("\u5f53\u524d\u5ba1\u7a3f\u610f\u89c1\u672a\u80fd\u89e3\u6790\u6210\u53ef\u52fe\u9009\u4efb\u52a1\uff0c\u672c\u6b21\u4fee\u6539\u7a3f\u5c06\u7ee7\u7eed\u4f7f\u7528\u5168\u6587\u6a21\u5f0f\u3002")
        st.code(st.session_state.review_feedback, language="markdown")
        render_editor_friendly_copy_button(st.session_state.review_feedback, "review_feedback_step4")

    if st.session_state.modified_article:
        st.divider()
        st.markdown("### \u5f53\u524d\u4fee\u6539\u7a3f\u9884\u89c8")
        render_wrapped_article_text(st.session_state.modified_article)
        render_editor_friendly_copy_button(st.session_state.modified_article, "modified_article_step4")

    accepted_review_items = st.session_state.get("accepted_review_items", []) or []
    filtered_review_feedback = build_selected_review_feedback(review_actions, accepted_review_items)

    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("\U0001f527 \u610f\u89c1\u592a\u6c34\uff0c\u91cd\u65b0\u5ba1\u67e5"):
            go_to_step(3)
            st.rerun()
    with col2:
        if st.button("\u23ed\ufe0f \u5ffd\u7565\u610f\u89c1\uff0c\u6cbf\u7528\u521d\u7a3f"):
            with st.spinner("\u6b63\u5728\u8df3\u8fc7\u4fee\u6539\uff0c\u51c6\u5907\u8fdb\u5165\u53bb AI \u5473\u6b65\u9aa4..."):
                st.session_state.modified_article = st.session_state.draft_article
                st.session_state.final_article = ""
                st.session_state.highlighted_article = ""
                st.session_state.spoken_script = ""
                append_article_version(st.session_state.modified_article, "\u5ffd\u7565\u610f\u89c1\u4fee\u6539\u7a3f", role=st.session_state.get("selected_role", ""), model=selected_model)
                notify_step_completed(defer_until_rerun=True)
                go_to_step(5)
                st.rerun()
    with col3:
        if st.button(f"\u2728 \u4f7f\u7528 {selected_model} \u63a5\u53d7\u610f\u89c1\u5e76\u751f\u6210\u4fee\u6539\u7a3f"):
            if review_actions and not accepted_review_items:
                st.warning('\u5f53\u524d\u6ca1\u6709\u52fe\u9009\u4efb\u4f55\u5ba1\u7a3f\u4efb\u52a1\u3002\u82e5\u60f3\u5b8c\u5168\u5ffd\u7565\u610f\u89c1\uff0c\u8bf7\u4f7f\u7528"\u5ffd\u7565\u610f\u89c1\uff0c\u6cbf\u7528\u521d\u7a3f"\u3002')
            else:
                mark_ai_stage_started("modification_generation")
                with st.spinner("\u7f16\u8f91\u6b63\u5728\u6839\u636e\u4e3b\u7f16\u610f\u89c1\u751f\u6210\u4fee\u6539\u7a3f..."):
                    global_instruction = prompts_data.get("global_instruction", "")
                    modification_banned_terms, modification_default_replacements, _, _ = resolve_active_term_rules("modification")
                    modification_term_rules_instruction = build_term_rules_instruction(
                        modification_banned_terms,
                        modification_default_replacements,
                    )
                    modification_prompt = build_modification_system_prompt(
                        global_instruction,
                        term_rules_instruction=modification_term_rules_instruction,
                    )

                    content_to_modify = build_modification_user_content(
                        filtered_review_feedback or st.session_state.review_feedback,
                        st.session_state.draft_article,
                        title_candidates=st.session_state.get("title_candidates", []),
                    )

                    modified_response = call_llm(
                        api_key=api_key,
                        base_url=current_base_url,
                        model_name=selected_model,
                        system_prompt=modification_prompt,
                        user_content=content_to_modify
                    )

                    modified_titles, _, modified_article_text = parse_article_generation_response(

                        modified_response,

                        st.session_state.get("title_candidates", []),

                    )

                    st.session_state.title_candidates = modified_titles

                    st.session_state.modified_article = modified_article_text
                    st.session_state.final_article = ""
                    st.session_state.highlighted_article = ""
                    st.session_state.spoken_script = ""
                    append_article_version(st.session_state.modified_article, "\u63a5\u53d7\u5ba1\u7a3f\u4fee\u6539\u7a3f", role=st.session_state.get("selected_role", ""), model=selected_model)
                    checkpoint_ai_stage("modification_generation", target_step=5)
                    save_draft()
                    notify_step_completed(defer_until_rerun=True)
                    go_to_step(5)
                    st.rerun()

# --- Step 5 (manual mode) ---
elif st.session_state.current_step == 5:
    current_role = st.session_state.get("selected_role", "")
    current_editor_prompt = prompts_data["editors"].get(current_role, "") if current_role in prompts_data["editors"] else ""
    de_ai_variant = st.session_state.get("de_ai_variant", DE_AI_VARIANT_DEFAULT)
    de_ai_button_suffix = (
        "（社区版）"
        if de_ai_variant == DE_AI_VARIANT_COMMUNITY
        else "（唠嗑版）"
        if de_ai_variant == DE_AI_VARIANT_CHAT
        else "（Humanizer版）"
        if de_ai_variant == DE_AI_VARIANT_HUMANIZER
        else ""
    )

    render_section_intro("去 AI 味", "在定稿前选择专用模型，把修改稿重写得更像真人专家输出。", "Step 05")
    render_context_strip([
        f"修改稿来源模型：{selected_model}",
        f"专用模型：{st.session_state.get('de_ai_model', DE_AI_MODELS[0])}",
        f"风格版本：{de_ai_variant}",
        f"Temperature：{st.session_state.get('de_ai_temperature', 0.75):.2f}",
        f"分镜脚本：{'开启' if enable_script else '关闭'}",
    ])

    st.markdown("### 当前修改稿")
    render_wrapped_article_text(st.session_state.modified_article)
    render_editor_friendly_copy_button(st.session_state.modified_article, "modified_article_step5")

    col_variant, col_model, col_temp = st.columns([1.15, 1.2, 1])
    with col_variant:
        st.selectbox(
            "去 AI 风格版本",
            DE_AI_VARIANTS,
            key="de_ai_variant",
            on_change=save_draft,
        )
    with col_model:
        st.selectbox(
            "去 AI 味专用大模型",
            DE_AI_MODELS,
            key="de_ai_model",
            on_change=save_draft,
        )
    with col_temp:
        st.slider(
            "Temperature",
            min_value=0.70,
            max_value=0.85,
            step=0.05,
            key="de_ai_temperature",
            on_change=save_draft,
        )

    with st.expander("本文临时词表微调", expanded=False):
        st.caption("这里适合写这篇稿子临时追加的禁用词、默认替换词和建议替换词。全局词表会一起合并显示。")
        st.text_area(
            "本文临时禁用词（每行一个）",
            key="article_banned_terms_text",
            height=110,
            placeholder="进一步来说\n值得一提的是",
            on_change=save_draft,
        )
        st.text_area(
            "本文临时默认替换词（每行一条：原词 => 新词）",
            key="article_default_replacement_terms_text",
            height=120,
            placeholder="值得关注 => 更值得细看的是\n进一步来说 => 往下看",
            on_change=save_draft,
        )
        st.text_area(
            "本文临时建议替换词（每行一条：原词 => 新词）",
            key="article_suggested_replacement_terms_text",
            height=120,
            placeholder="核心优势在于 => 真正有优势的地方在于\n显而易见 => 这点其实很清楚",
            on_change=save_draft,
        )
        _, preview_default_invalid_lines = parse_replacement_terms_text(st.session_state.get("article_default_replacement_terms_text", ""))
        _, preview_suggested_invalid_lines = parse_replacement_terms_text(st.session_state.get("article_suggested_replacement_terms_text", ""))
        if preview_default_invalid_lines:
            st.warning("以下本文临时默认替换词格式无法识别：" + "；".join(preview_default_invalid_lines[:6]))
        if preview_suggested_invalid_lines:
            st.warning("以下本文临时建议替换词格式无法识别：" + "；".join(preview_suggested_invalid_lines[:6]))
        preview_banned_terms, preview_default_replacements, preview_suggested_replacements, _ = resolve_active_term_rules(respect_enabled=False)
        preview_term_instruction = build_term_rules_preview_text(
            preview_banned_terms,
            preview_default_replacements,
            preview_suggested_replacements,
        )
        if preview_term_instruction:
            st.code(preview_term_instruction, language="markdown")
        else:
            st.caption("当前这篇稿子还没有生效的词表规则。")

    de_ai_variant = st.session_state.get("de_ai_variant", DE_AI_VARIANT_DEFAULT)
    de_ai_button_suffix = (
        "（社区版）"
        if de_ai_variant == DE_AI_VARIANT_COMMUNITY
        else "（唠嗑版）"
        if de_ai_variant == DE_AI_VARIANT_CHAT
        else "（Humanizer版）"
        if de_ai_variant == DE_AI_VARIANT_HUMANIZER
        else ""
    )
    de_ai_banned_terms, de_ai_default_replacements, _, _ = resolve_active_term_rules("de_ai")
    de_ai_term_rules_instruction = build_term_rules_instruction(
        de_ai_banned_terms,
        de_ai_default_replacements,
    )
    st.session_state.de_ai_prompt_template = build_de_ai_prompt_template(
        current_role,
        current_editor_prompt,
        st.session_state.get("source_content", ""),
        variant=de_ai_variant,
        term_rules_instruction=de_ai_term_rules_instruction,
    )

    with st.expander("查看本次去 AI 味专用 Prompt 模板（只读）", expanded=False):
        st.text_area(
            "去 AI 味 Prompt 模板",
            value=st.session_state.de_ai_prompt_template,
            height=360,
            disabled=True,
            key="de_ai_prompt_template_preview",
        )

    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("返回修改稿步骤"):
            go_to_step(4)
            st.rerun()
    with col2:
        if st.button("跳过去 AI 味，直接定稿"):
            spinner_msg = f"正在将修改稿设为定稿，并生成【{script_duration}口播及分镜脚本】..." if enable_script else "正在将修改稿设为定稿..."
            with st.spinner(spinner_msg):
                skip_titles, _, skip_article_text = parse_article_generation_response(
                    st.session_state.modified_article,
                    st.session_state.get("title_candidates", []),
                )
                st.session_state.title_candidates = skip_titles
                st.session_state.final_article = skip_article_text
                st.session_state.highlighted_article = ""
                append_article_version(st.session_state.final_article, "跳过去AI味定稿", role=current_role, model=selected_model)
                if enable_script:
                    generate_script_for_current_article(api_key, current_base_url, selected_model, script_duration)
                else:
                    st.session_state.spoken_script = ""

                if st.session_state.get("podcast_enabled"):
                    generate_podcast_script_for_current_article(
                        api_key,
                        current_base_url,
                        selected_model,
                        st.session_state.get("podcast_duration", "5分钟"),
                    )
                else:
                    reset_podcast_outputs(delete_audio=True)
                save_draft()
                notify_step_completed(defer_until_rerun=True)
                go_to_step(6)
                st.rerun()
    with col3:
        if st.button(f"使用 {st.session_state.get('de_ai_model', DE_AI_MODELS[0])} 去 AI 味重写{de_ai_button_suffix}"):
            mark_ai_stage_started("de_ai_generation")
            spinner_msg = (
                f"正在使用 {st.session_state.get('de_ai_model', DE_AI_MODELS[0])} 去 AI 味{de_ai_button_suffix}，并生成【{script_duration}口播及分镜脚本】..."
                if enable_script
                else f"正在使用 {st.session_state.get('de_ai_model', DE_AI_MODELS[0])} 去 AI 味重写{de_ai_button_suffix}..."
            )
            with st.spinner(spinner_msg):
                de_ai_response = call_llm(
                    api_key=api_key,
                    base_url=current_base_url,
                    model_name=st.session_state.get('de_ai_model', DE_AI_MODELS[0]),
                    system_prompt=st.session_state.de_ai_prompt_template,
                    user_content=st.session_state.modified_article,
                    temperature=st.session_state.get('de_ai_temperature', 0.75),
                )
                pure_titles, pure_article, highlighted_article = parse_de_ai_dual_output(
                    de_ai_response,
                    fallback_titles=st.session_state.get("title_candidates", []),
                )
                st.session_state.title_candidates = pure_titles
                st.session_state.final_article = build_structured_article_text(pure_titles, pure_article) or (de_ai_response or "").strip()
                st.session_state.highlighted_article = highlighted_article
                if de_ai_variant == DE_AI_VARIANT_COMMUNITY:
                    de_ai_stage_label = "去AI味定稿（社区版）"
                elif de_ai_variant == DE_AI_VARIANT_CHAT:
                    de_ai_stage_label = "去AI味定稿（唠嗑版）"
                elif de_ai_variant == DE_AI_VARIANT_HUMANIZER:
                    de_ai_stage_label = "去AI味定稿（Humanizer版）"
                else:
                    de_ai_stage_label = "去AI味定稿"
                append_article_version(
                    st.session_state.final_article,
                    de_ai_stage_label,
                    role=current_role,
                    model=st.session_state.get('de_ai_model', DE_AI_MODELS[0]),
                    highlighted_article=highlighted_article,
                )
                save_draft()
                if enable_script:
                    generate_script_for_current_article(api_key, current_base_url, selected_model, script_duration)
                else:
                    st.session_state.spoken_script = ""

                if st.session_state.get("podcast_enabled"):
                    generate_podcast_script_for_current_article(
                        api_key,
                        current_base_url,
                        selected_model,
                        st.session_state.get("podcast_duration", "5分钟"),
                    )
                else:
                    reset_podcast_outputs(delete_audio=True)
                if not highlighted_article:
                    st.warning("高亮版生成失败，本次仅保留纯净定稿。")
                checkpoint_ai_stage("de_ai_generation", target_step=6)
                save_draft()
                notify_step_completed(defer_until_rerun=True)
                go_to_step(6)
                st.rerun()
# --- Step 6：终极版分栏 UI ---
# --- Step 6：终极版分栏 UI ---
# --- Step 6：终极版分栏 UI ---

elif st.session_state.current_step == 6:
    refresh_obsidian_influence_map()
    refresh_evidence_map()
    render_section_intro("分发工作台", "在统一界面完成定稿审阅、脚本联动、搜图建议、导出分发和后续精修。", "Step 06")
    display_final_article = build_display_article_text(
        st.session_state.get("final_article", ""),
        st.session_state.get("title_candidates", []),
    )
    final_check_banned_terms, final_check_default_replacements, final_check_suggested_replacements, final_check_invalid_lines = resolve_active_term_rules("final_check")
    if final_check_banned_terms or final_check_default_replacements or final_check_suggested_replacements:
        st.session_state.term_scan_result = scan_article_terms(
            st.session_state.get("final_article", ""),
            final_check_banned_terms,
            final_check_default_replacements,
            final_check_suggested_replacements,
        )
        st.session_state.term_scan_summary = summarize_term_scan(st.session_state.get("term_scan_result", []))
    else:
        st.session_state.term_scan_result = []
        st.session_state.term_scan_summary = summarize_term_scan([])

    publish_quality_gate = build_publish_quality_gate_report(
        display_final_article,
        title_candidates=st.session_state.get("title_candidates", []),
        highlighted_article=st.session_state.get("highlighted_article", ""),
        term_scan_summary=st.session_state.get("term_scan_summary", {}),
        target_words=globals().get("get_target_article_words", lambda: 1500)(),
    )

    render_context_strip([
        f"最终角色：{st.session_state.selected_role if 'selected_role' in st.session_state else '自动路由'}",
        f"当前模型：{selected_model}",
        f"脚本状态：{'已生成' if st.session_state.spoken_script else '未生成'}",
    ])
    st.markdown("<p class='toolbar-note'>左侧保持主稿、复制、导出与分发链路；右侧上方用于观察 Obsidian 影响，下方保留精修对话。</p>", unsafe_allow_html=True)
    left_col, right_col = st.columns([1.45, 0.98])
    with left_col:
        st.markdown("### 稿件版本时间线")
        versions = st.session_state.get("article_versions", [])
        if versions:
            version_map = {item.get("id"): item for item in versions if item.get("id")}
            version_options = [item.get("id") for item in reversed(versions) if item.get("id")]
            active_version_id = st.session_state.get("active_article_version_id")
            default_version_index = version_options.index(active_version_id) if active_version_id in version_options else 0

            def format_version_option(version_id):
                version_item = version_map.get(version_id, {})
                return f"{version_id} · {version_item.get('stage', '未命名阶段')} · {version_item.get('created_at', '')} · {version_item.get('word_count', 0)} 字"

            selected_version_id = st.selectbox(
                "选择版本节点",
                options=version_options,
                index=default_version_index,
                format_func=format_version_option,
                key="article_version_selector",
            )
            selected_version = version_map.get(selected_version_id)

            if selected_version_id != st.session_state.get("active_article_version_id"):
                st.session_state.active_article_version_id = selected_version_id
                save_draft()

            if selected_version:
                st.caption(f"当前阶段：{selected_version.get('stage', '未命名')}｜角色：{selected_version.get('role', '未记录')}｜模型：{selected_version.get('model', '未记录')}")

                if st.button("将此版本设为当前定稿", key="use_selected_version_as_final", use_container_width=True):
                    restore_article_version_to_session(
                        selected_version,
                        fallback_titles=st.session_state.get("title_candidates", []),
                    )
                    save_draft()
                    notify_step_completed()
                    st.success(f"已切换到版本 {selected_version_id}")
                with st.expander("查看选中版本全文", expanded=False):
                    render_wrapped_article_text(selected_version.get("content", ""))
                    render_editor_friendly_copy_button(selected_version.get("content", ""), f"version_{selected_version_id}")
        else:
            st.info("暂无版本记录。生成初稿或定稿后会自动写入版本时间线。")

        st.divider()
        st.markdown("### 主稿面板（当前定稿）")
        render_wrapped_article_text(display_final_article)
        render_editor_friendly_copy_button(display_final_article, "final_article_step6")

        st.divider()
        st.markdown("### 发布前质量闸门")
        with st.container(border=True):
            if publish_quality_gate.get("overall_status") == "pass":
                st.success("当前稿件已通过发布前结构闸门。")
            elif publish_quality_gate.get("overall_status") == "fail":
                st.error("当前稿件还没过发布前闸门，建议先修完红项再发。")
            else:
                st.warning("当前稿件基本可用，但还有一些建议优化项。")

            expected_range = publish_quality_gate.get("expected_h2_range", (3, 5))
            st.caption(
                f"红项 {publish_quality_gate.get('fail_count', 0)} 个｜黄项 {publish_quality_gate.get('warn_count', 0)} 个｜当前 `##` 数量 {publish_quality_gate.get('h2_count', 0)}｜建议范围 {expected_range[0]}-{expected_range[1]}"
            )
            status_label_map = {
                "pass": "通过",
                "warn": "提醒",
                "fail": "拦截",
            }
            for item in publish_quality_gate.get("items", []):
                st.markdown(
                    f"- **[{status_label_map.get(item.get('status'), '提示')}] {item.get('label', '')}**：{item.get('detail', '')}"
                )

        st.divider()
        st.markdown("### 词表检查")
        with st.container(border=True):
            term_rule_scope = st.session_state.get("term_rules_scope", TERM_RULE_SCOPE_DEFAULT)
            if not st.session_state.get("term_rules_enabled", False):
                st.caption("当前未启用词表规则。开启后，系统会在修改稿 / 去 AI 阶段提前施压，并在最终稿阶段自动检查。")
            elif "final_check" not in term_rule_scope:
                st.caption("当前没有对最终稿启用词表检查。你可以在侧边栏的【个人写作词表】里打开这个范围。")
            elif final_check_invalid_lines:
                st.warning("以下替换词格式无法识别：" + "；".join(final_check_invalid_lines[:6]))
            elif not final_check_banned_terms and not final_check_default_replacements and not final_check_suggested_replacements:
                st.caption("当前没有生效的禁用词、默认替换词或建议替换词。")
            elif st.session_state.get("term_scan_result"):
                term_scan_summary = st.session_state.get("term_scan_summary", {})
                st.warning(
                    f"命中 {term_scan_summary.get('matched_terms', 0)} 个词表规则，共 {term_scan_summary.get('total_hits', 0)} 处；"
                    f"其中禁用词 {term_scan_summary.get('banned_terms', 0)} 个，默认替换 {term_scan_summary.get('default_replacement_terms', 0)} 个，建议替换 {term_scan_summary.get('suggested_replacement_terms', 0)} 个。"
                )
                st.caption("段落编号按纯净正文计算，不包含标题组。")
                term_scan_result = st.session_state.get("term_scan_result", [])
                banned_hits = [item for item in term_scan_result if item.get("type") == "banned"]
                default_hits = [
                    item for item in term_scan_result
                    if item.get("type") == "replacement" and item.get("level") == "default"
                ]
                suggested_hits = [
                    item for item in term_scan_result
                    if item.get("type") == "replacement" and item.get("level") == "suggested"
                ]

                def render_term_hit_group(title, items, empty_text):
                    st.markdown(f"**{title}**")
                    if not items:
                        st.caption(empty_text)
                        return
                    for item in items:
                        paragraph_text = "、".join([f"第{index}段" for index in item.get("paragraph_indexes", [])]) or "正文未知位置"
                        replacement_text = item.get("replacement") or "无默认替换建议"
                        st.markdown(
                            f"- `{item.get('term', '')}` · 命中 **{item.get('count', 0)}** 次 · {paragraph_text} · 建议：{replacement_text}"
                        )

                render_term_hit_group("禁用词命中", banned_hits, "当前没有命中禁用词。")
                render_term_hit_group("默认替换命中", default_hits, "当前没有命中默认替换词。")
                render_term_hit_group("建议替换命中", suggested_hits, "当前没有命中建议替换词。")
            else:
                st.success("这篇最终稿没有命中当前词表规则。")
        st.divider()
        st.markdown("### 高亮阅读版")
        render_highlighted_article_panel(st.session_state.get("highlighted_article", ""))

        if st.session_state.spoken_script:
            st.divider()
            st.markdown(f"### 分镜脚本 · {script_duration}")
            st.code(st.session_state.spoken_script, language="markdown")
            render_editor_friendly_copy_button(st.session_state.spoken_script, "spoken_script_step6")
        if st.session_state.get("podcast_enabled"):
            st.divider()
            st.markdown(f"### 播客解说稿 · {st.session_state.get('podcast_duration', '5分钟')}")
            podcast_script_segments = st.session_state.get("podcast_script_segments", []) or []
            podcast_script_json = json.dumps(podcast_script_segments, ensure_ascii=False, indent=2) if podcast_script_segments else ""

            podcast_btn_col1, podcast_btn_col2, podcast_btn_col3 = st.columns(3)
            with podcast_btn_col1:
                if st.button("重生成播客稿", use_container_width=True):
                    with st.spinner("正在生成单人播客解说稿..."):
                        success = generate_podcast_script_for_current_article(
                            api_key,
                            current_base_url,
                            selected_model,
                            st.session_state.get("podcast_duration", "5分钟"),
                        )
                        if success:
                            st.success("播客解说稿已更新。")
                        else:
                            st.error(st.session_state.get("podcast_last_error", "播客解说稿生成失败。"))
                        st.rerun()
            with podcast_btn_col2:
                synthesize_clicked = st.button("合成播客音频", use_container_width=True)
            with podcast_btn_col3:
                if st.button("清除播客音频", use_container_width=True):
                    audio_path = st.session_state.get("podcast_audio_path", "")
                    if audio_path and os.path.exists(audio_path):
                        try:
                            os.remove(audio_path)
                        except OSError:
                            pass
                    st.session_state.podcast_audio_path = ""
                    st.session_state.podcast_audio_manifest = {}
                    st.session_state.podcast_last_error = ""
                    save_draft()
                    st.rerun()

            if podcast_script_json:
                st.code(podcast_script_json, language="json")
                render_editor_friendly_copy_button(podcast_script_json, "podcast_script_step6")
            else:
                st.info("当前还没有播客解说稿。你可以先点击“重生成播客稿”。")

            if synthesize_clicked:
                if not podcast_script_segments:
                    st.error("请先生成播客解说稿，再进行音频合成。")
                elif not (podcast_api_key or "").strip():
                    st.error("请先在侧边栏填写阿里云 DashScope API Key。")
                else:
                    with st.spinner("正在合成播客音频..."):
                        try:
                            ensure_podcast_runtime_dirs()
                            output_path, manifest = synthesize_podcast(
                                segments=podcast_script_segments,
                                voice=st.session_state.get("podcast_voice", DEFAULT_TTS_VOICE),
                                api_key=podcast_api_key.strip(),
                                output_dir=PODCAST_OUTPUT_DIR,
                                cache_dir=PODCAST_CACHE_DIR,
                                model=DEFAULT_TTS_MODEL,
                            )
                            st.session_state.podcast_audio_path = output_path
                            st.session_state.podcast_audio_manifest = manifest
                            st.session_state.podcast_last_error = ""
                            save_draft()
                            st.success("播客音频已生成。")
                            st.rerun()
                        except PodcastAudioError as exc:
                            st.session_state.podcast_last_error = str(exc)
                            save_draft()
                            st.error(str(exc))

            podcast_error = st.session_state.get("podcast_last_error", "")
            if podcast_error:
                st.error(podcast_error)

            podcast_audio_path = st.session_state.get("podcast_audio_path", "")
            if podcast_audio_path and os.path.exists(podcast_audio_path):
                with open(podcast_audio_path, "rb") as podcast_audio_file:
                    podcast_audio_bytes = podcast_audio_file.read()
                st.audio(podcast_audio_bytes, format="audio/mp3")
                st.download_button(
                    label="下载播客 MP3",
                    data=podcast_audio_bytes,
                    file_name=os.path.basename(podcast_audio_path),
                    mime="audio/mpeg",
                    use_container_width=True,
                )
                manifest = st.session_state.get("podcast_audio_manifest", {}) or {}
                if manifest:
                    voice = manifest.get("voice", "")
                    st.caption(
                        f"音色：{PODCAST_VOICE_LABELS.get(voice, voice)}｜"
                        f"片段数：{manifest.get('segment_count', 0)}｜缓存命中：{manifest.get('cache_hits', 0)}"
                    )

        st.divider()
        st.markdown("### 智能配图助手")
        st.info("需要为文章寻找真实、高质的配图？点击下方按钮，AI 将根据文章核心内容提取 10 个精确的 Google 图片搜索关键词。")

        if st.button("💡 提取 10 个 Google 搜图关键词", use_container_width=True):
            with st.spinner("正在深度分析文章，提取精确的搜图词汇..."):
                keyword_prompt = """请根据以下文章内容，提取 10 个最适合在 Google 图片（Google Images）中搜索配图的精准关键词组合。
                
                【核心要求】：
                1. 必须精准输出 10 个。
                2. 为了在 Google 搜出最高质量的图，请尽量采用【中英文混合】或【纯英文】的专业搜索词（例如："Genshin Impact UI design", "Sensor Tower SLG revenue chart 2024", "Tencent Games logo transparent"）。
                3. 场景具体化：不要只搜游戏名，要加上明确的修饰词（如实机演示、数据图表、买量素材、应用商店截图等）。
                4. 直接用编号 1-10 列出，不要有任何废话解释。
                
                【文章定稿内容】：
                """ + get_article_body_text(st.session_state.final_article)

                st.session_state.image_keywords = call_llm(
                    api_key=api_key,
                    base_url=current_base_url,
                    model_name=selected_model,
                    system_prompt="你是一个专业的游戏媒体视觉编辑，熟知如何通过高级检索词在 Google 找到极具说服力的行业配图。",
                    user_content=keyword_prompt
                )
                save_draft()
                notify_step_completed()

        if st.session_state.image_keywords:
            st.success("✅ 关键词提取成功！你可以直接复制这些词去 Google 搜图：")
            st.code(st.session_state.image_keywords, language="markdown")
            render_editor_friendly_copy_button(st.session_state.image_keywords, "image_keywords_step6")
            
        st.divider()

        def create_docx(article_text, script_text=None):
            doc = Document()
            doc.add_heading('【最终成稿】', level=1)
            doc.add_paragraph(article_text)
            
            if script_text:
                doc.add_heading('【短视频 AI 分镜脚本】', level=1)
                doc.add_paragraph(script_text)
            
            bio = io.BytesIO()
            doc.save(bio)
            return bio.getvalue()
            
        docx_data = create_docx(display_final_article, st.session_state.spoken_script if st.session_state.spoken_script else None)
        
        btn_col1, btn_col2, btn_col3 = st.columns(3)
        with btn_col1:
             st.download_button(
                label="📄 导出 Word 文档" if not st.session_state.spoken_script else "📄 导出图文与脚本(Word)",
                data=docx_data,
                file_name="公众号文章_定稿.docx" if not st.session_state.spoken_script else "公众号与短视频脚本_定稿.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
        with btn_col2:
            if st.button("✈️ 推送通知到飞书群", use_container_width=True):
                with st.spinner("正在推送到飞书..."):
                    success, msg = push_to_feishu(display_final_article, st.session_state.spoken_script if st.session_state.spoken_script else None)
                    if success:
                        st.success("🎉 飞书推送成功！")
                        notify_step_completed()
                    else:
                        st.error(f"❌ 推送失败：{msg}")
                            
        with btn_col3:
            if st.button("🔄 开启新一篇工作流", use_container_width=True):
                clear_draft()
                create_task_from_current_state(clone_current=False)
                st.rerun()


    
    with right_col:
        should_show_evidence_map = bool(
            (st.session_state.get("final_article", "") or "").strip()
            and (st.session_state.get("source_content", "") or "").strip()
        )
        should_show_obsidian_influence = bool(
            st.session_state.get("obsidian_enabled")
            and (st.session_state.get("final_article", "") or "").strip()
            and (st.session_state.get("obsidian_hits", []) or [])
        )

        if should_show_evidence_map:
            with st.container(border=True):
                render_evidence_map_panel(
                    st.session_state.get("evidence_map", []),
                    st.session_state.get("evidence_summary", ""),
                )

        if should_show_obsidian_influence:
            with st.container(border=True):
                render_obsidian_influence_panel(
                    st.session_state.get("obsidian_influence_map", []),
                    st.session_state.get("obsidian_influence_summary", ""),
                )

        with st.container(border=True):
            render_section_intro("精修对话区", "保留消息历史与输入框，用于继续追问出处或改写段落。", "对话")
            chat_container = st.container(height=360)
    
            with chat_container:
                for msg in st.session_state.chat_history:
                    with st.chat_message(msg["role"]):
                        st.markdown(msg["content"])
    
            if user_query := st.chat_input("输入修改指令、追问出处或改写要求，回车发送..."):
                st.session_state.chat_history.append({"role": "user", "content": user_query})
                with chat_container:
                    with st.chat_message("user"):
                        st.markdown(user_query)
    
                    with st.chat_message("assistant"):
                        with st.spinner("思考中..."):
                            knowledge_context = build_chat_knowledge_context()
                            chat_sys_prompt = f"""You are an expert article refinement and source-tracing assistant.
    
                            [Reference source library - the only factual source for tracing]
                            {st.session_state.source_content}
    
                            [Current final article]
                            {display_final_article}
    
                            [Background knowledge library]
                            {knowledge_context}
    
                            [Tasks]
                            1. If the user asks for sourcing, only use the reference source library to locate the original passage. Do not treat the knowledge library as a news-fact source.
                            2. If the user asks for a rewrite, output the replacement paragraph directly with no extra chatter.
                            3. If the user asks for derived analysis, combine the article and the background knowledge carefully. Mention the note title first when you rely on the knowledge library.
                            """
    
                            history_to_send = st.session_state.chat_history[:-1]
    
                            ai_response = call_llm(
                                api_key=api_key,
                                base_url=current_base_url,
                                model_name=selected_model,
                                system_prompt=chat_sys_prompt,
                                user_content=user_query,
                                history=history_to_send
                            )
                            st.markdown(ai_response)
    
                st.session_state.chat_history.append({"role": "assistant", "content": ai_response})
                save_draft()
                st.rerun()






































































